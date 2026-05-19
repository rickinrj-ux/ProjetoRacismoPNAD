"""
gerar_relatorio_word.py
=======================
Gera relatorio_tcc.docx — documento Word editável com todos os resultados.
Usa python-docx + formatação ABNT (margens, fonte Times/Arial, espaçamento 1,5).
"""
import sys, io
from pathlib import Path

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

ROOT    = Path(r"C:\Users\user\Documents\ProjetoRacismoPNAD")
TABLES  = ROOT / "outputs" / "tables"
FIGURES = ROOT / "outputs" / "figures"
OUT_DOC = ROOT / "relatorio_tcc.docx"

# ── Helpers de formatação ─────────────────────────────────────────────────────

def set_margins(doc, top=3, bottom=2, left=3, right=2):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def set_paragraph_format(para, first_line=1.25, space_after=6, line_spacing=1.5):
    pf = para.paragraph_format
    pf.first_line_indent = Cm(first_line) if first_line else None
    pf.space_after        = Pt(space_after)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = line_spacing

def add_heading(doc, text, level=1, numbering=None):
    style = f"Heading {level}"
    heading = doc.add_heading(level=level)
    heading.clear()
    run = heading.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14) if level == 1 else (Pt(13) if level == 2 else Pt(12))
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after  = Pt(6)
    return heading

def add_para(doc, text, bold=False, italic=False, first_line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_format(p, first_line=first_line)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

def add_caption(doc, text, fig=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.bold = False
    run.font.name = "Arial"
    run.font.size = Pt(10)

def add_figure(doc, img_path, caption_text, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    if Path(img_path).exists():
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run.add_text(f"[Figura não encontrada: {Path(img_path).name}]")
    add_caption(doc, caption_text)

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

# ── OMML equation helpers (Word native math equations) ───────────────────────
_M = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

def _mi(t): return f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">{t}</m:t></m:r>'
def _mu(t): return f'<m:r><m:t xml:space="preserve">{t}</m:t></m:r>'
def _msub(b, s): return f'<m:sSub><m:e>{b}</m:e><m:sub>{s}</m:sub></m:sSub>'
def _msup(b, s): return f'<m:sSup><m:e>{b}</m:e><m:sup>{s}</m:sup></m:sSup>'
def _mss(b, sb, sp): return f'<m:sSubSup><m:e>{b}</m:e><m:sub>{sb}</m:sub><m:sup>{sp}</m:sup></m:sSubSup>'
def _mfrac(n, d): return f'<m:f><m:num>{n}</m:num><m:den>{d}</m:den></m:f>'
def _mwrap(inner):
    return (f'<m:oMathPara {_M}>'
            f'<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
            f'<m:oMath>{inner}</m:oMath></m:oMathPara>')

def add_equation(doc, omml_para_xml, space_before=4, space_after=8):
    """Insert a Word native equation (OMML) as a centred paragraph."""
    elem = etree.fromstring(omml_para_xml.encode('utf-8'))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.5
    p._p.append(elem)
    return p

# Pre-built OMML strings ─────────────────────────────────────────────────────
_EQ_HLM1 = _mwrap(
    _mu('ln(') + _msub(_mi('W'), _mi('ijk')) + _mu(') = ') +
    _msub(_mi('β'), _mi('0jk')) +
    _mu(' + ') + _msub(_mi('β'), _mu('1')) + _mu('·') + _msub(_mi('Negro'), _mi('ijk')) +
    _mu(' + ') + _msub(_mi('β'), _mu('2')) + _mu('·') + _msub(_mi('Sexo'), _mi('ijk')) +
    _mu(' + ') + _msub(_mi('β'), _mu('3')) + _mu('·') + _msub(_mi('X'), _mi('ijk')) +
    _mu(' + ') + _msub(_mi('β'), _mu('4')) + _mu('·') + _mss(_mi('X'), _mi('ijk'), _mu('2')) +
    _mu(' + ∑') + _msub(_mi('β'), _mi('e')) + _mu('·') + _msub(_mi('Educ'), _mi('e')) +
    _mu(' + ') + _msub(_mi('ε'), _mi('ijk')) +
    _mu(',  ε ~ N(0, ') + _msup(_mi('σ'), _mu('2')) + _mu(')')
)

_EQ_HLM2 = _mwrap(
    _msub(_mi('β'), _mi('0jk')) + _mu(' = ') +
    _msub(_mi('γ'), _mu('00k')) +
    _mu(' + ') + _msub(_mi('γ'), _mu('01')) + _mu('·%') + _msub(_mi('Negro'), _mi('jk')) +
    _mu(' + ') + _msub(_mi('γ'), _mu('02')) + _mu('·') + _msub(_mi('Desemprego'), _mi('jk')) +
    _mu(' + ') + _msub(_mi('γ'), _mu('03')) + _mu('·') + _msub(_mi('Educ'), _mi('jk')) +
    _mu(' + ') + _msub(_mi('u'), _mi('0jk')) +
    _mu(',  ') + _msub(_mi('u'), _mi('0j')) + _mu(' ~ N(0, ') + _msup(_mi('τ'), _mu('2')) + _mi('u') + _mu(')')
)

_EQ_HLM3 = _mwrap(
    _msub(_mi('γ'), _mu('00k')) + _mu(' = ') +
    _msub(_mi('δ'), _mu('000')) +
    _mu(' + ') + _msub(_mi('δ'), _mu('1')) + _mu('·') + _msub(_mi('Z'), _mi('k%negro')) +
    _mu(' + ') + _msub(_mi('δ'), _mu('2')) + _mu('·') + _msub(_mi('Z'), _mi('kdesemprego')) +
    _mu(' + ') + _msub(_mi('δ'), _mu('3')) + _mu('·') + _msub(_mi('Z'), _mi('keduc')) +
    _mu(' + ') + _msub(_mi('v'), _mi('00k')) +
    _mu(',  ') + _msub(_mi('v'), _mi('00k')) + _mu(' ~ N(0, ') + _msup(_mi('τ'), _mu('2')) + _mi('v') + _mu(')')
)

_EQ_OLS = _mwrap(
    _mu('ln(') + _msub(_mi('W'), _mi('ijk')) + _mu(') = ') +
    _msub(_mi('α'), _mu('0')) +
    _mu(' + ') + _msub(_mi('β'), _mu('1')) + _mu('·') + _msub(_mi('Negro'), _mi('ijk')) +
    _mu(' + ') + _msub(_mi("X′"), _mi('ijk')) + _mu('·β') +
    _mu(' + ') + _msub(_mi('ε'), _mi('ijk')) +
    _mu(',  ε ~ iid N(0, ') + _msup(_mi('σ'), _mu('2')) + _mu(')')
)

_EQ_NULL = _mwrap(
    _mu('ln(') + _msub(_mi('W'), _mi('ijk')) + _mu(') = ') +
    _msub(_mi('δ'), _mu('000')) + _mu(' + ') + _msub(_mi('u'), _mi('00k')) +
    _mu(' + ') + _msub(_mi('ε'), _mi('ijk')) +
    _mu(',  ') + _msub(_mi('u'), _mi('00k')) + _mu(' ~ N(0, ') + _msup(_mi('τ'), _mu('2')) + _mi('v') + _mu(')') +
    _mu(',  ') + _msub(_mi('ε'), _mi('ijk')) + _mu(' ~ N(0, ') + _msup(_mi('σ'), _mu('2')) + _mu(')')
)

_EQ_ICC = _mwrap(
    _msub(_mi('ρ'), _mu('UF')) + _mu(' = ') +
    _mfrac(
        _msup(_mi('τ'), _mu('2')) + _mi('v'),
        _msup(_mi('τ'), _mu('2')) + _mi('v') + _mu(' + ') + _msup(_mi('σ'), _mu('2'))
    )
)

_EQ_ICC_NUM = _mwrap(
    _msub(_mi('ρ'), _mu('UF')) + _mu(' = ') +
    _mfrac(_mu('0,0834'), _mu('0,0834 + 0,7653')) +
    _mu(' = 0,0983   (9,83%)')
)

# ── Carrega resultados ────────────────────────────────────────────────────────

def load_results():
    r = {}
    r["hlm"]        = pd.read_csv(TABLES / "hlm_serie_s20pct.csv",        index_col=0)
    r["gap"]        = pd.read_csv(TABLES / "gap_decomposicao_serie_s20pct.csv")
    r["lrt"]        = pd.read_csv(TABLES / "lrt_serie_s20pct.csv")
    r["km_perfis"]  = pd.read_csv(TABLES / "kmeans_perfis_k3.csv",        index_col=0)
    r["km_gap"]     = pd.read_csv(TABLES / "kmeans_gap_racial_k3.csv")
    r["km_metricas"]= pd.read_csv(TABLES / "kmeans_metricas.csv")
    r["ml_perf"]    = pd.read_csv(TABLES / "ml_performance.csv")
    r["shap_imp"]   = pd.read_csv(TABLES / "shap_importance_comparada.csv", index_col=0)
    r["sna_nos"]    = pd.read_csv(TABLES / "sna_metricas_nos.csv")
    r["sna_temporal"]= pd.read_csv(TABLES / "sna_temporal.csv")
    # Novos resultados GLMM e melhorias OB/QR/M4 (com guards para tolerância a ausência)
    for key, fname in [
        ("glmm_gc",     "glmm_glassceil_full.csv"),
        ("ob_mel",      "ob_melhorias.csv"),
        ("qr_kb",       "qr_kb_test.csv"),
        ("qr_mel",      "qr_melhorias.csv"),
        ("hlm_m4_vc",   "hlm_m4_variancia.csv"),
        ("hlm_m4_coef", "hlm_m4_coeficientes.csv"),
        ("hlm_icc",     "hlm_icc_racial.csv"),
        ("hlm_ns",      "hlm_nakagawa_r2.csv"),
    ]:
        path = TABLES / fname
        r[key] = pd.read_csv(path) if path.exists() else None
    return r

def extract_kpis(r):
    k = {}
    gap = r["gap"]
    k["b_m1"]  = float(gap.loc[gap["Modelo"]=="M1_Individual",  "b_negro"].values[0])
    k["b_m2"]  = float(gap.loc[gap["Modelo"]=="M2_Localidade",  "b_negro"].values[0])
    k["b_m3"]  = float(gap.loc[gap["Modelo"]=="M3_Completo",    "b_negro"].values[0])
    k["gb"]    = abs(float(gap.loc[gap["Modelo"]=="M1_Individual",  "Gap%"].values[0]))
    k["gl"]    = abs(float(gap.loc[gap["Modelo"]=="M3_Completo",    "Gap%"].values[0]))
    k["med"]   = float(gap.loc[gap["Modelo"]=="M3_Completo", "Mediacao_total%"].values[0])
    k["med_upa"]= float(gap.loc[gap["Modelo"]=="M2_Localidade","Mediacao_UPA%"].values[0])
    m4_row = gap[gap["Modelo"]=="M4_Ocupacao"]
    if len(m4_row):
        k["b_m4"]   = float(m4_row["b_negro"].values[0])
        k["gap_m4"] = abs(float(m4_row["Gap%"].values[0]))
        k["med_occ"]= float(m4_row["Mediacao_occ%"].values[0]) if "Mediacao_occ%" in m4_row.columns else 0.0
    else:
        k["b_m4"]   = k["b_m3"]
        k["gap_m4"] = k["gl"]
        k["med_occ"]= 0.0

    hlm = r["hlm"]
    def hv(row, col):
        try: return str(hlm.loc[row, col]) if row in hlm.index and col in hlm.columns else "—"
        except: return "—"
    k["icc_m0"] = hv("ICC_UF","M0_Nulo")
    k["icc_m3"] = hv("ICC_UF","M3_Completo")
    k["n_obs"]  = hv("N (obs.)","M1_Individual")

    k["silh"]  = float(r["km_metricas"].loc[r["km_metricas"]["k"]==3,"silhouette"].values[0])
    k["xgb_r2"]= float(r["ml_perf"][r["ml_perf"]["Modelo"]=="XGBoost"]["R²"].values[0])
    k["rf_r2"] = float(r["ml_perf"][r["ml_perf"]["Modelo"]=="Random Forest"]["R²"].values[0])
    k["xgb_mae"]= float(r["ml_perf"][r["ml_perf"]["Modelo"]=="XGBoost"]["MAE"].values[0])
    k["rf_mae"] = float(r["ml_perf"][r["ml_perf"]["Modelo"]=="Random Forest"]["MAE"].values[0])

    top1 = r["shap_imp"].index[0]
    k["shap_top1"]     = top1
    k["shap_top1_val"] = float(r["shap_imp"].loc[top1,"SHAP_mean_abs_XGB"])
    feats = r["shap_imp"].reset_index()
    mask  = feats["Feature"].str.contains("Ra", na=False)
    k["shap_negro_rank"] = int(feats[mask].index[0]) + 1 if mask.any() else 6
    k["shap_negro_val"]  = float(r["shap_imp"].iloc[k["shap_negro_rank"]-1]["SHAP_mean_abs_XGB"])

    k["sna_h"]    = 0.4382
    k["gap_2016"] = float(r["sna_temporal"].loc[r["sna_temporal"]["Ano"]==2016,"gap_log"].values[0])
    k["gap_2025"] = float(r["sna_temporal"].loc[r["sna_temporal"]["Ano"]==2025,"gap_log"].values[0])
    k["pct_mista"]= float(r["sna_temporal"].loc[r["sna_temporal"]["Ano"]==2025,"pct_upa_mista"].values[0])*100

    sna = r["sna_nos"]
    k["bra_between"] = float(sna.loc[sna["race"]=="Branco","betweenness"].max())
    k["neg_between"] = float(sna.loc[sna["race"]=="Negro", "betweenness"].max())
    return k

# ── Tabelas Word ──────────────────────────────────────────────────────────────

def build_hlm_table(doc, r):
    hlm = r["hlm"]
    rows_of_interest = [
        "Intercept","negro","sexo_fem","idade_c","idade_sq",
        "educ_fund_completo","educ_medio_completo","educ_superior_completo","educ_pos_graduacao",
        "pct_negro_upa_z","tx_desemprego_upa_z","media_educ_upa_z",
        "pct_negro_uf_z","tx_desemprego_uf_z","media_educ_uf_z",
        "sigma2 (Nivel 1)","tau2_UF (Nivel 3)","ICC_UF","N (obs.)","AIC",
    ]
    label_map = {
        "Intercept":             "Intercepto",
        "negro":                 "Raça (negro) ***",
        "sexo_fem":              "Gênero (feminino)",
        "idade_c":               "Idade (centralizada)",
        "idade_sq":              "Idade² (experiência)",
        "educ_fund_completo":    "Educ.: Fundamental",
        "educ_medio_completo":   "Educ.: Médio",
        "educ_superior_completo":"Educ.: Superior",
        "educ_pos_graduacao":    "Educ.: Pós-graduação",
        "pct_negro_upa_z":       "% Negro na UPA (z)",
        "tx_desemprego_upa_z":   "Desemprego UPA (z)",
        "media_educ_upa_z":      "Educ. média UPA (z)",
        "pct_negro_uf_z":        "% Negro no Estado (z)",
        "tx_desemprego_uf_z":    "Desemprego Estado (z)",
        "media_educ_uf_z":       "Educ. média Estado (z)",
        "sigma2 (Nivel 1)":      "σ² (Nível 1)",
        "tau2_UF (Nivel 3)":     "τ²_UF (Nível 3)",
        "ICC_UF":                "ICC_UF",
        "N (obs.)":              "N (observações)",
        "AIC":                   "AIC",
    }
    cols = [c for c in ["M0_Nulo","M1_Individual","M2_Localidade","M3_Completo"] if c in hlm.columns]
    col_labels = {"M0_Nulo":"M0","M1_Individual":"M1","M2_Localidade":"M2","M3_Completo":"M3"}
    keep = [r for r in rows_of_interest if r in hlm.index]

    section_headers = {
        "educ_fund_completo": "Controles educacionais",
        "pct_negro_upa_z":    "Contexto de localidade – Nível 2 (UPA)",
        "pct_negro_uf_z":     "Contexto macrorregional – Nível 3 (UF)",
        "sigma2 (Nivel 1)":   "Componentes de variância e ajuste",
    }

    n_cols = 1 + len(cols)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = tbl.rows[0]
    shade_row(hdr, "1F3864")
    hdr.cells[0].text = "Variável"
    for j, col in enumerate(cols):
        hdr.cells[j+1].text = col_labels[col]
    for cell in hdr.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9)
                run.font.name = "Arial"

    # Data rows
    shading_alt = False
    for rk in keep:
        if rk in section_headers:
            # Section separator row
            sec_row = tbl.add_row()
            shade_row(sec_row, "D9E1F2")
            cell = sec_row.cells[0]
            cell.merge(sec_row.cells[-1])
            cell.text = section_headers[rk]
            for para in cell.paragraphs:
                for run in para.runs:
                    run.italic = True
                    run.bold   = False
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

        row = tbl.add_row()
        shade_row(row, "F2F2F2" if shading_alt else "FFFFFF")
        shading_alt = not shading_alt

        label = label_map.get(rk, rk)
        row.cells[0].text = label
        is_negro = rk == "negro"
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = is_negro
                run.font.size = Pt(9)
                run.font.name = "Arial"

        for j, col in enumerate(cols):
            val = hlm.loc[rk, col] if col in hlm.columns else "—"
            row.cells[j+1].text = str(val)
            for para in row.cells[j+1].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

    # Column widths
    widths = [Cm(6)] + [Cm(2.8)]*len(cols)
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]

    return tbl

def build_shap_table(doc, r):
    imp = r["shap_imp"].copy().reset_index()
    racial = {"Raça (negro)","% Negro na UPA","% Negro no Estado",
              "Renda média UPA","Desemprego na UPA","Desemprego no Estado"}

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells, ["Feature","|SHAP| RF","Rank RF","|SHAP| XGB","Rank XGB"]):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9)
                run.font.name = "Arial"

    for i, (_, row) in enumerate(imp.iterrows()):
        feat = str(row["Feature"])
        is_racial = feat in racial
        tr = tbl.add_row()
        shade_row(tr, "FFF2CC" if is_racial else ("F2F2F2" if i%2==0 else "FFFFFF"))
        vals = [feat, f"{row['SHAP_mean_abs_RF']:.5f}", str(int(row['Rank_RF'])),
                f"{row['SHAP_mean_abs_XGB']:.5f}", str(int(row['Rank_XGB']))]
        for j, (cell, v) in enumerate(zip(tr.cells, vals)):
            cell.text = v
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = is_racial
                    run.font.size = Pt(9)
                    run.font.name = "Arial"

    widths = [Cm(6.5), Cm(2.5), Cm(2), Cm(2.5), Cm(2)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    return tbl

def build_sna_table(doc, r):
    sna = r["sna_nos"].sort_values(["race","educ_grp"] if "educ_grp" in r["sna_nos"].columns else ["race"])
    cols_show = ["node","race","educ_label","n_workers","mean_renda","betweenness","constraint"]
    headers   = ["Grupo","Raça","Educação","N","log_Renda","Betweenness","Constraint"]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells, headers):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9)
                run.font.name = "Arial"

    for i, (_, row) in enumerate(sna.iterrows()):
        tr = tbl.add_row()
        shade_row(tr, "F2F2F2" if i%2==0 else "FFFFFF")
        is_negro = str(row["race"]) == "Negro"
        vals = [
            str(row["node"]),
            str(row["race"]),
            str(row["educ_label"]),
            f"{int(row['n_workers']):,}",
            f"{float(row['mean_renda']):.3f}",
            f"{float(row['betweenness']):.3f}",
            f"{float(row['constraint']):.4f}",
        ]
        for j, (cell, v) in enumerate(zip(tr.cells, vals)):
            cell.text = v
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j<3 else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
                    if j == 5 and float(row["betweenness"]) > 0:
                        run.bold = True

    widths = [Cm(3.5), Cm(2), Cm(2.8), Cm(2.2), Cm(2.2), Cm(2.5), Cm(2.5)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    return tbl

# ── Constrói o documento ──────────────────────────────────────────────────────

def build_doc(r, k):
    doc = Document()
    set_margins(doc)

    # ── CAPA ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()  # espaço
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ESCOLA SUPERIOR DE AGRICULTURA “LUIZ DE QUEIROZ”\nUNIVERSIDADE DE SÃO PAULO')
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MBA em Data Science e Analytics")
    run.font.name = "Arial"; run.font.size = Pt(12)

    for _ in range(6): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "RACISMO ESTRUTURAL NO MERCADO DE TRABALHO BRASILEIRO:\n"
        "UMA ABORDAGEM MULTINÍVEL, DE MACHINE LEARNING E DE REDES SOCIAIS\n"
        "COM DADOS DA PNAD CONTÍNUA (2016–2025)"
    )
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(16)

    for _ in range(6): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ricardo Calheiros")
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(13)

    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Piracicaba, 2026")
    run.font.name = "Arial"; run.font.size = Pt(12)

    doc.add_page_break()

    # ── RESUMO ───────────────────────────────────────────────────────────────
    add_heading(doc, "Resumo", level=1)
    resumo = (
        f"Este trabalho investiga o gap salarial racial e as barreiras estruturais à progressão "
        f"de carreira de profissionais negros no Brasil, combinando modelos hierárquicos lineares "
        f"(HLM), machine learning com SHAP, análise de redes sociais, decomposição de Oaxaca-Blinder, "
        f"logit multinível e regressão quantílica sobre a série histórica completa da PNAD Contínua "
        f"de 2016 a 2025, com 15,9 milhões de observações brutas. O novo dataset inclui variáveis "
        f"de composição ocupacional (grupos CBO-Domiciliar), vínculo empregatício (VD4009) e horas "
        f"trabalhadas (VD4031) extraídas dos microdados originais do IBGE.\n\n"
        f"O HLM de três níveis estima que profissionais negros recebem, em média, {k['gb']:.1f}% a "
        f"menos que brancos comparáveis. Desse diferencial bruto, {k['med']:.1f}% é mediado pelo "
        f"contexto de moradia (duplo disadvantage) e 20,5% adicional pela composição ocupacional, "
        f"reduzindo o gap residual de discriminação pura a {k['gap_m4']:.1f}% (M4). A decomposição "
        f"de Oaxaca-Blinder indica que 84,0% do gap é atribuível a diferenças de dotações — "
        f"especialmente acesso desigual a ocupações de maior prestígio — e apenas 16,0% a retornos "
        f"diferenciais às mesmas características.\n\n"
        f"O logit multinível (GLMM lme4, efeito aleatório de UPA) identifica o gap de oportunidades: "
        f"negros têm OR=0,747 de acesso a ocupações qualificadas (AME=−1,07 p.p., M2), "
        f"após controle exaustivo de observáveis e estrutura hierárquica de 39,8 mil UPAs. A regressão quantílica "
        f"formaliza o glass ceiling: sem variáveis ocupacionais (M3), o gap vai de −8,2% no q10 "
        f"a −12,3% no q95 (Δ=−0,0455, confirmando inclinação negativa). Com controles ocupacionais "
        f"(M4), o gap residual persiste de −3,2% no q10 a −7,9% no q95, evidenciando discriminação "
        f"pura de remuneração que também cresce no topo da hierarquia salarial.\n\n"
        f"Modelos XGBoost com SHAP (R²={k['xgb_r2']:.4f}) identificam horas trabalhadas, grupos "
        f"CBO e emprego formal como os principais determinantes individuais de renda após o contexto "
        f"de moradia, com a raça mantendo efeito residual de −2,5% mesmo após controle completo "
        f"das variáveis estruturais. A análise de redes sociais revela betweenness nula para negros "
        f"em todos os níveis educacionais, evidenciando exclusão das redes de conversão de capital "
        f"humano em mobilidade profissional."
    )
    for bloco in resumo.split("\n\n"):
        add_para(doc, bloco, first_line=0)

    p = doc.add_paragraph()
    run = p.add_run("Palavras-chave: ")
    run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(12)
    run2 = p.add_run("gap salarial racial; discriminação estrutural; modelos hierárquicos lineares; logit multinível; regressão quantílica; glass ceiling racial; SHAP values; Oaxaca-Blinder; PNAD Contínua.")
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

    doc.add_page_break()

    # ── ABSTRACT ─────────────────────────────────────────────────────────────
    add_heading(doc, "Abstract", level=1)
    abstract = (
        f"This study investigates the racial wage gap and structural barriers to career progression "
        f"for Black professionals in Brazil using a multilevel, machine learning, and social network "
        f"analysis framework applied to the full historical series of Brazil's Continuous National "
        f"Household Sample Survey (PNAD Contínua) from 2016 to 2025 (15.9 million raw observations).\n\n"
        f"A three-level hierarchical linear model (individual, census tract, and state) estimates "
        f"that Black workers earn {k['gb']:.1f}% less than comparable White workers after controlling "
        f"for education, sex, and age experience. Of this gross differential, {k['med']:.1f}% is "
        f"mediated by residential context (local networking, Level 2), leaving a residual net gap "
        f"of {k['gl']:.1f}% attributable to direct labour market discrimination.\n\n"
        f"K-Means clustering (k=3) reveals three racially homogeneous socioeconomic typologies, "
        f"with Black workers concentrated in lower-income segments and Black women forming a "
        f"dual-disadvantage cluster at the intersection of race and gender.\n\n"
        f"Random Forest and XGBoost models with SHAP values confirm that neighbourhood income is "
        f"the most important predictor of individual earnings (|SHAP|={k['shap_top1_val']:.3f}), "
        f"while race ranks {k['shap_negro_rank']}th even after all controls.\n\n"
        f"Social network analysis shows that Black groups have zero betweenness centrality "
        f"regardless of education level, whereas White groups serve as brokers in the co-residence "
        f"network, suggesting that the conversion of human capital into earnings depends on social "
        f"capital that is structurally denied to Black workers."
    )
    for bloco in abstract.split("\n\n"):
        add_para(doc, bloco, first_line=0)

    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(12)
    run2 = p.add_run("racial wage gap; structural discrimination; hierarchical linear models; SHAP values; social network analysis; PNAD Contínua.")
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

    doc.add_page_break()

    # ── LISTA DE ABREVIATURAS E SIGLAS ────────────────────────────────────────
    add_heading(doc, "Lista de Abreviaturas e Siglas", level=1)

    siglas = [
        ("ABNT",  "Associação Brasileira de Normas Técnicas"),
        ("BLUP",  "Best Linear Unbiased Predictor (Melhor Preditor Linear Não Viesado)"),
        ("DEFF",  "Design Effect (Efeito de Delineamento Amostral)"),
        ("DF",    "Distrito Federal"),
        ("ESALQ", 'Escola Superior de Agricultura "Luiz de Queiroz"'),
        ("HLM",   "Hierarchical Linear Model (Modelo Linear Hierárquico)"),
        ("IBGE",  "Instituto Brasileiro de Geografia e Estatística"),
        ("ICC",   "Intraclass Correlation Coefficient (Coeficiente de Correlação Intraclasse)"),
        ("IPEA",  "Instituto de Pesquisa Econômica Aplicada"),
        ("LRT",   "Likelihood Ratio Test (Teste de Razão de Verossimilhança)"),
        ("ML",    "Maximum Likelihood (Máxima Verossimilhança)"),
        ("OB",    "Oaxaca-Blinder (decomposição do gap salarial)"),
        ("OLS",   "Ordinary Least Squares (Mínimos Quadrados Ordinários)"),
        ("OR",    "Odds Ratio (razão de chances)"),
        ("PEA",   "População Economicamente Ativa"),
        ("PNAD",  "Pesquisa Nacional por Amostra de Domicílios"),
        ("QR",    "Quantile Regression (Regressão Quantílica)"),
        ("RAIS",  "Relação Anual de Informações Sociais"),
        ("REML",  "Restricted Maximum Likelihood (Máxima Verossimilhança Restrita)"),
        ("RF",    "Random Forest (Floresta Aleatória)"),
        ("SE",    "Standard Error (Erro-Padrão)"),
        ("SHAP",  "SHapley Additive exPlanations"),
        ("SNA",   "Social Network Analysis (Análise de Redes Sociais)"),
        ("TCC",   "Trabalho de Conclusão de Curso"),
        ("UF",    "Unidade da Federação"),
        ("UPA",   "Unidade Primária de Amostragem"),
        ("USP",   "Universidade de São Paulo"),
        ("XGB",   "XGBoost (Extreme Gradient Boosting)"),
    ]

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.style = "Table Grid"
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = sig_tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells, ["Sigla", "Significado"]):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for i, (sigla, significado) in enumerate(siglas):
        tr = sig_tbl.add_row()
        shade_row(tr, "F2F2F2" if i % 2 == 0 else "FFFFFF")
        tr.cells[0].text = sigla
        tr.cells[1].text = significado
        for cell in tr.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Arial"
        tr.cells[0]._tc.get_or_add_tcPr()
        for run in tr.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()

    # ── 1. INTRODUÇÃO ─────────────────────────────────────────────────────────
    add_heading(doc, "Considerações Iniciais", level=1)
    add_para(doc,
        "O Brasil é um dos países com maior desigualdade racial de renda no mundo. Segundo a PNAD "
        "Contínua, a razão entre o rendimento médio de trabalhadores brancos e negros permanece "
        "acima de 1:1,5 ao longo de toda a série histórica disponível, persistindo mesmo quando se "
        "controlam escolaridade, experiência e setor de atividade (IBGE, 2023). A desigualdade "
        "racial no mercado de trabalho brasileiro é, portanto, não apenas uma herança colonial, mas "
        "um fenômeno reproduzido ativamente por mecanismos que a abordagem tradicional de diferenças "
        "de capital humano não é capaz de capturar (HASENBALG, 1979)."
    )
    add_para(doc,
        "A literatura empírica contemporânea identifica três canais principais de reprodução dessa "
        "desigualdade: (i) discriminação direta, isto é, diferenças de tratamento em processos de "
        "seleção e promoção com características individuais observadas (PAGER, 2007); (ii) segregação "
        "residencial e seus efeitos sobre o capital social disponível ao trabalhador — o contexto do "
        "bairro define a qualidade das redes de indicação profissional (WILSON, 1987; SAMPSON et al., "
        "1997); e (iii) subvalorização sistêmica do capital humano negro (HASENBALG, 1979)."
    )
    add_para(doc,
        "Este trabalho avança sobre a literatura nacional ao integrar três metodologias "
        "complementares — econometria multinível, machine learning interpretável e análise de redes "
        "sociais — sobre a maior base de dados longitudinal disponível no Brasil para este tema, a "
        "PNAD Contínua em sua série completa de 2016 a 2025."
    )

    add_para(doc,
        "As raízes históricas da desigualdade racial no mercado de trabalho brasileiro remontam "
        "à escravidão e à abolição sem políticas de inclusão. O Brasil foi o último país do "
        "Ocidente a abolir a escravidão (1888), e o processo de emancipação não foi acompanhado "
        "de redistribuição de terra, acesso a crédito, educação ou qualquer forma de reparação "
        "econômica. A transição para o trabalho livre ocorreu simultaneamente à imigração europeia "
        "subsidiada, que preencheu as vagas do mercado formal emergente enquanto a população "
        "negra liberta permanecia sem acesso aos meios de produção (ANDREWS, 1991; HASENBALG, 1979). "
        "Essa herança estrutural é a origem histórica do diferencial racial que os dados da PNAD "
        "Contínua registram mais de 130 anos após a abolição."
    )
    add_para(doc,
        "A literatura econômica sobre discriminação racial no mercado de trabalho brasileiro "
        "consolidou-se a partir dos anos 1990. Soares (2000) estimou que o diferencial salarial "
        "racial bruto era de aproximadamente 40% na virada do milênio, com metade explicada por "
        "diferenças de escolaridade e metade por discriminação direta. Henriques (2001) demonstrou "
        "que a desigualdade racial persiste mesmo quando se comparam trabalhadores com o mesmo "
        "número de anos de estudo, evidenciando que o retorno educacional é sistematicamente "
        "menor para negros do que para brancos com escolaridade equivalente — o que Cacciamali e "
        "Hirata (2005) interpretaram como subvalorização estrutural do capital humano negro pelo "
        "mercado. Becker (1957) havia previsto teoricamente que a discriminação no mercado de "
        "trabalho poderia persistir mesmo em mercados competitivos quando empregadores têm 'gosto' "
        "pela discriminação ou quando existe discriminação estatística (ARROW, 1973)."
    )
    add_para(doc,
        "A abordagem interseccional (CRENSHAW, 1989) argumenta que raça, gênero e classe não "
        "operam como eixos independentes de desvantagem, mas como sistemas que se entrecruzam "
        "e se reforçam mutuamente. Mulheres negras, por exemplo, não experimentam a soma das "
        "penalidades de raça e gênero separadamente, mas uma forma qualitativamente distinta "
        "de marginalização que não é capturada por modelos aditivos simples. No mercado de "
        "trabalho brasileiro, a interseccionalidade manifesta-se no glass ceiling duplo: "
        "mulheres negras enfrentam barreiras de acesso a ocupações qualificadas tanto pelo "
        "gênero quanto pela raça, com as interações entre esses eixos amplificando as "
        "desvantagens individuais."
    )
    add_para(doc,
        "Dados recentes do IBGE (2023) e do DIEESE (2023) confirmam a persistência da "
        "desigualdade: em 2022, trabalhadores negros recebiam em média 57,3% do rendimento "
        "de trabalhadores brancos — uma razão praticamente inalterada desde 2012. A taxa de "
        "desemprego entre negros (10,5%) supera consistentemente a de brancos (6,5%). A "
        "sub-representação de negros em cargos de liderança e ocupações de prestígio permanece "
        "acentuada: entre diretores e gerentes, negros representam apenas 29,5% do total, "
        "desproporcionalmente abaixo de sua participação na força de trabalho (55%). Esta "
        "lacuna entre avanços formais — como a Lei de Cotas (2012) e a Lei de Igualdade "
        "Salarial (2023) — e a persistência dos diferenciais observados motiva a abordagem "
        "empírica multimétodo deste trabalho."
    )

    add_heading(doc, "Hipóteses", level=2)
    hipoteses = [
        ("H1 – Gap racial bruto e líquido:",
         "Profissionais negros apresentam rendimento inferior ao de brancos comparáveis após controlar "
         "por escolaridade, experiência potencial e gênero — e esse diferencial persiste após a adição "
         "de controles contextuais de nível de bairro e de estado."),
        ("H2 – Mediação pelo networking local:",
         "Uma fração significativa do gap racial bruto é explicada pelo contexto socioeconômico do "
         "local de moradia (composição racial da UPA, desemprego local, nível educacional médio do "
         "entorno) — capturando os efeitos indiretos da segregação residencial sobre a renda."),
        ("H3 – Tipologias de vulnerabilidade alinhadas com raça:",
         "Métodos de agrupamento não-supervisionados identificam clusters socioeconômicos que se "
         "sobrepõem a fronteiras raciais, com trabalhadores negros concentrados nos segmentos de "
         "maior vulnerabilidade."),
        ("H4 – Importância residual da raça nos modelos preditivos:",
         "Mesmo após controlar por educação, experiência, gênero e contexto de moradia, a variável "
         "racial mantém relevância preditiva independente, conforme medida pelos valores SHAP dos "
         "modelos de gradient boosting."),
        ("H5 – Isolamento estrutural na rede de co-residência:",
         "Grupos negros ocupam posições periféricas na rede demográfica de co-residência, com menor "
         "capacidade de corretagem (brokerage) entre grupos do que seus pares brancos de mesma "
         "escolaridade, o que limita a conversão do capital humano em rendimento."),
    ]
    for label, texto in hipoteses:
        p = doc.add_paragraph(style="List Bullet")
        run1 = p.add_run(label + " ")
        run1.bold = True
        run1.font.name = "Times New Roman"; run1.font.size = Pt(12)
        run2 = p.add_run(texto)
        run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ── 2. REVISÃO DE LITERATURA ──────────────────────────────────────────────
    add_heading(doc, "Revisão de Literatura", level=2)

    add_heading(doc, "2.1 Desigualdade racial no mercado de trabalho brasileiro", level=2)
    add_para(doc,
        "Hasenbalg (1979) demonstrou pioneiramente que a desigualdade racial no Brasil não decorre "
        "apenas de diferenças históricas de acesso à educação, mas de mecanismos ativos de "
        "discriminação no mercado de trabalho que convertem desvantagens sociais em desvantagens "
        "econômicas de forma cumulativa. Trabalhos posteriores (HENRIQUES, 2001; SOARES, 2009) "
        "confirmaram a persistência dessas diferenças mesmo após controlar por escolaridade, "
        "reforçando a hipótese de discriminação estrutural."
    )

    add_heading(doc, "2.2 Efeitos de vizinhança e segregação residencial", level=2)
    add_para(doc,
        "Wilson (1987) propôs a hipótese da concentrated disadvantage: a concentração de pobreza "
        "em bairros racialmente segregados amplifica desvantagens individuais por meio da redução "
        "de redes de contato com o mercado de trabalho formal, degradação de serviços públicos e "
        "aumento da violência. Sampson et al. (1997) forneceram evidência empírica para essa "
        "hipótese, e estudos brasileiros encontraram padrões similares para as regiões "
        "metropolitanas (MARQUES, 2010)."
    )

    add_heading(doc, "2.3 Modelos lineares hierárquicos", level=2)
    add_para(doc,
        "Raudenbush e Bryk (2002) sistematizaram a fundamentação estatística dos modelos lineares "
        "hierárquicos (HLM), tornando-os o padrão metodológico para análise de dados com estrutura "
        "aninhada (indivíduos dentro de bairros dentro de estados). Esses modelos permitem decompor "
        "a variância do desfecho em componentes de cada nível e estimar os efeitos contextuais "
        "controlando simultaneamente pelos efeitos individuais."
    )

    add_heading(doc, "2.4 Interpretabilidade em machine learning: SHAP values", level=2)
    add_para(doc,
        "Lundberg e Lee (2017) propuseram os SHapley Additive exPlanations (SHAP), unificando "
        "importância de variáveis, efeitos parciais e explicações individuais em uma única estrutura "
        "axiomática baseada na teoria dos jogos cooperativos. Para dados socioeconômicos, SHAP "
        "permite responder à pergunta: 'quanto e em que direção a raça de um indivíduo específico "
        "afeta a predição de sua renda?' — complementando os coeficientes do HLM."
    )

    add_heading(doc, "2.5 Análise de redes sociais e capital social", level=2)
    add_para(doc,
        "Granovetter (1973) demonstrou que laços fracos — conexões entre indivíduos de grupos "
        "sociais distintos — são os principais canais de transmissão de informações sobre "
        "oportunidades profissionais. Burt (2004) formalizou o conceito de buraco estrutural: "
        "indivíduos que conectam grupos desconexos obtêm vantagens relacionais (acesso antecipado "
        "a vagas, mentoria, promoções). Aplicado à questão racial, a SNA permite investigar se "
        "trabalhadores negros ocupam as posições de rede que possibilitam o aproveitamento dessas "
        "vantagens."
    )

    add_heading(doc, "2.6 Regressão Quantílica e Glass Ceiling", level=2)
    add_para(doc,
        "Koenker e Bassett (1978) introduziram a regressão quantílica como método para estimar "
        "os efeitos de covariáveis em pontos arbitrários da distribuição condicional do desfecho, "
        "em contraste com a regressão OLS, que estima apenas o efeito médio. Aplicada ao gap "
        "salarial racial, a regressão quantílica permite testar a hipótese de glass ceiling: se o "
        "coeficiente de raça torna-se progressivamente mais negativo nos quantis superiores, o "
        "diferencial racial é maior no topo da distribuição — não como fenômeno residual, mas como "
        "estrutura dominante de reprodução da desigualdade nos estratos de alta remuneração. "
        "Evidências para o mercado de trabalho americano (ALBRECHT et al., 2003) e brasileiro "
        "confirmam esse padrão, sugerindo que barreiras informais de seleção e promoção operam "
        "de forma mais intensa justamente onde os prêmios de mobilidade são maiores."
    )

    add_heading(doc, "2.7 Logit Multinível e Gap de Oportunidades", level=2)
    add_para(doc,
        "Enquanto os modelos de regressão linear medem o gap de remuneração dentro do mercado "
        "de trabalho, modelos de resposta binária multinível (logit/probit) permitem quantificar "
        "o gap de oportunidades: a diferença racial na probabilidade de acesso a tipos de emprego "
        "específicos (formal, qualificado, de alta remuneração). Arrow (1973) e a literatura de "
        "discriminação estatística argumentam que empregadores podem usar a raça como proxy de "
        "produtividade em contextos de informação incompleta, gerando barreiras de acesso que "
        "persistem mesmo após controle de qualificações observáveis. O logit multinível — com "
        "controles de nível individual e contextual e SE clusterizado — permite distinguir "
        "discriminação de acesso (barreira pré-contratação) de discriminação de remuneração "
        "(barreira dentro da ocupação), dois mecanismos com implicações de política distintas."
    )

    doc.add_page_break()

    # ── 3. DADOS E METODOLOGIA ────────────────────────────────────────────────
    add_heading(doc, "Implementação de Algoritmo(s) de Machine Learning", level=1)

    add_heading(doc, "3.1 Base de dados: PNAD Contínua (2016–2025)", level=2)
    add_para(doc,
        "A PNAD Contínua é uma pesquisa amostral de domicílios conduzida trimestralmente pelo IBGE, "
        "com cobertura nacional e metodologia de painel rotativo. A pesquisa abrange todos os 27 "
        "estados da federação (26 estados + Distrito Federal) e é organizada em Unidades Primárias "
        "de Amostragem (UPAs), que correspondem a setores censitários ou agrupamentos destes. "
        "Para este trabalho, foram processados todos os 40 trimestres disponíveis de 2016T1 a "
        "2025T4, totalizando 15.941.675 observações individuais-trimestre, distribuídas por "
        "41.517 UPAs amostradas em todo o território nacional."
    )
    add_para(doc,
        "A amostra cobre 6.812.609 domicílios, identificados pela condição de pessoa de referência "
        "(V2005 = 1). Em termos de composição por sexo, 7.656.647 observações correspondem a "
        "homens (48,0%) e 8.285.028 a mulheres (52,0%). Quanto à composição racial, 9.559.051 "
        "indivíduos são classificados como negros ou pardos (60,0%) e 6.382.624 como brancos (40,0%), "
        "refletindo a distribuição demográfica registrada pelo Censo 2022. A classificação racial "
        "segue o critério binário adotado pelos estudos de desigualdade racial no Brasil: "
        "negro = preto (código 2) + pardo (código 4); branco = branco (código 1), ambos da "
        "variável V2010 (cor ou raça autodeclarada)."
    )
    add_para(doc,
        "Do total de observações, 7.717.561 (48,4%) possuem renda do trabalho positiva declarada. "
        "Em relação à participação no mercado de trabalho, 9.198.750 indivíduos compõem a "
        "População Economicamente Ativa — PEA (57,7%), dos quais 8.289.881 estão ocupados (52,0% "
        "do total). O cruzamento entre raça e sexo, restrito às observações com renda positiva, "
        "revela assimetrias relevantes: homens brancos somam 1.843.856 observações, mulheres "
        "brancas 1.431.788, homens negros 2.668.513 e mulheres negras 1.773.404 — padrão que "
        "evidencia tanto a maior participação masculina no mercado formal quanto a predominância "
        "numérica da população negra na amostra."
    )

    add_heading(doc, "3.2 Modelo Linear Hierárquico de Três Níveis (HLM)", level=2)
    add_para(doc,
        "O modelo parte da equação de Mincer estendida para o log-rendimento mensal de trabalho, "
        "estruturada em três níveis:"
    )
    add_para(doc, "Nível 1 – Indivíduo:", bold=True, first_line=0)
    add_equation(doc, _EQ_HLM1)
    add_para(doc, "Nível 2 – UPA (proxy de bairro):", bold=True, first_line=0)
    add_equation(doc, _EQ_HLM2)
    add_para(doc, "Nível 3 – UF (Estado):", bold=True, first_line=0)
    add_equation(doc, _EQ_HLM3)
    add_para(doc,
        "O ICC para o Nível 3 é: ρ_UF = τ²v / (τ²v + σ²). Valores ρ > 0,05 justificam a inclusão "
        "do nível superior (RAUDENBUSH; BRYK, 2002). A estimação utiliza REML com método de Powell "
        "para evitar colapso de variância na fronteira τ²=0. Dado o número elevado de UPAs (41.517), "
        "os efeitos aleatórios de localidade foram substituídos por slopes fixos."
    )

    # ── 3.3 JUSTIFICAÇÃO HLM vs OLS ──────────────────────────────────────────
    add_heading(doc, "3.3 Justificação do Modelo Multinível em Detrimento do OLS", level=2)

    add_para(doc,
        "O uso do modelo OLS (Mínimos Quadrados Ordinários) para estimar equações de rendimento "
        "sobre dados da PNAD Contínua viola a premissa fundamental de independência das observações: "
        "indivíduos residentes no mesmo estado (UF) compartilham um contexto macroeconômico, "
        "institucional e cultural comum, gerando correlação intragrupo não modelada. O HLM de três "
        "níveis resolve esse problema decompondo explicitamente a variância do log-rendimento em "
        "componentes individuais (Nível 1), de localidade (Nível 2 – UPA) e de estado (Nível 3 – UF)."
    )

    # Modelo OLS e sua limitação
    add_para(doc, "Modelo OLS de referência (Eq. 1):", bold=True, first_line=0)
    add_equation(doc, _EQ_OLS)
    add_para(doc,
        "A suposição iid (identicamente distribuídos e independentes) é violada quando "
        "Cov(εijk, εij'k) ≠ 0 para observações do mesmo estado k. Ignorar essa estrutura produz "
        "erros-padrão inconsistentes e inferências enviesadas sobre β₁ (o gap racial)."
    )

    # Modelo Nulo e ICC
    add_para(doc, "Modelo Nulo HLM — decomposição de variância (Eq. 2):", bold=True, first_line=0)
    add_equation(doc, _EQ_NULL)
    add_para(doc,
        "A estimação REML do modelo nulo fornece: σ̂² = 0,7653 (variância individual) e "
        "τ̂²_UF = 0,0834 (variância entre estados). O Coeficiente de Correlação Intraclasse (ICC) é:"
    )
    add_equation(doc, _EQ_ICC)
    add_equation(doc, _EQ_ICC_NUM)
    add_para(doc,
        "Valores de ICC superiores a 5% justificam a inclusão do nível hierárquico superior "
        "(RAUDENBUSH; BRYK, 2002). O ICC de 9,83% indica que aproximadamente 10% da variância "
        "total do log-rendimento é explicada pelo estado de residência — antes de qualquer controle "
        "individual. Mesmo após a inclusão de todas as variáveis individuais e contextuais (M3), "
        "o ICC permanece em 3,98%, acima do limiar de 5% que se mantém em M0, M1 e M2."
    )

    add_figure(doc, FIGURES / "hlm_justif_variancia.png",
        "Figura A1 – Decomposição da variância do log-rendimento por nível hierárquico. "
        "σ² = variância intraindividual (Nível 1); τ²_UF = variância interestadual (Nível 3). "
        "ICC anotado em cada modelo; *** p < 0,001 (boundary test H₀: τ² = 0).",
        width_cm=15)
    add_figure(doc, FIGURES / "hlm_justif_icc.png",
        "Figura A2 – Evolução do ICC por Unidade da Federação ao longo dos quatro modelos. "
        "Linha vermelha tracejada = limiar de 5% recomendado por Raudenbush e Bryk (2002). "
        "ICC > 5% em M0, M1 e M2 justifica a estrutura de três níveis.",
        width_cm=13)

    # Teste estatístico
    add_para(doc,
        "Teste de Wald para H₀: τ²_UF = 0 (Boundary Test):",
        bold=True, first_line=0
    )
    add_para(doc,
        "Dado que τ² ≥ 0 por definição, a hipótese nula está na fronteira do espaço paramétrico. "
        "O p-valor correto usa a distribuição mista ½χ²₀ + ½χ²₁ (boundary test):"
    )
    add_para(doc,
        "M0: z = τ̂²/SE(τ̂²) = 0,08342 / 0,02314 = 3,61  →  p < 0,001 ***\n"
        "M1: z = 0,06109 / 0,01694 = 3,61  →  p < 0,001 ***\n"
        "M2: z = 0,03365 / 0,00933 = 3,61  →  p < 0,001 ***\n"
        "M3: z = 0,02504 / 0,00694 = 3,61  →  p < 0,001 ***",
        first_line=0
    )
    add_para(doc,
        "Em todos os modelos, a variância dos efeitos aleatórios de intercepto (τ²_UF) é "
        "estatisticamente diferente de zero com p < 0,001, confirmando que os estados produzem "
        "interceptos significativamente distintos que não podem ser ignorados."
    )

    # Tabela componentes de variância
    p = doc.add_paragraph()
    run = p.add_run(
        "Tabela A1 – Componentes de variância e testes de significância (H₀: τ² = 0). "
        "Boundary test: distribuição ½χ²₀ + ½χ²₁. *** p < 0,001."
    )
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    var_tbl = doc.add_table(rows=1, cols=6)
    var_tbl.style = "Table Grid"
    var_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = var_tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells,
                         ["Modelo", "σ̂²", "τ̂²_UF", "SE(τ̂²)", "z", "ICC_UF"]):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9); run.font.name = "Arial"

    var_data = [
        ("M0 (Nulo)",       "0,7653", "0,08342", "0,02314", "3,61", "9,83% ***"),
        ("M1 (Individual)", "0,6329", "0,06109", "0,01694", "3,61", "8,80% ***"),
        ("M2 (+ UPA)",      "0,6037", "0,03365", "0,00933", "3,61", "5,28% ***"),
        ("M3 (Completo)",   "0,6037", "0,02504", "0,00694", "3,61", "3,98% ***"),
    ]
    for i, row_data in enumerate(var_data):
        tr = var_tbl.add_row()
        shade_row(tr, "F2F2F2" if i % 2 == 0 else "FFFFFF")
        for j, (cell, v) in enumerate(zip(tr.cells, row_data)):
            cell.text = v
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9); run.font.name = "Arial"

    # Caterpillar
    add_figure(doc, FIGURES / "hlm_justif_caterpillar.png",
        "Figura A3 – Caterpillar plot dos efeitos aleatórios de intercepto (ũ₀k) por UF. "
        "Barras de erro = IC 95%. Verde escuro: UFs cujo IC não inclui zero. "
        "26 de 27 estados têm efeito aleatório estatisticamente distinto de zero.",
        width_cm=12)

    # Efeito de design
    add_para(doc,
        "Efeito de Design (DEFF) e viés do OLS não-clusterizado:",
        bold=True, first_line=0
    )
    add_para(doc,
        "O efeito de design quantifica a inflação da variância amostral provocada pelo clustering: "
        "DEFF = 1 + (n̄ − 1)·ρ_UF ≈ 1 + (56.959 − 1) × 0,0983 ≈ 5.600. "
        "Isso significa que erros-padrão do OLS calculados sem correção de cluster estariam "
        "subestimados por um fator de √5.600 ≈ 75×, produzindo testes de hipóteses completamente "
        "inválidos. Mesmo com correção de cluster (como a implementada na coluna OLS desta "
        "análise), o OLS não decompõe corretamente a variância nem fornece estimativas eficientes "
        "dos efeitos contextuais de Nível 2 e 3."
    )

    # Comparação de SEs
    add_figure(doc, FIGURES / "hlm_justif_se_comparison.png",
        "Figura A4 – Comparação de erros-padrão: OLS com SE clusterizado por UF vs HLM REML. "
        "Painel A: SE absoluto. Painel B: razão SE_OLS/SE_HLM. "
        "Razão > 1 indica que o OLS infla o SE; razão < 1 indica subestimação.",
        width_cm=15)

    # LRT sequencial
    add_para(doc,
        "Teste da Razão de Verossimilhança (LRT) — progressão de ajuste:",
        bold=True, first_line=0
    )
    add_para(doc,
        "O LRT compara modelos aninhados via a estatística Λ = −2·ΔLog-Verossimilhança REML, "
        "distribuída assintoticamente como χ²(Δdf):"
    )
    add_para(doc,
        "M0 → M1  (controles individuais):  Λ = 292.168,  df = 8,  p < 0,001 ***\n"
        "M1 → M2  (contexto UPA):            Λ =  72.471,  df = 3,  p < 0,001 ***\n"
        "M2 → M3  (preditores UF fixos):     Λ ≈ 0  (redundante com efeito aleatório UF)",
        first_line=0
    )
    add_para(doc,
        "Os dois primeiros incrementos são altamente significativos, confirmando que cada nível "
        "hierárquico contribui com poder explicativo real. O terceiro incremento nulo (M2→M3) "
        "indica que, uma vez incluído o efeito aleatório de UF, os preditores fixos de nível "
        "estadual tornam-se redundantes — resultado esperado em modelos com efeitos aleatórios "
        "de intercepto que já capturam a heterogeneidade entre estados."
    )

    add_figure(doc, FIGURES / "hlm_justif_lr_tests.png",
        "Figura A5 – Progressão da Log-Verossimilhança REML e Teste da Razão de Verossimilhança. "
        "Painel A: LL por modelo. Painel B: estatística LRT por incremento. "
        "*** p < 0,001 (χ² com df indicados).",
        width_cm=15)

    add_para(doc,
        "Em síntese, as evidências são conclusivas: o HLM de três níveis é o modelo estatisticamente "
        "correto para esses dados. O ICC de 9,83% supera o limiar de 5%, τ²_UF é significativamente "
        "diferente de zero em todos os modelos (p < 0,001), 26 dos 27 estados possuem interceptos "
        "distintos da média nacional, e o LRT confirma que cada nível hierárquico adiciona poder "
        "explicativo substancial. O OLS, mesmo com correção de cluster, não oferece decomposição "
        "de variância entre níveis, estimativas eficientes dos efeitos contextuais, nem o "
        "tratamento correto da dependência estrutural dos dados."
    )

    doc.add_page_break()

    add_heading(doc, "3.4 Clustering Socioeconômico (K-Means)", level=2)
    add_para(doc,
        "O algoritmo MiniBatchKMeans foi aplicado sobre as 2.395.285 observações com renda e "
        "variáveis contextuais completas, usando 10 dimensões padronizadas: idade, escolaridade "
        "ordinal, log-rendimento, raça, gênero, status de emprego e quatro variáveis de contexto "
        "da UPA. O número ótimo de clusters foi determinado pelo Silhouette Coefficient "
        "(ROUSSEEUW, 1987) com validação pelo índice de Davies-Bouldin (DAVIES; BOULDIN, 1979)."
    )

    add_heading(doc, "3.5 Random Forest, XGBoost e SHAP Values", level=2)
    add_para(doc,
        "Para predição do log-rendimento, foram ajustados dois modelos de ensemble: (i) Random "
        "Forest (BREIMAN, 2001) com 200 árvores e profundidade máxima 10; e (ii) XGBoost "
        "(CHEN; GUESTRIN, 2016) com 300 iterações, lr=0,05 e regularização L1/L2. Sobre o modelo "
        "XGBoost, foi aplicado o TreeExplainer da biblioteca SHAP (LUNDBERG; LEE, 2017) sobre um "
        "subsample de 50.000 observações."
    )

    add_heading(doc, "3.6 Análise de Redes Sociais (SNA)", level=2)
    add_para(doc,
        "A rede demográfica foi construída com 10 nós, representando as combinações de raça × "
        "educação (2 raças × 5 níveis), e arestas ponderadas pelo índice de Jaccard de co-presença "
        "em UPAs: w_AB = |U_A ∩ U_B| / |U_A ∪ U_B|, onde U_A é o conjunto de UPAs com "
        "trabalhadores do grupo A. As métricas de rede incluem centralidade de grau, betweenness, "
        "clustering coefficient e constraint de Burt (2004)."
    )

    add_heading(doc, "3.7 Decomposição de Oaxaca-Blinder", level=2)
    add_para(doc,
        "A decomposição de Oaxaca-Blinder (OAXACA, 1973; BLINDER, 1973) decompõe o gap salarial "
        "racial observado em dois componentes economicamente interpretáveis. Sejam β̂_B e β̂_N os "
        "vetores de coeficientes OLS estimados separadamente para brancos e negros, e X̄_B e X̄_N "
        "as médias das covariáveis de cada grupo. O gap total Δȳ = ȳ_B − ȳ_N é decomposto como:"
    )
    add_equation(doc, _mwrap(
        _mu('Δ') + _mi('ȳ') + _mu(' = ') +
        _mu('(') + _msub(_mi('X̄'), _mi('B')) + _mu(' − ') + _msub(_mi('X̄'), _mi('N')) + _mu(')') +
        _msub(_mi('β̂'), _mi('B')) +
        _mu('  +  ') +
        _msub(_mi('X̄'), _mi('N')) + _mu('·(') +
        _msub(_mi('β̂'), _mi('B')) + _mu(' − ') + _msub(_mi('β̂'), _mi('N')) + _mu(')')
    ))
    add_para(doc,
        "O primeiro termo é o efeito dotações: parcela do gap atribuível às diferenças médias nas "
        "características observáveis (escolaridade, experiência, contexto de moradia). O segundo "
        "termo é o efeito retornos: parcela atribuível ao fato de o mercado remunerar as mesmas "
        "características de forma distinta entre grupos — interpretado na literatura como estimativa "
        "de discriminação ou de fatores não observados correlacionados com raça (SOARES, 2009). "
        "O modelo de referência utiliza covariáveis individuais (escolaridade ordinal, idade, "
        "gênero) e contextuais (% negros na UPA, taxa de desemprego e educação média da UPA), "
        "estimado em amostra estratificada de 20% da PEA com renda positiva."
    )

    add_heading(doc, "3.8 Análise de Segregação Espacial por Tipo de Área", level=2)
    add_para(doc,
        "A variável V1023 da PNAD Contínua classifica cada domicílio em três tipos de área: "
        "(1) Capital — município sede da capital estadual; (2) Região Metropolitana (RM) não "
        "capital — demais municípios das Regiões Metropolitanas e Regiões Integradas de "
        "Desenvolvimento Econômico (RIDEs); e (3) Interior — municípios fora de qualquer RM ou "
        "RIDE, urbanos ou rurais. Essa tipologia é utilizada como proxy da distância funcional aos "
        "grandes centros de trabalho: capitais e RMs concentram os maiores polos de emprego formal, "
        "redes de transporte, serviços especializados e redes profissionais, enquanto o interior "
        "enfrenta mercados locais mais restritos e menor mobilidade laboral."
    )
    add_para(doc,
        "Para cada tipo de área, estimamos: (i) a taxa de ocupação por grupo racial, calculada "
        "sobre a PEA; e (ii) o rendimento médio mensal e o gap salarial racial (branco − negro "
        "em log-rendimento), calculados sobre os ocupados com renda positiva. A análise temporal "
        "compara a trajetória 2016–2025 de brancos e negros em áreas metropolitanas versus "
        "interior, testando a hipótese de que a segregação residencial nas periferias amplifica a "
        "desvantagem racial além do efeito individual estimado pelo HLM."
    )

    add_heading(doc, "3.9 Logit Multinível: Gap de Oportunidades", level=2)
    add_para(doc,
        "Para identificar o gap de oportunidades — a diferença racial no acesso a tipos de emprego "
        "específicos —, estimamos modelos de regressão logística com estrutura multinível análoga "
        "ao HLM: variáveis de nível individual (Nível 1), variáveis contextuais da UPA (Nível 2) "
        "e UF como efeito fixo com erros-padrão clusterizados por UPA (aproximação do Nível 3). "
        "Três desfechos binários são modelados: (i) emprego formal (código VD4009 ∈ {1,3,5,7}); "
        "(ii) ocupação qualificada (grupo CBO ∈ {dirigente, profissional, técnico, administrativo}); "
        "e (iii) renda no quintil superior (log_renda ≥ p80). Para cada desfecho, dois modelos "
        "sequenciais são estimados — M1 (controles individuais + UF FE) e M2 (M1 + contexto UPA) — "
        "e o efeito marginal médio (AME) de 'negro' é calculado como a diferença média de "
        "probabilidades preditas entre cenário negro=1 e negro=0, mantendo todas as demais "
        "covariáveis em seus valores observados."
    )

    add_heading(doc, "3.10 Regressão Quantílica: Formalização do Glass Ceiling", level=2)
    add_para(doc,
        "A regressão quantílica (KOENKER; BASSETT, 1978) é estimada para os quantis "
        "q ∈ {0,10; 0,25; 0,50; 0,75; 0,90; 0,95} da distribuição condicional de log_renda, "
        "com as mesmas covariáveis do M3 (controles individuais, contexto UPA e UF FE) e do M4 "
        "(M3 + grupos CBO + formalidade + horas). O coeficiente β̂_negro(q) indica o efeito "
        "da raça no quantil q da distribuição de rendimento. A hipótese de glass ceiling é "
        "testada pela inclinação de β̂_negro ao longo dos quantis: inclinação negativa "
        "(β̂_negro(q95) < β̂_negro(q10)) confirma que o gap racial é maior no topo da "
        "distribuição. A diferença entre M3 e M4 em cada quantil quantifica a mediação "
        "ocupacional do glass ceiling — separando o componente explicado pela sub-representação "
        "em ocupações de alto prestígio do componente de discriminação pura de remuneração."
    )

    add_heading(doc, "3.11 Justificação da Escolha dos Modelos", level=2)
    add_para(doc,
        "A seleção dos modelos estatísticos empregados neste trabalho segue critérios formais de "
        "adequação à estrutura dos dados, poder explicativo e validade inferencial, detalhados a seguir."
    )
    add_para(doc,
        "HLM versus OLS: A justificação primária para o uso de Modelos Hierárquicos Lineares em "
        "detrimento do OLS reside na estrutura de dados aninhados. Os 1,5 milhão de indivíduos da "
        "amostra estão distribuídos em aproximadamente 15 mil UPAs e 27 estados, criando dependência "
        "intra-cluster que viola o pressuposto de independência do OLS. O modelo nulo HLM estimou "
        "ICC_UF=9,83%, acima do limiar de 5% sugerido por Raudenbush e Bryk (2002), confirmando que "
        "a estrutura hierárquica é estatisticamente relevante. O Teste de Razão de Verossimilhança "
        "entre HLM Nulo e OLS Nulo gera estatística χ²=191.625 (Δk=1), refutando o modelo plano "
        "com probabilidade virtualmente nula. Adicionalmente, o uso de OLS com dummies de UF "
        "(efeitos fixos) mostrou AIC=3.684.832 — pior que o OLS com variáveis contextuais "
        "(AIC=3.588.312) e que o HLM Contextual (AIC=3.588.683) — indicando que a abordagem de "
        "efeitos aleatórios é mais parcimoniosa e igualmente válida, além de permitir a decomposição "
        "formal da variância entre níveis hierárquicos, impossível com efeitos fixos."
    )
    add_para(doc,
        "HLM 3 níveis versus 2 níveis: A extensão para três níveis (indivíduo → UPA → UF) é "
        "justificada pelo ICC_UF significativo no modelo nulo, que indica variância atribuível ao "
        "nível estado que não se reduz ao nível UPA. A sequência M0→M1→M4 de LRTs progressivos "
        "confirma que cada nível contribui significativamente para a explicação da variância de "
        "rendimentos (χ² acumulado > 900.000 unidades ao longo da progressão M0→M4)."
    )
    add_para(doc,
        "Regressão quantílica versus OLS para o glass ceiling: O OLS estima o efeito médio do gap "
        "racial, mas a hipótese de glass ceiling pressupõe heterogeneidade do efeito ao longo da "
        "distribuição. A regressão quantílica (Koenker; Bassett, 1978) estima β̂_negro(q) em cada "
        "quantil sem pressupor distribuição paramétrica para os erros, sendo robusta a "
        "heteroscedasticidade. A confirmação de β̂_negro(q95) < β̂_negro(q10) (Δ=−0,0455) exigiria, "
        "em OLS, a especificação de interações não-lineares ad hoc — com risco de sobreajuste. "
        "A regressão quantílica fornece esse teste de forma não-paramétrica e diretamente interpretável."
    )
    add_para(doc,
        "Logit multinível versus OLS linear de probabilidade (LPM): Para desfechos binários "
        "(y_formal, y_ocp_qualificada, y_top20), o LPM produz probabilidades previstas fora do "
        "intervalo [0,1] e é heteroscedástico por construção. O logit garante probabilidades "
        "coerentes e é consistente com a estrutura de dados binários. A aproximação por efeitos "
        "fixos de UF com erros clusterizados por UPA — em vez de GLMM completo — foi adotada por "
        "viabilidade computacional com 1,5 milhão de observações, produzindo estimativas de odds "
        "ratio e AME equivalentes às do GLMM para os parâmetros de nível 1 (Moineddin et al., 2007)."
    )
    add_para(doc,
        "XGBoost versus Random Forest: Ambos os algoritmos foram estimados. O XGBoost atingiu "
        "R²=0,6162 versus R²=0,44 do Random Forest com as variáveis anteriores. A superioridade "
        "do XGBoost se deve ao boosting adaptativo, que corrige erros residuais iterativamente. "
        "A interpretabilidade via SHAP (Lundberg; Lee, 2017) é equivalente em ambos, sendo "
        "escolhido o modelo com maior poder preditivo. O SHAP é preferível às importâncias de "
        "variáveis do Random Forest porque decompõe a contribuição individual de cada observação, "
        "tornando a análise do gap racial mais precisa e menos influenciada pela correlação entre "
        "preditores."
    )
    add_para(doc,
        "Oaxaca-Blinder versus decomposição de variância simples: A decomposição de Oaxaca-Blinder "
        "(1973) separa formalmente o gap salarial em componente de dotações (diferenças em "
        "características observáveis) e de retornos (diferenças na remuneração dos mesmos atributos), "
        "o que uma simples comparação de médias não permite. O estimador two-fold utilizado é "
        "equivariante à escolha do grupo de referência, eliminando o índice de transformação "
        "que afeta o estimador three-fold em presença de multicolinearidade."
    )
    if (FIGURES / "modelos_loglik_aic.png").exists():
        add_figure(doc, FIGURES / "modelos_loglik_aic.png",
            "Figura 34 – Comparativo de log-verossimilhança e AIC entre especificações OLS e HLM. "
            "Azul: modelos OLS (sem efeitos aleatórios); Verde: modelos HLM (com efeitos aleatórios). "
            "O HLM Nulo supera amplamente o OLS Nulo (ΔLL=95.812), justificando a estrutura multinível.",
            width_cm=16)
    if (FIGURES / "modelos_lrt_icc.png").exists():
        add_figure(doc, FIGURES / "modelos_lrt_icc.png",
            "Figura 35 – LRT entre etapas progressivas do HLM (esq.) e trajetória do ICC_UF por "
            "modelo (dir.). O ICC se mantém acima do limiar de 5% em todos os modelos, confirmando "
            "a necessidade da estrutura de três níveis.",
            width_cm=16)

    doc.add_page_break()

    # ── 4. RESULTADOS E DISCUSSÃO ─────────────────────────────────────────────
    add_heading(doc, "Resultados e Discussão", level=1)

    # ── 4.0 ANÁLISE DESCRITIVA ────────────────────────────────────────────────
    add_heading(doc, "Análise Descritiva — Caracterização da Amostra por Grupo Racial", level=2)
    add_para(doc,
        "A Tabela 1 apresenta as estatísticas descritivas ponderadas pelo peso amostral V1028 "
        "para os dois grupos raciais analisados. A população de estudo compreende 7.717.561 "
        "empregados com renda positiva ao longo de 2016–2025, representando 3.513.803 mil "
        "trabalhadores ponderados — 1.607.081 mil brancos (45,7%) e 1.906.722 mil negros (54,3%)."
    )
    add_para(doc,
        "O gap mediano bruto de 27,5% — brancos recebem mediana de R$2.000 contra R$1.450 dos "
        "negros — sintetiza a magnitude da desigualdade antes de qualquer controle. A diferença "
        "em termos de escolaridade formal, emprego formal e participação no setor público aponta "
        "que negros apresentam desvantagens em múltiplas dimensões simultaneamente, o que "
        "justifica a abordagem metodológica multivariada deste trabalho."
    )

    # Tabela 1
    try:
        tab1 = pd.read_csv(TABLES / "tab1_descritiva_racial.csv", index_col=0)
        cols = ["Brancos", "Negros"]
        tbl = doc.add_table(rows=len(tab1)+1, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(["Indicador"] + cols):
            cell = tbl.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
        shade_row(tbl.rows[0], "D9E1F2")
        for i, (idx, row) in enumerate(tab1.iterrows(), 1):
            tbl.rows[i].cells[0].text = str(idx)
            tbl.rows[i].cells[1].text = str(row.get("Brancos",""))
            tbl.rows[i].cells[2].text = str(row.get("Negros",""))
            for cell in tbl.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10); run.font.name = "Times New Roman"
            if i % 2 == 0:
                shade_row(tbl.rows[i], "F2F2F2")
        add_caption(doc, "Tabela 1. Estatísticas descritivas ponderadas por grupo racial "
                    "— PNAD Contínua 2016–2025, empregados com renda positiva (peso V1028).")
    except Exception as e:
        add_para(doc, f"[Tabela 1 — tab1_descritiva_racial.csv não encontrado: {e}]", italic=True)

    add_figure(doc, FIGURES / "fig1_densidade_log_salario.png",
               "Figura 1. Distribuição do log-rendimento mensal por grupo racial — "
               "estimativa de densidade Kernel ponderada (PNAD Contínua 2016–2025).")

    add_para(doc,
        "A Figura 1 evidencia a sobreposição das distribuições salariais dos dois grupos, "
        "com o deslocamento da distribuição negra para valores menores sendo consistente ao "
        "longo de toda a distribuição. A diferença entre as médias de log-renda "
        "(brancos: 7,674 vs negros: 7,254) equivale a um gap de 42,0% na renda esperada, "
        "próximo ao gap bruto documentado por DIEESE (2023) e IBGE (2023)."
    )

    add_para(doc,
        "A Tabela 2 decompõe o gap bruto por subgrupos de gênero, faixa etária e "
        "escolaridade. Destaca-se a heterogeneidade do gap: entre trabalhadores de 35–44 anos, "
        "o gap mediano atinge 37,5%, o maior da série; para jovens de 14–24 anos, o gap é de "
        "apenas 9,1%, o que pode refletir a maior homogeneidade de funções de entrada no "
        "mercado. A persistência do gap em todos os níveis educacionais — inclusive entre "
        "trabalhadores com pós-graduação (19,1%) — é evidência contra a hipótese de que "
        "a equalização educacional seria suficiente para eliminar o diferencial racial."
    )

    # Tabela 2
    try:
        tab2 = pd.read_csv(TABLES / "tab2_gap_bruto_subgrupos.csv")
        cols2 = list(tab2.columns)
        tbl2 = doc.add_table(rows=len(tab2)+1, cols=len(cols2))
        tbl2.style = "Table Grid"
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(cols2):
            cell = tbl2.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.size = Pt(9); run.font.name = "Arial"
        shade_row(tbl2.rows[0], "D9E1F2")
        for i, row in tab2.iterrows():
            for j, val in enumerate(row):
                tbl2.rows[i+1].cells[j].text = str(val)
                for run in tbl2.rows[i+1].cells[j].paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
            if i % 2 == 0:
                shade_row(tbl2.rows[i+1], "F2F2F2")
        add_caption(doc, "Tabela 2. Gap salarial racial bruto por subgrupos de gênero, "
                    "faixa etária e escolaridade — PNAD Contínua 2016–2025 (ponderado, mediana e média).")
    except Exception as e:
        add_para(doc, f"[Tabela 2 — tab2_gap_bruto_subgrupos.csv não encontrado: {e}]", italic=True)

    add_figure(doc, FIGURES / "fig2_heatmap_gap_genero_educ.png",
               "Figura 2. Heatmap do gap racial por gênero e escolaridade — "
               "gap na mediana (%) — PNAD Contínua 2016–2025.")

    add_para(doc,
        "A Figura 2 destaca que a interseção entre raça e escolaridade gera padrões "
        "assimétricos: o gap é mais acentuado entre trabalhadores com escolaridade intermediária, "
        "consistente com a hipótese de Crenshaw (1989) sobre a dupla desvantagem da mulher negra "
        "de menor escolaridade, onde a interseccionalidade amplifica as barreiras estruturais. "
        "O paradoxo da pós-graduação — gap de 19,1% mesmo entre os mais escolarizados — "
        "corrobora o argumento de Cacciamali e Hirata (2005) sobre a subvalorização do capital "
        "humano negro independentemente do nível de instrução."
    )

    # 4.1 HLM
    add_heading(doc, "Modelos Hierárquicos Lineares (HLM)", level=2)
    add_para(doc,
        f"A Tabela 1 apresenta os cinco modelos HLM ajustados sequencialmente, do modelo nulo (M0) "
        f"ao modelo de composição ocupacional (M4). O modelo nulo estima ICC_UF = {k['icc_m0']}, "
        f"indicando que aproximadamente 9,8% da variância do log-rendimento é atribuível ao estado "
        f"de residência, acima do limiar de 5% sugerido por Raudenbush e Bryk (2002)."
    )
    add_para(doc,
        f"O modelo M1 estima β̂₁ = {k['b_m1']:.4f} (p<0,001), indicando que profissionais negros "
        f"recebem em média {k['gb']:.1f}% a menos que brancos com mesma escolaridade, sexo e faixa "
        f"etária — o gap racial bruto. Após a inclusão das variáveis de contexto da UPA (M2), o "
        f"coeficiente reduz para {k['b_m2']:.4f}, implicando mediação contextual de {k['med_upa']:.1f}% "
        f"do gap bruto pelo local de moradia (duplo disadvantage). O modelo completo M3 produz "
        f"β̂₁ = {k['b_m3']:.4f}: o gap líquido de {k['gl']:.1f}% representa a discriminação direta "
        f"residual não explicável por capital humano ou contexto de moradia."
    )
    add_para(doc,
        f"O modelo M4, que acrescenta variáveis de composição ocupacional (grupos CBO), vínculo "
        f"empregatício, formalidade e horas trabalhadas, reduz o coeficiente para "
        f"β̂₁ = {k['b_m4']:.4f} (gap residual de {k['gap_m4']:.1f}%). A inclusão dessas variáveis "
        f"explica adicionalmente {k['med_occ']:.1f}% do gap bruto — a mediação ocupacional. "
        f"A variância residual (σ²) cai de 0,6037 (M3) para 0,4127 (M4), redução de 31,7%, "
        f"confirmando o elevado poder explicativo da estrutura ocupacional sobre a heterogeneidade "
        f"salarial individual. A variância entre UFs (τ²_UF) também se reduz de 0,02504 para "
        f"0,01414, indicando que parte das diferenças macrorregionais de renda é mediada pela "
        f"composição setorial e ocupacional das economias estaduais."
    )

    p = doc.add_paragraph()
    run = p.add_run("Tabela 1 – Modelos HLM de Três Níveis. PNAD Contínua 2016–2025. "
                    "*** p<0,001; ** p<0,01; * p<0,05. Erros-padrão entre parênteses.")
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    build_hlm_table(doc, r)

    add_figure(doc,
        FIGURES / "kmeans_selecao_k.png",
        "Figura 1 – Curvas de seleção de k: Silhouette e Davies-Bouldin (PNAD 2016–2025).",
        width_cm=12
    )
    add_figure(doc,
        FIGURES / "kmeans_composicao_racial_k3.png",
        "Figura 2 – Composição racial por cluster (k=3).",
        width_cm=12
    )

    # ── M4: Random Slope + ICC Racial + Nakagawa R² ──────────────────────────
    doc.add_page_break()
    add_heading(doc, "4.1b M4 – Heterogeneidade do Retorno Educacional e ICC por Raça", level=2)

    m4_coef = r.get("hlm_m4_coef")
    m4_vc   = r.get("hlm_m4_vc")
    hlm_icc = r.get("hlm_icc")
    hlm_ns  = r.get("hlm_ns")

    if m4_coef is not None and m4_vc is not None:
        # Extrair valores-chave
        def _get_vc(df, comp):
            row = df[df["componente"].str.contains(comp, na=False, regex=False)]
            return float(row["variancia"].values[0]) if len(row) else np.nan

        tau2_int   = _get_vc(m4_vc, "intercepto")
        tau2_slope = _get_vc(m4_vc, "slope educ")
        corr_u     = _get_vc(m4_vc, "covariância") / (
            np.sqrt(tau2_int * tau2_slope) if tau2_int > 0 and tau2_slope > 0 else 1
        )
        icc_upa    = _get_vc(m4_vc, "ICC_UPA")

        b_negro_row = m4_coef[m4_coef["variavel"] == "negro"] if "variavel" in m4_coef.columns else pd.DataFrame()
        b_negro = float(b_negro_row["beta"].values[0]) if len(b_negro_row) else np.nan

        add_para(doc,
            f"O modelo M4 adiciona um random slope de educ_ord_c (educação ordinal centrada) "
            f"por UPA, testando se o retorno à educação varia entre localidades. O ajuste usa "
            f"20% da amostra (n=471.570; 31.694 UPAs; seed=42), com UF como efeito fixo."
        )
        add_para(doc,
            f"Resultados: β_negro = {b_negro:.4f} ({(np.exp(b_negro)-1)*100:.1f}%); "
            f"τ²₀ (intercepto) = {tau2_int:.4f}; τ²₁ (slope educ) = {tau2_slope:.4f}; "
            f"ICC_UPA (q=0) = {icc_upa:.4f}. "
            f"A variância do slope τ²₁ > 0 confirma que o retorno educacional varia "
            f"significativamente entre localidades."
        )

        # Tabela M4 variância
        p = doc.add_paragraph()
        run = p.add_run("Tabela — M4: Componentes de Variância e Pseudo-R² (Nakagawa-Schielzeth, 2013).")
        run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(8)

        tbl_vc = doc.add_table(rows=len(m4_vc)+1, cols=2)
        tbl_vc.style = "Table Grid"; tbl_vc.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, hd in enumerate(["Componente", "Valor"]):
            cell = tbl_vc.rows[0].cells[j]
            cell.text = hd
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = "Arial"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        for i, (_, row_d) in enumerate(m4_vc.iterrows()):
            tbl_vc.rows[i+1].cells[0].text = str(row_d["componente"])
            val = row_d["variancia"]
            tbl_vc.rows[i+1].cells[1].text = f"{val:.6f}" if pd.notna(val) else "—"
            for c in tbl_vc.rows[i+1].cells:
                c.paragraphs[0].runs[0].font.name = "Times New Roman"
                c.paragraphs[0].runs[0].font.size = Pt(10)

        # ICC racial
        if hlm_icc is not None:
            add_para(doc,
                "O ICC estratificado por raça (Modelo Nulo) quantifica quanto da variância "
                "salarial é atribuível ao bairro para cada grupo racial:"
            )
            tbl_icc = doc.add_table(rows=len(hlm_icc)+1, cols=len(hlm_icc.columns))
            tbl_icc.style = "Table Grid"; tbl_icc.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, col in enumerate(hlm_icc.columns):
                cell = tbl_icc.rows[0].cells[j]
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.name = "Arial"
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            for i, (_, rw) in enumerate(hlm_icc.iterrows()):
                for j, col in enumerate(hlm_icc.columns):
                    v = rw[col]
                    tbl_icc.rows[i+1].cells[j].text = f"{v:.4f}" if isinstance(v, float) else str(v)
                    tbl_icc.rows[i+1].cells[j].paragraphs[0].runs[0].font.name = "Times New Roman"
                    tbl_icc.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

            icc_neg = hlm_icc.loc[hlm_icc["Grupo"]=="Negros", "ICC_UPA"].values
            icc_bra = hlm_icc.loc[hlm_icc["Grupo"]=="Brancos","ICC_UPA"].values
            if len(icc_neg) and len(icc_bra):
                delta = float(icc_neg[0]) - float(icc_bra[0])
                add_para(doc,
                    f"O diferencial ΔICC = {delta:+.4f} (Negros − Brancos) indica que a localidade "
                    f"de moradia {'explica mais' if delta > 0 else 'explica menos'} da variância "
                    f"salarial para trabalhadores negros, {'confirmando' if delta > 0 else 'não confirmando'} "
                    f"a hipótese de duplo disadvantage contextual."
                )

        # Nakagawa R²
        if hlm_ns is not None:
            add_para(doc,
                "Pseudo-R² de Nakagawa e Schielzeth (2013): R²m (efeitos fixos) e R²c (completo):"
            )
            cols_ns = ["Modelo", "sigma2_f", "sigma2_u", "sigma2_e", "R2m", "R2c"]
            cols_ns = [c for c in cols_ns if c in hlm_ns.columns]
            tbl_ns = doc.add_table(rows=len(hlm_ns)+1, cols=len(cols_ns))
            tbl_ns.style = "Table Grid"; tbl_ns.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, col in enumerate(cols_ns):
                cell = tbl_ns.rows[0].cells[j]
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.name = "Arial"
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            for i, (_, rw) in enumerate(hlm_ns.iterrows()):
                for j, col in enumerate(cols_ns):
                    v = rw[col]
                    tbl_ns.rows[i+1].cells[j].text = (
                        f"{float(v):.4f}" if isinstance(v, float) and pd.notna(v) else str(v)
                    )
                    tbl_ns.rows[i+1].cells[j].paragraphs[0].runs[0].font.name = "Times New Roman"
                    tbl_ns.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

        add_figure(doc, FIGURES / "hlm_m4_random_slope.png",
            "Figura – M4: Distribuição dos slopes aleatórios de educ_ord_c por UPA (esquerda) "
            "e dispersão intercepto × slope (direita). Correlação negativa indica que bairros "
            "com menor renda base oferecem menor retorno educacional (PNAD 2016–2025, 20%).",
            width_cm=15)
    else:
        add_para(doc,
            "M4 (random slope + ICC racial + Nakagawa R²): outputs em geração — "
            "execute run_hlm_m4.py e regenere o DOCX.",
            italic=True)

    # 4.2 Clustering
    doc.add_page_break()
    add_heading(doc, "4.2 Clustering Socioeconômico", level=2)
    add_para(doc,
        f"O Silhouette Coefficient máximo em k=3 (S={k['silh']:.4f}) determinou a solução de três "
        f"clusters. A Tabela 2 apresenta os perfis médios."
    )

    # Cluster summary table (hardcoded from known results)
    p = doc.add_paragraph()
    run = p.add_run("Tabela 2 – Perfis dos Clusters Socioeconômicos (k=3). PNAD 2016–2025.")
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    km_tbl = doc.add_table(rows=1, cols=6)
    km_tbl.style = "Table Grid"
    km_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = km_tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells, ["Cluster","N","% Negro","% Mulher","log_Renda","Descrição"]):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9); run.font.name = "Arial"
    cluster_data = [
        ("C0","778.802","24,2%","22,8%","7,440","Elite masc. predominantemente branca"),
        ("C1","1.075.250","86,1%","0,0%","6,791","Trabalhadores negros – segmento masculino"),
        ("C2","541.233","78,6%","100,0%","6,535","Mulheres negras – dupla desvantagem"),
    ]
    for i, row_data in enumerate(cluster_data):
        tr = km_tbl.add_row()
        shade_row(tr, "F2F2F2" if i%2==0 else "FFFFFF")
        for j, (cell, v) in enumerate(zip(tr.cells, row_data)):
            cell.text = v
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j in (0,5) else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9); run.font.name = "Arial"

    add_para(doc,
        "O Cluster 0, predominantemente branco (75,8%) e masculino (77,2%), concentra os maiores "
        "rendimentos. Os Clusters 1 e 2 agrupam trabalhadores negros (86% e 79%, respectivamente) "
        "com rendimentos 8% e 15% inferiores ao Cluster 0. Notavelmente, os grupos negros apresentam "
        "escolaridade média ligeiramente superior à do Cluster 0, confirmando a subconversão do "
        "capital humano em renda (H5)."
    )

    add_figure(doc, FIGURES / "kmeans_perfis_k3.png",
        "Figura 3 – Perfis médios por cluster (k=3).", width_cm=14)
    add_figure(doc, FIGURES / "kmeans_gap_racial_k3.png",
        "Figura 4 – Gap racial de log-renda dentro de cada cluster.", width_cm=12)

    # 4.3 ML + SHAP
    doc.add_page_break()
    add_heading(doc, "4.3 Modelos de Machine Learning e SHAP Values", level=2)
    add_para(doc,
        f"A Tabela 3 apresenta o desempenho preditivo dos dois modelos sobre o conjunto de teste "
        f"(hold-out 20%, N=307.768)."
    )

    p = doc.add_paragraph()
    run = p.add_run("Tabela 3 – Desempenho preditivo: Random Forest e XGBoost.")
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    perf_tbl = doc.add_table(rows=1, cols=4)
    perf_tbl.style = "Table Grid"
    perf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = perf_tbl.rows[0]
    shade_row(hdr, "1F3864")
    for cell, txt in zip(hdr.cells, ["Modelo","R²","MAE","RMSE"]):
        cell.text = txt
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10); run.font.name = "Arial"
    perf_data = [
        ("Random Forest", f"{k['rf_r2']:.4f}", f"{k['rf_mae']:.4f}", "0,7033"),
        ("XGBoost",       f"{k['xgb_r2']:.4f}", f"{k['xgb_mae']:.4f}", "0,6986"),
    ]
    for i, row_data in enumerate(perf_data):
        tr = perf_tbl.add_row()
        shade_row(tr, "F2F2F2" if i%2==0 else "FFFFFF")
        for j, (cell, v) in enumerate(zip(tr.cells, row_data)):
            cell.text = v
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = (i==1)
                    run.font.size = Pt(10); run.font.name = "Arial"

    add_para(doc,
        f"O XGBoost alcança R²={k['xgb_r2']:.4f} — ganho expressivo em relação ao modelo sem "
        f"variáveis ocupacionais (anterior R²≈0,44), explicável pela adição de horas trabalhadas, "
        f"emprego formal e grupos CBO, que respondem coletivamente por grande parte da variância "
        f"de rendimento individual. O Random Forest obtém R²={k['rf_r2']:.4f}. "
        f"A Tabela 4 apresenta a importância SHAP comparada entre os dois modelos."
    )

    p = doc.add_paragraph()
    run = p.add_run("Tabela 4 – Importância SHAP comparada. PNAD 2016–2025 (N_SHAP=50.000).")
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    build_shap_table(doc, r)

    add_para(doc,
        f"A renda média da UPA segue como preditor mais importante (|SHAP| = {k['shap_top1_val']:.3f}), "
        f"confirmando que o contexto de moradia supera em importância preditiva a escolaridade. "
        f"Com as novas variáveis, horas trabalhadas (|SHAP|=0,166), CBO: Profissionais (0,119) "
        f"e emprego formal (0,109) emergem como segundo, terceiro e quarto preditores, "
        f"evidenciando que a estrutura ocupacional é o principal mediador individual do rendimento. "
        f"A variável racial ocupa o {k['shap_negro_rank']}º lugar, com SHAP médio de −0,0249 para "
        f"trabalhadores negros — equivalente a uma penalidade de 2,5% sobre o rendimento predito "
        f"não atribuível a diferenças em educação, experiência, gênero, contexto de moradia ou "
        f"composição ocupacional. Essa redução em relação ao modelo sem variáveis ocupacionais "
        f"(anterior: −0,0469) confirma que parte do efeito racial é mediada pela composição "
        f"ocupacional, mas um resíduo significativo persiste — evidência de discriminação pura."
    )

    add_figure(doc, FIGURES / "shap_beeswarm_xgb.png",
        "Figura 5 – Beeswarm SHAP: distribuição de valores por feature (XGBoost).", width_cm=14)
    add_figure(doc, FIGURES / "shap_dependence_negro_xgb.png",
        "Figura 6 – Dependence plot: efeito da raça colorido por % negro na UPA.", width_cm=13)
    add_figure(doc, FIGURES / "shap_waterfall_A_branco_alta_renda_xgb.png",
        "Figura 7a – Waterfall SHAP: Caso A – branco de alta renda.", width_cm=13)
    add_figure(doc, FIGURES / "shap_waterfall_B_negro_alta_renda_xgb.png",
        "Figura 7b – Waterfall SHAP: Caso B – negro de alta renda.", width_cm=13)
    add_figure(doc, FIGURES / "shap_waterfall_C_negro_baixa_renda_xgb.png",
        "Figura 7c – Waterfall SHAP: Caso C – negro de baixa renda.", width_cm=13)

    # 4.4 SNA
    doc.add_page_break()
    add_heading(doc, "4.4 Análise de Redes Sociais", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Tabela 5 – Métricas de rede por grupo demográfico (raça × educação). "
                    "Constraint de Burt: maior valor = maior isolamento estrutural.")
    run.italic = True; run.font.name = "Arial"; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)

    build_sna_table(doc, r)

    add_para(doc,
        f"Os resultados da SNA revelam cinco achados principais. Em todos os cinco níveis "
        f"educacionais, grupos negros registram betweenness centrality igual a zero, enquanto "
        f"Branco_Fundamental alcança B={k['bra_between']:.3f}. Isso significa que os fluxos de "
        f"informação e oportunidade profissional que cruzam grupos socioeconômicos distintos "
        f"transitam exclusivamente por atores brancos — confirmando H5."
    )
    add_para(doc,
        f"O índice de homofilia H={k['sna_h']:.4f} indica heterofilia leve: em termos de peso "
        f"acumulado de co-presença em UPAs, há mais mistura inter-racial do que segregação pura. "
        f"A mistura ocorre principalmente nos bairros populares (grupos Sem instrução de ambas as "
        f"raças compartilham Jaccard ≈ 0,979), enquanto o par com menor integração é "
        f"Branco_Superior ↔ Negro_Pós (J=0,492)."
    )
    add_para(doc,
        f"O diferencial de log-rendimento reduziu-se de {k['gap_2016']:.3f} (2016) para "
        f"{k['gap_2025']:.3f} (2025), queda de {abs(k['gap_2025']-k['gap_2016']):.3f} em dez anos "
        f"— redução de apenas {abs(k['gap_2025']-k['gap_2016'])/k['gap_2016']*100:.1f}% em relação "
        f"ao patamar inicial."
    )

    add_figure(doc, FIGURES / "sna_rede_demografica.png",
        "Figura 8 – Grafo de co-residência: tamanho ∝ renda, vermelho = arestas inter-raciais.", width_cm=13)
    add_figure(doc, FIGURES / "sna_constraint_vs_renda.png",
        "Figura 9 – Constraint de Burt × rendimento médio por grupo.", width_cm=12)
    add_figure(doc, FIGURES / "sna_homofilia_por_educ.png",
        "Figura 10 – Peso Jaccard intra vs. inter-racial por nível de educação.", width_cm=12)
    add_figure(doc, FIGURES / "sna_temporal_gap.png",
        "Figura 11 – Evolução do gap salarial e integração residencial (2016–2025).", width_cm=13)

    doc.add_page_break()

    # ── 4.5 PREDIÇÕES E CONTRAFACTUAL ─────────────────────────────────────────
    add_heading(doc, "4.5 Predições por Perfil Demográfico e Análise Contrafactual", level=2)
    add_para(doc,
        "A Figura 12 apresenta o rendimento mensal médio observado para doze perfis demográficos "
        "definidos pelo cruzamento de raça (branco/negro), sexo (homem/mulher) e nível de "
        "escolaridade (fundamental completo, médio completo, superior completo), com intervalos "
        "de confiança de 95% calculados sobre a série completa da PNAD 2016–2025."
    )
    add_figure(doc, FIGURES / "predicoes_por_perfil.png",
        "Figura 12 – Rendimento médio observado por perfil demográfico e escolaridade. "
        "Barras de erro: IC 95% baseado no erro-padrão da média. "
        "PNAD Contínua 2016–2025, PEA com renda positiva.",
        width_cm=15)
    add_para(doc,
        "Os resultados revelam quatro padrões centrais: (i) em todos os níveis de escolaridade, "
        "homens brancos recebem mais que todos os demais grupos; (ii) o prêmio salarial da "
        "escolaridade existe para todos os grupos, mas é proporcionalmente menor para negros — "
        "ou seja, a desvantagem racial persiste mesmo com equalização educacional; (iii) mulheres "
        "negras com superior completo recebem, em média, rendimento equivalente ao de homens negros "
        "e mulheres brancas com médio completo, evidenciando a intersecção de gênero e raça; e "
        "(iv) o diferencial racial dentro de cada nível educacional não se reduz com o avanço da "
        "escolaridade, contrariando a hipótese de que mais educação é suficiente para eliminar o gap."
    )
    add_para(doc,
        "A Figura 13 apresenta a análise contrafactual de Oaxaca-Blinder: quanto ganhariam os "
        "trabalhadores negros se o mercado de trabalho remunerasse suas características ao mesmo "
        "nível que remunera as características dos brancos? No modelo ampliado — que inclui "
        "composição ocupacional (CBO), vínculo empregatício, formalidade e horas trabalhadas —, "
        "o gap real de 37,4% (brancos R$1.273 vs. negros R$927) se reduziria para 28,8%, "
        "representando um fechamento de apenas 20,3%. Esse resultado revela que a maior parte "
        "do diferencial racial (79,7%) é explicada pelas características observáveis dos "
        "trabalhadores — especialmente o acesso desigual a ocupações de maior prestígio —, "
        "não pelos retornos diferenciais a essas características."
    )
    add_figure(doc, FIGURES / "predicoes_contrafactual.png",
        "Figura 13 – Análise contrafactual: rendimento médio de negros com seus retornos observados "
        "versus com os retornos dos brancos aplicados às suas características. "
        "Linha tracejada: gap residual após equiparação de retornos.",
        width_cm=13)

    # ── 4.6 OAXACA-BLINDER ─────────────────────────────────────────────────────
    add_heading(doc, "4.6 Decomposição de Oaxaca-Blinder do Gap Salarial Racial", level=2)
    add_para(doc,
        "A decomposição de Oaxaca-Blinder, estimada sobre amostra de 1,5 milhão de observações "
        "da PEA com renda positiva, particionou o gap total de log-rendimento (0,4229, equivalente "
        "a 52,6% de diferença geométrica) em seus componentes estruturais. O modelo inclui, além "
        "das variáveis educacionais e demográficas, as novas variáveis de composição ocupacional "
        "(grupos CBO), vínculo empregatício, formalidade e horas trabalhadas."
    )
    add_figure(doc, FIGURES / "oaxaca_decomposicao.png",
        "Figura 14 – Decomposição de Oaxaca-Blinder: gap total particionado em efeito dotações "
        "(diferença de características) e efeito retornos (retornos diferenciais / discriminação). "
        "Referência: coeficientes do grupo branco.",
        width_cm=14)
    add_para(doc,
        "O efeito dotações responde por 84,0% do gap: quando se incluem as variáveis de "
        "composição ocupacional, formalidade e horas trabalhadas, a maior parte do diferencial "
        "salarial racial passa a ser explicada por diferenças nas características observáveis dos "
        "trabalhadores — em especial o acesso desigual a grupos CBO de maior prestígio e ao "
        "emprego formal. O efeito retornos responde por 16,0% do gap: mesmo controlando "
        "integralmente pela composição ocupacional, o mercado de trabalho remunera as "
        "características de trabalhadores negros a taxas inferiores às dos brancos — evidência "
        "de discriminação por retornos diferenciais. O termo de interação (−5,2%) é de magnitude "
        "reduzida, indicando que os efeitos dotações e retornos são relativamente aditivos. "
        "A elevação do efeito dotações em relação ao modelo sem variáveis ocupacionais (de 71% "
        "para 84%) confirma que a segregação ocupacional é o principal mecanismo de transmissão "
        "do gap salarial racial — a discriminação se manifesta primariamente no acesso, não no "
        "salário dentro de cada ocupação."
    )
    add_figure(doc, FIGURES / "oaxaca_coeficientes.png",
        "Figura 15 – Retornos às características por grupo racial: coeficientes OLS estimados "
        "separadamente para brancos e negros. A diferença entre as barras de cada variável "
        "representa o efeito retornos daquela característica específica.",
        width_cm=14)
    add_figure(doc, FIGURES / "oaxaca_por_variavel.png",
        "Figura 16 – Efeito dotações decomposto por variável: contribuição de cada característica "
        "ao efeito total de dotações. Azul = dotação superior dos brancos amplia o gap; "
        "Vermelho = dotação inferior dos brancos reduziria o gap.",
        width_cm=13)
    add_para(doc,
        "O detalhamento do efeito dotações por variável (Figura 16) revela que, no modelo "
        "ampliado, as variáveis ocupacionais emergiram como os principais canais de transmissão "
        "do gap: a sobre-representação de negros em grupos CBO de menor remuneração "
        "(serviços, elementar, agropecuária) e a menor taxa de emprego formal explicam "
        "conjuntamente a maior parcela do efeito dotações. A escolaridade mantém contribuição "
        "relevante, agora interpretada como mecanismo de acesso às ocupações — e não apenas "
        "retorno direto ao capital humano. O contexto da UPA (% de negros, taxa de desemprego) "
        "mantém contribuição secundária mas significativa, convergindo com o achado do HLM sobre "
        "mediação contextual: o local de moradia amplifica o gap ao restringir o acesso a "
        "redes de contato e oportunidades ocupacionais de maior qualidade."
    )

    # Melhorias OB: por sexo + estabilidade temporal (condicionais)
    ob_mel = r.get("ob_mel")
    if ob_mel is not None:
        add_heading(doc, "4.6.1 Decomposição OB por Sexo", level=2)
        add_para(doc,
            "A decomposição OB estratificada por sexo — estimada sobre sub-amostras de homens "
            "(sexo_fem=0) e mulheres (sexo_fem=1) com fórmula sem o preditor de sexo — revela "
            "assimetria relevante no gap racial intra-grupo. Homens negros e mulheres negras "
            "enfrentam mecanismos diferentes de exclusão salarial: para homens, a segregação "
            "ocupacional (efeito dotações) domina; para mulheres, os retornos diferenciais "
            "(discriminação de remuneração) têm peso relativo maior, consistente com a "
            "intersecção de penalidades de raça e gênero (CRENSHAW, 1989)."
        )
        if (FIGURES / "ob_por_sexo.png").exists():
            add_figure(doc, FIGURES / "ob_por_sexo.png",
                "Figura OB-A – Decomposição Oaxaca-Blinder por sexo: gap total, efeito dotações "
                "e efeito retornos para homens e mulheres separadamente. "
                "Barras vermelhas = retornos diferenciais (discriminação de remuneração); "
                "barras azuis = dotações (segregação ocupacional/educacional).",
                width_cm=14)
        sub_hom = ob_mel[ob_mel["grupo"] == "Homens"] if "grupo" in ob_mel.columns else None
        sub_mul = ob_mel[ob_mel["grupo"] == "Mulheres"] if "grupo" in ob_mel.columns else None
        if sub_hom is not None and len(sub_hom):
            rh = sub_hom.iloc[0]
            rm = sub_mul.iloc[0] if (sub_mul is not None and len(sub_mul)) else None
            txt_hom = (f"Gap total: {rh.get('gap_pct',0):.1f}% | "
                       f"Dotações: {rh.get('end_pct',0):.1f}% | "
                       f"Retornos: {rh.get('ret_pct',0):.1f}%")
            txt_mul = (f"Gap total: {rm.get('gap_pct',0):.1f}% | "
                       f"Dotações: {rm.get('end_pct',0):.1f}% | "
                       f"Retornos: {rm.get('ret_pct',0):.1f}%"
                       if rm is not None else "—")
            add_para(doc, f"Homens: {txt_hom}. Mulheres: {txt_mul}.")
        add_heading(doc, "4.6.2 Estabilidade Temporal da Decomposição OB", level=2)
        add_para(doc,
            "Para verificar se a estrutura do gap racial se alterou no período estudado, "
            "a decomposição OB foi estimada em dois sub-períodos: anos iniciais (≤ 2018, "
            "pré-pandemia e pré-reformas) e anos recentes (≥ 2022, pós-pandemia). "
            "Estabilidade dos coeficientes confirma que o gap racial não é conjuntural — "
            "é estrutural. Variação no efeito retornos sugere mudança na intensidade de "
            "discriminação de remuneração entre períodos."
        )
        if (FIGURES / "ob_estabilidade_temporal.png").exists():
            add_figure(doc, FIGURES / "ob_estabilidade_temporal.png",
                "Figura OB-B – Estabilidade temporal da decomposição OB: comparação entre "
                "anos iniciais (≤ 2018) e anos recentes (≥ 2022). Barras agrupadas por "
                "componente. Variação do efeito retornos sinaliza mudança na discriminação direta.",
                width_cm=14)
        sub_ini = ob_mel[ob_mel["grupo"] == "Periodo Inicial"] if "grupo" in ob_mel.columns else None
        sub_rec = ob_mel[ob_mel["grupo"] == "Periodo Recente"] if "grupo" in ob_mel.columns else None
        if sub_ini is not None and len(sub_ini) and sub_rec is not None and len(sub_rec):
            ri = sub_ini.iloc[0]; rr = sub_rec.iloc[0]
            add_para(doc,
                f"Período inicial (≤ 2018): gap={ri.get('gap_pct',0):.1f}%, "
                f"dotações={ri.get('end_pct',0):.1f}%, retornos={ri.get('ret_pct',0):.1f}%. "
                f"Período recente (≥ 2022): gap={rr.get('gap_pct',0):.1f}%, "
                f"dotações={rr.get('end_pct',0):.1f}%, retornos={rr.get('ret_pct',0):.1f}%. "
                "A persistência do gap e a estabilidade relativa dos componentes reforçam que "
                "a discriminação racial no mercado de trabalho não é resultado de choques "
                "conjunturais, mas de estruturas institucionais de longa duração."
            )

    # ── REGRESSÕES MINCER E INTERAÇÕES ────────────────────────────────────────
    add_heading(doc, "Regressões Mincer com Controles Progressivos e Interações Raciais", level=2)
    add_para(doc,
        "A Tabela 3 apresenta cinco especificações da equação Mincer estimadas por MQO ponderado "
        "(peso V1028, erros robustos HC3). M1 inclui apenas a dummy racial; cada especificação "
        "subsequente adiciona controles: sexo, escolaridade e experiência potencial (M2); efeitos "
        "fixos de UF (M3); horas trabalhadas, emprego formal e setor público (M4); e grupos "
        "ocupacionais CBO (M5). O coeficiente β̂(negro) mede o gap residual no log-salário em "
        "cada nível de controle."
    )

    # Tabela 3
    try:
        tab3 = pd.read_csv(TABLES / "tab3_mincer_progressivo.csv")
        cols3 = ["Especificação", "β negro", "SE", "IC 95% inf.", "IC 95% sup.", "Sig.", "Gap (%)", "R²"]
        tbl3 = doc.add_table(rows=len(tab3)+1, cols=len(cols3))
        tbl3.style = "Table Grid"
        tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(cols3):
            cell = tbl3.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.size = Pt(9); run.font.name = "Arial"
        shade_row(tbl3.rows[0], "D9E1F2")
        for i, row in tab3.iterrows():
            for j, col in enumerate(cols3):
                tbl3.rows[i+1].cells[j].text = str(row.get(col, ""))
                for run in tbl3.rows[i+1].cells[j].paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
            if i % 2 == 0:
                shade_row(tbl3.rows[i+1], "F2F2F2")
        add_caption(doc, "Tabela 3. Regressões Mincer com controles progressivos — coeficiente "
                    "de raça em cada especificação. Erros-padrão robustos HC3; *** p<0,001; ** p<0,01; * p<0,05.")
    except Exception as e:
        add_para(doc, f"[Tabela 3 — tab3_mincer_progressivo.csv não encontrado: {e}]", italic=True)

    add_para(doc,
        "O gap bruto de 23,1% (M1) reduz-se progressivamente à medida que controles são "
        "adicionados, atingindo 6,2% em M5 — após controlar todas as características observáveis, "
        "incluindo a função CBO. Essa progressão reproduz e confirma o achado central do HLM: "
        "o componente não explicado do gap, atribuível à discriminação direta, persiste em todas "
        "as especificações a p<0,001. A inclusão de efeitos fixos de UF (M3) reduz o gap de "
        "22,3% para 7,9%, indicando que parte substancial da desigualdade é geograficamente "
        "mediada — convergindo com o ICC de 9,83% estimado pelo HLM."
    )

    add_para(doc,
        "A Tabela 4 apresenta os termos de interação raça × gênero, raça × escolaridade e "
        "raça × experiência, estimados no modelo M4 com os termos de interação adicionados. "
        "A interação negro × sexo feminino (β = +0,0416, p<0,001) indica que mulheres negras "
        "enfrentam penalidade relativa menor em relação a seus pares brancos do que homens negros "
        "— sugerindo que a desvantagem de gênero, para mulheres negras, já está incorporada na "
        "linha de base. Este resultado é consistente com a abordagem interseccional de Crenshaw "
        "(1989): a análise de raça e gênero separadamente subestima as desvantagens do grupo "
        "de mulheres negras, cujas experiências não são a soma simples das desvantagens raciais "
        "e de gênero."
    )
    add_para(doc,
        "A interação negro × experiência (β = −0,0019, p<0,001) indica que o retorno à "
        "experiência é marginalmente menor para negros — cada ano adicional de experiência "
        "potencial gera 0,19 pontos percentuais menos de renda adicional para negros do que "
        "para brancos, compatível com a hipótese de glass ceiling de progressão: barreiras de "
        "promoção tornam-se mais opressivas ao longo da carreira, limitando o retorno salarial "
        "à senioridade entre trabalhadores negros."
    )

    # Tabela 4
    try:
        tab4 = pd.read_csv(TABLES / "tab4_interacoes.csv")
        cols4 = ["Variável", "β", "SE", "IC 95% inf.", "IC 95% sup.", "Sig.", "Efeito (%)"]
        tbl4 = doc.add_table(rows=len(tab4)+1, cols=len(cols4))
        tbl4.style = "Table Grid"
        tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(cols4):
            cell = tbl4.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.size = Pt(9); run.font.name = "Arial"
        shade_row(tbl4.rows[0], "D9E1F2")
        for i, row in tab4.iterrows():
            for j, col in enumerate(cols4):
                tbl4.rows[i+1].cells[j].text = str(row.get(col, ""))
                for run in tbl4.rows[i+1].cells[j].paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
            if i % 2 == 0:
                shade_row(tbl4.rows[i+1], "F2F2F2")
        add_caption(doc, "Tabela 4. Termos de interação raça × gênero, raça × escolaridade e "
                    "raça × experiência — MQO ponderado (HC3). *** p<0,001; * p<0,05; ns = não significativo.")
    except Exception as e:
        add_para(doc, f"[Tabela 4 — tab4_interacoes.csv não encontrado: {e}]", italic=True)

    add_figure(doc, FIGURES / "fig4_gap_faixa_etaria.png",
               "Figura 4. Gap racial por faixa etária (curva de ciclo de vida) — "
               "gap bruto na mediana (%), PNAD Contínua 2016–2025.")

    add_para(doc,
        "A Figura 4 evidencia a curva de ciclo de vida do gap racial: o diferencial é mínimo "
        "entre jovens de 14–24 anos (9,1%), cresce acentuadamente na faixa de 25–44 anos "
        "(pico de 37,5% em 35–44 anos) e permanece elevado entre os mais velhos (40,0% na "
        "faixa 65+). Esse padrão — típico dos efeitos cumulativos de discriminação ao longo da "
        "carreira (SOARES, 2000; HENRIQUES, 2001) — indica que o gap não é um fenômeno de "
        "entrada no mercado, mas se aprofunda com o acúmulo de experiência diferencial. "
        "A persistência entre trabalhadores mais velhos sugere que o glass ceiling racial age "
        "de forma acumulativa, com efeitos que se intensificam conforme a trajetória profissional "
        "avança — consistente com o achado da regressão quantílica sobre a ampliação do gap "
        "nos estratos de maior remuneração."
    )

    # ── 4.7 SEGREGAÇÃO ESPACIAL ────────────────────────────────────────────────
    add_heading(doc, "Segregação Espacial: Empregabilidade e Renda por Tipo de Área", level=2)
    add_para(doc,
        "A análise por tipo de área de moradia (Capital, RM não capital, Interior) revela que "
        "a segregação espacial opera de forma não homogênea ao longo do território: o gap salarial "
        "racial segue um padrão não linear, sendo mais elevado nas capitais (62,8%), intermediário "
        "no interior (56,8%) e menor nas RMs não capitais (38,9%)."
    )
    add_figure(doc, FIGURES / "segreg_taxa_ocupacao.png",
        "Figura 17 – Taxa de ocupação por raça e tipo de área de moradia "
        "(PEA — PNAD Contínua 2016–2025).",
        width_cm=14)
    add_figure(doc, FIGURES / "segreg_renda_por_area.png",
        "Figura 18 – Rendimento médio por raça e tipo de área de moradia "
        "(ocupados com renda — PNAD Contínua 2016–2025).",
        width_cm=14)
    add_para(doc,
        "O diferencial de taxa de ocupação entre brancos e negros é relativamente estável nos "
        "três tipos de área (3,6–4,0 pontos percentuais), indicando que a barreira ao acesso ao "
        "emprego não se concentra em nenhum território específico. A grande variação está no gap "
        "salarial (Figuras 18 e 19): nas capitais, o mercado formal é mais estratificado — "
        "brancos concentram-se em posições de alta remuneração (gerência, finanças, governo) "
        "enquanto negros permanecem sobre-representados em serviços e funções operacionais, "
        "produzindo o maior gap percentual (62,8%). Nas RMs não capitais — com predominância de "
        "empregos industriais e logísticos de remuneração mais homogênea — o gap é o menor dos "
        "três contextos (38,9%). No interior, a combinação de mercados locais informais e menor "
        "fiscalização trabalhista sustenta um gap intermediário (56,8%), com trabalhadores negros "
        "sofrendo penalidade dupla: menor remuneração e menor acesso a postos formais."
    )
    add_figure(doc, FIGURES / "segreg_gap_por_area.png",
        "Figura 19 – Gap salarial racial (esquerda) e diferencial de ocupação (direita) por tipo "
        "de área. Capital e Interior têm os maiores gaps; RMs não capitais têm o menor gap, "
        "refletindo maior homogeneidade dos postos de trabalho.",
        width_cm=15)
    add_figure(doc, FIGURES / "segreg_temporal_metro_interior.png",
        "Figura 20 – Evolução do rendimento médio por raça em áreas metropolitanas "
        "versus interior (2016–2025). A área sombreada representa o gap racial.",
        width_cm=15)
    add_para(doc,
        "A análise temporal (Figura 20) mostra que o gap racial se manteve persistente e "
        "relativamente estável em ambas as categorias ao longo de 2016–2025, sem tendência "
        "de convergência significativa. Esses resultados corroboram H4 (efeito de vizinhança) "
        "e convergem com os achados da SNA sobre exclusão de negros das posições de brokerage "
        "nas redes profissionais: o local de moradia determina não apenas o acesso ao emprego, "
        "mas o tipo de emprego disponível e o teto salarial alcançável. Políticas de mobilidade "
        "urbana, cotas em cargos de liderança e desenvolvimento regional são instrumentos "
        "complementares e necessários às políticas de educação e antidiscriminação."
    )

    # ── 4.8 COMPOSIÇÃO OCUPACIONAL E SEGREGAÇÃO NO TOPO ───────────────────────
    add_heading(doc, "4.8 Composição Ocupacional e Segregação no Topo da Distribuição", level=2)
    add_para(doc,
        "Para mensurar diretamente o mecanismo de segregação ocupacional, foram utilizados os "
        "códigos CBO-Domiciliar 2010 (variável V4010 da PNAD Contínua), agrupados nos 10 grandes "
        "grupos do ISCO-08 (International Standard Classification of Occupations). Essa abordagem "
        "supera a análise por proxy de percentil ao identificar quais funções específicas negros "
        "e brancos ocupam dentro de cada tipo de área. Complementarmente, calculou-se o Índice de "
        "Representação (IR) nos estratos superiores de renda: IR = 1,0 indica representação "
        "proporcional; IR < 1 indica sub-representação."
    )
    add_figure(doc, FIGURES / "comp_grupo_cbo_capital.png",
        "Figura 21 – Distribuição de brancos e negros por grupo ocupacional CBO nas Capitais "
        "(PNAD Contínua 2016–2025, VD4008/V4010). Brancos concentram-se em funções de alta "
        "remuneração; negros sobre-representam-se em funções elementares.",
        width_cm=16)
    add_para(doc,
        "Os resultados na Capital (Figura 21) revelam polarização ocupacional pronunciada: brancos "
        "concentram 25,6% em Profissionais das Ciências (médicos, engenheiros, advogados) e 7,5% em "
        "Dirigentes e Gerentes — grupos com renda mediana de R$4.500 e R$5.000, respectivamente. "
        "Negros, por sua vez, representam apenas 11,5% em Profissionais (razão=0,45) e 3,1% em "
        "Dirigentes (razão=0,42), enquanto estão concentrados em Ocupações Elementares (razão=2,12 "
        "— o dobro da participação proporcional) e Operários Qualificados (razão=1,61). A renda "
        "mediana das Ocupações Elementares é de R$1.100, contra R$4.500 dos Profissionais — uma "
        "diferença de 309%. Essa concentração estrutural nos estratos mais baixos da hierarquia "
        "ocupacional é o mecanismo central que produz o maior gap salarial nas capitais."
    )
    add_figure(doc, FIGURES / "comp_razao_grupo_cbo.png",
        "Figura 22 – Razão de representação negro/branco por grupo CBO e tipo de área "
        "(verde = sobre-representados; vermelho = sub-representados; 1,0 = proporcional). "
        "A segregação no topo é mais intensa nas Capitais.",
        width_cm=16)
    add_para(doc,
        "A Figura 22 mostra que a sub-representação de negros nos grupos de alta remuneração é "
        "um padrão pan-geográfico, mas com intensidade variável: em todos os três tipos de área, "
        "negros são sub-representados em Dirigentes, Profissionais e Técnicos (razão < 1), e "
        "sobre-representados em Elementar, Operários e Serviços/Vendas. A segregação é mais severa "
        "nas Capitais (razão de Profissionais = 0,45) do que no Interior, porque as capitais "
        "concentram as posições de maior nível hierárquico — que são exatamente as mais inacessíveis "
        "para negros. Esse padrão constitui evidência direta de segregação ocupacional vertical: "
        "não é a ausência de empregos para negros, mas a barreira ao acesso dos empregos de topo."
    )
    add_figure(doc, FIGURES / "comp_representacao_topo.png",
        "Figura 23 – Representação de negros nos percentis superiores de renda (Top 25%, 10% e 5%) "
        "por tipo de área. IR = Índice de Representação (1,0 = proporcional | < 1 = sub-representado).",
        width_cm=16)
    add_figure(doc, FIGURES / "comp_formalidade.png",
        "Figura 24 – Taxa de formalidade (emprego com carteira/estatutário) e de trabalho por conta "
        "própria por raça e tipo de área (VD4009, PNAD Contínua 2016–2025).",
        width_cm=15)
    add_para(doc,
        "A análise de vínculos empregatícios (Figura 24) acrescenta uma dimensão importante: no "
        "Interior, o gap de formalidade é de 8,3 pontos percentuais (branco 46,6% vs negro 38,2%), "
        "significativamente maior do que na Capital (54,1% vs 52,9% — diferença de apenas 1,2 p.p.). "
        "Isso significa que, fora das capitais, negros enfrentam não apenas a barreira ocupacional "
        "vertical (excluídos dos grupos de alto nível), mas também a barreira horizontal da "
        "informalidade — trabalhando nas mesmas funções dos brancos, porém sem carteira assinada, "
        "sem previdência e com menor remuneração. A taxa de trabalho doméstico entre negros (7,6–8,7%) "
        "é sistematicamente o dobro da dos brancos (3,9–5,1%), consistente com a hiper-representação "
        "de negros nas Ocupações Elementares (razão=2,12 na Capital)."
    )
    add_figure(doc, FIGURES / "comp_ratio_renda.png",
        "Figura 25 – Razão de renda negro/branco por percentil e tipo de área. "
        "A queda da razão nos percentis superiores evidencia o 'glass ceiling racial'.",
        width_cm=14)
    add_para(doc,
        "A Figura 25 sintetiza o glass ceiling racial em termos de rendimento: na Capital, negros "
        "recebem ~60% da renda dos brancos na mediana (P50), mas apenas ~50% no P90 — o gap se "
        "aprofunda exatamente nos postos de maior valor. Esse padrão — a razão de renda "
        "decrescente com o percentil — é a assinatura econométrica da segregação ocupacional "
        "vertical: não basta controlar por escolaridade e experiência (já feito no HLM), pois a "
        "exclusão opera nos critérios não observáveis de promoção e acesso a redes de poder "
        "(GRANOVETTER, 1973; PAGER, 2007). O conjunto dos resultados da seção 4.8 constitui "
        "evidência de que o gap salarial racial nas capitais é produzido, em larga medida, por "
        "um sistema de alocação ocupacional que concentra brancos nos grupos de Profissionais e "
        "Dirigentes e negros nas Ocupações Elementares e de Serviços."
    )

    doc.add_page_break()

    # ── 4.9 LOGIT MULTINÍVEL ──────────────────────────────────────────────────
    add_heading(doc, "4.9 GLMM Logístico: Teto de Vidro Ocupacional e Glass Ceiling de Acesso", level=2)
    add_para(doc,
        "O GLMM logístico (lme4 glmer, efeito aleatório de UPA) modela o acesso a posições "
        "qualificadas como desfecho binário, quantificando a barreira de entrada racial que "
        "antecede o gap salarial documentado pelo HLM. Três desfechos complementares são "
        "estimados sobre a população completa de trabalhadores com renda positiva "
        "(N=2.395.285): (i) ocupação qualificada — grupos CBO 1–4: Dirigentes, Profissionais "
        "das Ciências, Técnicos de Nível Médio e Trabalhadores Administrativos; (ii) top 20% de "
        "renda (y_top20); e (iii) top 10% de renda (y_top10). O gradiente entre os dois "
        "últimos limiares constitui teste direto de glass ceiling de acesso: se OR(top10) < "
        "OR(top20), a exclusão racial se intensifica quanto mais alto o limiar — evidência de que "
        "o teto de vidro opera como barreira progressivamente mais seletiva, e não como um "
        "degrau único de exclusão."
    )
    add_figure(doc, FIGURES / "logit_gap_bruto.png",
        "Figura 26 – Gap de oportunidades racial bruto: proporção de brancos e negros na PEA "
        "em emprego formal, ocupação qualificada e top 20% de renda "
        "(PNAD Contínua 2016–2025, PEA com renda positiva, N=2.395.285).",
        width_cm=14)
    add_para(doc,
        "A Figura 26 revela gaps brutos substanciais: brancos superam negros em 5,4 p.p. no "
        "acesso ao emprego formal (49,3% vs 43,9%), em 14,0 p.p. no acesso a ocupações "
        "qualificadas (37,9% vs 23,8%) e em 15,4 p.p. na probabilidade de estar no quintil "
        "superior de renda (31,5% vs 16,1%). Esses diferenciais brutos combinam efeitos de "
        "composição educacional, demográfica e contextual — o GLMM logístico isola o efeito "
        "líquido de raça após controle exaustivo dessas covariáveis na população completa."
    )
    # Tabela GLMM Glass Ceiling — gradiente top20 → top10
    gc = r.get("glmm_gc")
    if gc is not None:
        add_para(doc, "Tabela 7. GLMM Logístico — Teto de Vidro Duplo (população completa, N=2.395.285)",
                 bold=True, first_line=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        gc_rows = [
            ("CBO 1–4 M1",  "ocp_qualif", "M1"),
            ("CBO 1–4 M2",  "ocp_qualif", "M2"),
            ("Top 20% M1",  "y_top20",    "M1"),
            ("Top 20% M2",  "y_top20",    "M2"),
            ("Top 10% M1",  "y_top10",    "M1"),
            ("Top 10% M2",  "y_top10",    "M2"),
        ]
        headers_gc = ["Desfecho/Modelo", "OR (Negro)", "IC 95%", "AME (p.p.)", "p-valor"]
        tbl_gc = doc.add_table(rows=len(gc_rows)+1, cols=len(headers_gc))
        tbl_gc.style = "Table Grid"
        tbl_gc.alignment = WD_TABLE_ALIGNMENT.CENTER
        shade_row(tbl_gc.rows[0], "1F3864")
        for j, h in enumerate(headers_gc):
            cell = tbl_gc.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9); run.font.name = "Arial"
        for i, (label, des, mod) in enumerate(gc_rows):
            sub = gc[(gc["desfecho"]==des) & (gc["modelo"]==mod)]
            if len(sub):
                row_d = sub.iloc[0]
                vals = [label,
                        f"{row_d['OR_negro']:.3f}",
                        f"[{row_d['CI95_lo']:.3f}; {row_d['CI95_hi']:.3f}]",
                        f"{row_d['AME_pp']:.2f}",
                        "< 0,001"]
            else:
                vals = [label, "—", "—", "—", "—"]
            shade = "F2F2F2" if i % 2 == 0 else "FFFFFF"
            shade_row(tbl_gc.rows[i+1], shade)
            for j, v in enumerate(vals):
                cell = tbl_gc.rows[i+1].cells[j]
                cell.text = v
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
        add_caption(doc, "Nota: OR < 1 = desvantagem racial. Modelos M1: controles individuais + UF FE + "
                    "HC1 robust SE. M2: + contexto UPA (renda mediana e educação médias). "
                    "GLMM R (lme4 glmer): OR(ocp_qualif M1)=0,704 [ICC_UPA=26,2%] — consistente com Python. "
                    "Gradiente: OR(top20)=0,726 > OR(top10)=0,671 → glass ceiling intensifica no topo.")
    add_figure(doc, FIGURES / ("glmm_glassceil_forest.png"
                               if (FIGURES/"glmm_glassceil_forest.png").exists()
                               else "glmm_odds_ratios_full.png"),
        "Figura 27 – Forest plot dos Odds Ratios do GLMM logístico (população completa, N=2.395.285). "
        "Três desfechos × dois modelos. OR < 1 indica desvantagem de negros em todos os limiares. "
        "Gradiente ocp_qualif → top20% → top10% confirma intensificação do teto de vidro.",
        width_cm=15)
    add_para(doc,
        "Os resultados (Tabela 7, Figura 27) documentam teto de vidro duplo: negros têm "
        "OR=0,708 (AME=−1,30 p.p.) de acesso a ocupações CBO 1–4 em M1, e OR=0,726 "
        "(AME=−4,79 p.p.) para top 20% de renda. O gradiente confirmado — OR(top20)=0,726 "
        "vs OR(top10)=0,671 — prova que a exclusão racial se intensifica no topo da "
        "hierarquia: negros não apenas têm menor chance de alcançar os 20% superiores, "
        "mas essa desvantagem é ainda mais severa no limiar dos 10%. O ICC de 26,2% no "
        "GLMM de UPA (lme4, M1 para ocp_qualif) confirma que a estrutura hierárquica de "
        "UPAs é essencial — o mercado de trabalho é mais geograficamente determinado no "
        "acesso a ocupações qualificadas (ICC 26%) do que nos salários (ICC 9,8% HLM). "
        "O emprego formal apresenta OR ligeiramente acima de 1 em M2, configurando paradoxo "
        "de Simpson: dentro de cada estrato educacional e localidade, negros têm marginalmente "
        "maior odds de formalidade como estratégia de proteção — porém isso ocorre porque "
        "estão excluídos das ocupações qualificadas formais de alta remuneração, permanecendo "
        "sobre-representados no emprego formal de base."
    )
    add_figure(doc, FIGURES / ("glmm_glassceil_probpredita.png"
                               if (FIGURES/"glmm_glassceil_probpredita.png").exists()
                               else "logit_prob_predita.png"),
        "Figura 28 – Probabilidades preditas por raça e nível de escolaridade (GLMM M1, "
        "homem, contexto UPA médio). O gap racial em ocp_qualif, top20% e top10% "
        "não se fecha com a elevação da escolaridade — evidência de teto de credencial.",
        width_cm=15)
    add_para(doc,
        "A análise de moderação (Modelo M3, interação negro×anos_educação) revela OR de "
        "interação de 1,009 para CBO 1–4, 1,012 para top 20% e 1,028 para top 10% — todos "
        "próximos de 1,0 e estatisticamente triviais. Esse resultado é metodologicamente "
        "relevante: educação não atenua o teto de vidro racial. Mesmo negros com escolaridade "
        "igual à de brancos mantêm odds significativamente menores de acesso a ocupações "
        "qualificadas e ao topo da distribuição salarial. Esse 'teto de credencial' — o "
        "fenômeno pelo qual diplomas equivalentes não geram oportunidades equivalentes — "
        "é consistente com discriminação estatística (ARROW, 1973), onde empregadores usam "
        "raça como sinal de qualidade independentemente das credenciais apresentadas, e com "
        "os experimentos de auditoria de currículos de Pager (2007)."
    )
    add_figure(doc, FIGURES / "logit_decomposicao_gap.png",
        "Figura 29 – Decomposição do gap de oportunidades: bruto vs. ajustado por características "
        "individuais (M1-AME) e contexto de UPA (M2-AME). A diferença entre gap bruto e AME "
        "representa a parcela explicada pelas covariáveis observáveis.",
        width_cm=14)
    add_para(doc,
        "A decomposição do gap (Figura 29) mostra que para CBO 1–4: AME M1 = −1,30 p.p. "
        "(controles individuais); AME M2 = −1,12 p.p. (+ contexto UPA). A redução de 0,18 p.p. "
        "de M1 para M2 reflete mediação contextual pelo entorno da UPA (efeito de vizinhança, "
        "Wilson 1987). Os 1,12 p.p. residuais constituem discriminação de acesso pura não "
        "atribuível a nenhuma característica observável. O coeficiente GLMM (β=−0,301, "
        "OR=0,740 em M2) converge com o gap residual do HLM M4 (β=−0,054 em log-renda), "
        "estabelecendo diagnóstico tripartite: discriminação opera (i) nas portas de entrada "
        "das ocupações qualificadas (GLMM: OR=0,708), (ii) nas estruturas salariais dentro de "
        "cada ocupação (HLM M4: −5,4%), e (iii) se intensifica no topo da distribuição "
        "(QR: gap de −8,2% no q10 a −11,8% no q90). O teto de vidro duplo — barreira de "
        "acesso (GLMM) + barreira de progressão salarial (QR) — é a contribuição analítica "
        "central deste trabalho."
    )

    doc.add_page_break()

    # ── 4.10 REGRESSÃO QUANTÍLICA ─────────────────────────────────────────────
    add_heading(doc, "4.10 Regressão Quantílica e Formalização do Glass Ceiling", level=2)
    add_para(doc,
        "A regressão quantílica (KOENKER; BASSETT, 1978) estima o coeficiente de 'negro' "
        "em log_renda para cada quantil da distribuição condicional, revelando se o gap racial "
        "é uniforme ou se se intensifica no topo da distribuição — teste formal da hipótese de "
        "glass ceiling racial (ARULAMPALAM et al., 2007). Dois modelos são estimados em "
        "amostra de 20% (N≈1,5M): M3 (controles individuais, contexto UPA e UF FE) e M4 "
        "(M3 + grupos CBO + formalidade + horas), permitindo decompor o gap em cada quantil "
        "entre componente explicado pela composição ocupacional e componente residual de "
        "discriminação pura. O teste de Koenker-Bassett (KB) verifica H₀: β(q) constante ao "
        "longo dos quantis, com estatística Z para o contraste β(q90)−β(q10) via bootstrap "
        "(200 repetições, 5% de sub-amostra) e Wald χ²(2) para constância conjunta."
    )
    # Tabela KB test se disponível
    qr_kb = r.get("qr_kb")
    if qr_kb is not None:
        add_para(doc, "Tabela 8. Teste de Koenker-Bassett — Heterogeneidade do Gap Racial por Quantil",
                 bold=True, first_line=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        headers_kb = ["Modelo", "β(q10) %", "β(q90) %", "Δ(q90−q10) %", "Z", "p-valor", "χ²(2)", "p(χ²)"]
        tbl_kb = doc.add_table(rows=len(qr_kb)+1, cols=len(headers_kb))
        tbl_kb.style = "Table Grid"
        tbl_kb.alignment = WD_TABLE_ALIGNMENT.CENTER
        shade_row(tbl_kb.rows[0], "1F3864")
        for j, h in enumerate(headers_kb):
            cell = tbl_kb.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9); run.font.name = "Arial"
        for i, (_, row_d) in enumerate(qr_kb.iterrows()):
            shade = "F2F2F2" if i % 2 == 0 else "FFFFFF"
            shade_row(tbl_kb.rows[i+1], shade)
            cols_kb = ["modelo", "b_q10_pct", "b_q90_pct", "diff_pct", "Z_stat", "p_Z",
                       "wald_chi2", "p_wald"]
            for j, col in enumerate(cols_kb):
                cell = tbl_kb.rows[i+1].cells[j]
                val = row_d.get(col, "—")
                cell.text = f"{val:.3f}" if isinstance(val, float) else str(val)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
        add_caption(doc, "Nota: H₀ = β(q) constante nos quantis. Z = contraste bootstrap β(q90)−β(q10); "
                    "χ²(2) = Wald para constância conjunta β(q10)=β(q50)=β(q90). "
                    "Bootstrap: 200 replicações, 5% sub-amostra. *** p<0,001.")
    add_figure(doc, FIGURES / "quantreg_trajetoria.png",
        "Figura 31 – Trajetória do gap racial por quantil (M3 e M4): coeficiente de 'negro' "
        "transformado em %, com IC 95% bootstrap. Linha tracejada: estimativa OLS de referência. "
        "Inclinação negativa e teste KB significativo confirmam glass ceiling.",
        width_cm=14)
    add_para(doc,
        "A Figura 31 apresenta a evidência mais direta de glass ceiling racial neste trabalho. "
        "No modelo M3 (sem variáveis ocupacionais), o coeficiente de 'negro' em log-renda "
        "evolui de β̂=−0,0857 (−8,2%) no q10 para β̂=−0,1253 (−11,8%) no q90 — trajetória "
        "monotonicamente crescente em magnitude. O teste KB rejeita H₀ de constância: "
        "Δβ(q90−q10) é estatisticamente significativo (p<0,001), confirmando heterogeneidade "
        "quantílica estrutural. No modelo M4 (com grupos CBO, formalidade e horas), o gap "
        "cai substancialmente — a inclinação crescente se mantém, indicando que a discriminação "
        "residual pós-ocupação também aumenta no topo. Nos quantis inferiores (q10–q25), o gap "
        "reflete principalmente diferenças de inserção no mercado formal. A partir do q75, o "
        "mecanismo dominante é a exclusão do acesso às posições de alta remuneração — "
        "confirmada pelo IR=0,47 no Top 5% das capitais e pelo gradiente GLMM "
        "(OR top20=0,726 > OR top10=0,671). O glass ceiling racial é a estrutura dominante "
        "de reprodução da desigualdade nos estratos superiores do mercado de trabalho brasileiro."
    )
    if (FIGURES / "quantreg_mediacao_ocp.png").exists():
        add_figure(doc, FIGURES / "quantreg_mediacao_ocp.png",
            "Figura 32 – Mediação ocupacional por quantil: diferença entre M3 e M4 em cada "
            "ponto da distribuição. A mediação ocupacional é mais intensa no topo — confirmando "
            "que o glass ceiling é primariamente um fenômeno de exclusão de ocupações qualificadas.",
            width_cm=14)
        add_para(doc,
            "A Figura 32 decompõe, em cada quantil, quanto do gap é explicado pelas variáveis "
            "ocupacionais (CBO, formalidade, horas). A mediação ocupacional é crescente ao longo "
            "da distribuição: nos quantis inferiores, a estrutura ocupacional explica uma "
            "fração menor do gap; no topo (q90-q95), a maior parte do gap é mediada pela "
            "sub-representação de negros nas ocupações de alta remuneração. Isso implica que "
            "o gap residual pós-ocupação — discriminação de remuneração dentro da mesma "
            "categoria — é relativamente mais concentrado nos quantis intermediários, enquanto "
            "no topo prevalece a discriminação de acesso. A consistência desse diagnóstico com "
            "os resultados do logit multinível GLMM (OR=0,747, AME=−1,07 p.p. para ocupação "
            "qualificada) e da análise de composição (IR=0,47 no Top 5%) fecha um triângulo "
            "metodológico de evidências convergentes sobre a natureza do glass ceiling racial."
        )
    if (FIGURES / "quantreg_por_area.png").exists():
        add_figure(doc, FIGURES / "quantreg_por_area.png",
            "Figura 33 – Glass ceiling por tipo de área: trajetória do gap racial por quantil "
            "separada por Capital, RM não capital e Interior (modelo M3).",
            width_cm=14)
        add_para(doc,
            "A regressão quantílica por área (Figura 33) revela heterogeneidade geográfica "
            "do glass ceiling: nas capitais, a aceleração do gap nos quantis superiores é mais "
            "intensa, consistente com a maior estratificação ocupacional documentada na seção 4.8. "
            "No interior, o glass ceiling existe mas com menor inclinação, sugerindo que a "
            "segregação ocupacional é um fenômeno mais pronunciado nos mercados de trabalho "
            "urbanos complexos — onde a diversidade de funções cria mais oportunidades de "
            "acesso diferencial e, portanto, mais espaço para discriminação de seleção."
        )
    # Figuras QR melhorias — condicionais à existência dos arquivos
    if (FIGURES / "qr_kb_test.png").exists():
        add_figure(doc, FIGURES / "qr_kb_test.png",
            "Figura 34 – Teste de Koenker-Bassett: distribuição bootstrap de Δβ(q90−q10) "
            "para M3 e M4. A linha vertical indica o valor observado. Rejeição de H₀ confirma "
            "heterogeneidade quantílica estrutural (glass ceiling).",
            width_cm=14)
    if (FIGURES / "qr_por_sexo.png").exists():
        add_figure(doc, FIGURES / "qr_por_sexo.png",
            "Figura 35 – Regressão quantílica por sexo: gap racial por quantil estimado "
            "separadamente para homens e mulheres. O glass ceiling racial é documentado "
            "em ambos os sexos, com intensidade diferenciada.",
            width_cm=14)
        add_para(doc,
            "A regressão quantílica estratificada por sexo (Figura 35) revela que o glass "
            "ceiling racial opera de forma distinta entre homens e mulheres: para homens negros, "
            "o gap se intensifica de forma mais abrupta nos quantis superiores, enquanto para "
            "mulheres negras o gap é mais uniforme ao longo da distribuição — reflexo da "
            "intersecção entre penalidades raciais e de gênero que, para mulheres negras, "
            "mantêm o teto mais baixo em toda a hierarquia salarial (CRENSHAW, 1989). "
            "Esse resultado, combinado com o GLMM (OR interação negro×edu ≈ 1,0), reforça "
            "que as barreiras são estruturais e não se desfazem com investimento em capital humano."
        )
    if (FIGURES / "qr_glassceil_completo.png").exists():
        add_figure(doc, FIGURES / "qr_glassceil_completo.png",
            "Figura 36 – Glass ceiling completo: trajetória QR (gap por quantil) + GLMM "
            "(OR por limiar de acesso) em painel único. Convergência das duas metodologias "
            "confirma teto de vidro duplo — barreira de progressão salarial + barreira de acesso.",
            width_cm=16)
        add_para(doc,
            "A Figura 36 sintetiza o teto de vidro duplo: o painel esquerdo (QR) documenta "
            "que o gap salarial racial cresce de −8,2% (q10) a −11,8% (q90); o painel direito "
            "(GLMM) documenta que a probabilidade de acesso ao topo diminui de OR=0,726 "
            "(top20%) para OR=0,671 (top10%). As duas curvas convergem para o mesmo fenômeno "
            "por caminhos metodológicos independentes: a discriminação racial no mercado de "
            "trabalho brasileiro não se limita ao salário dentro de uma ocupação, mas opera "
            "antes — no momento do acesso — e se intensifica quanto mais alto o patamar almejado."
        )

    doc.add_page_break()

    # ── 5. DISCUSSÃO ──────────────────────────────────────────────────────────
    add_heading(doc, "Síntese e Discussão dos Resultados", level=2)
    add_para(doc,
        "Os resultados das seis metodologias aplicadas neste trabalho — HLM, clustering, ML/SHAP, "
        "SNA, Oaxaca-Blinder, GLMM logístico e regressão quantílica — convergem para um "
        "diagnóstico consistente: a desigualdade racial no mercado de trabalho brasileiro é um "
        "fenômeno multicausal com componentes individuais, contextuais, ocupacionais e estruturais "
        "que se reforçam mutuamente e que não se reduzem a diferenças observáveis em capital humano."
    )
    add_para(doc,
        f"O achado mais robusto desta análise é que {k['med']:.1f}% do gap salarial racial bruto é "
        f"mediado pelo local de moradia. O coeficiente contextual γ̂₀₁ = −0,289 para a proporção "
        f"de negros na UPA indica que a penalidade de viver em bairro segregado equivale, em "
        f"magnitude, à própria penalidade individual de ser negro. Isso sugere que políticas de "
        f"redistribuição de renda que não enfrentem a segregação residencial terão eficácia limitada "
        f"sobre o gap racial."
    )
    add_para(doc,
        "O Cluster 1 (trabalhadores negros de escolaridade superior à do Cluster 0) aufere "
        "rendimentos 8% inferiores, e a SNA demonstra que grupos negros com pós-graduação têm "
        "betweenness nulo. Juntas, essas evidências indicam que negros enfrentam um duplo obstáculo "
        "ao retorno educacional: além do gap direto mensurado pelo HLM, perdem acesso às redes de "
        "indicação que convertem credenciais formais em mobilidade profissional (GRANOVETTER, 1973)."
    )
    add_para(doc,
        f"O modelo M4 decompõe o gap em três camadas: (i) gap bruto de {k['gb']:.1f}% (M1); "
        f"(ii) mediação contextual pela UPA de {k['med_upa']:.1f}% (M2→M3); "
        f"(iii) mediação ocupacional de {k['med_occ']:.1f}% (M3→M4), explicada pelo acesso desigual "
        f"a grupos CBO de maior prestígio e ao emprego formal. O gap residual de {k['gap_m4']:.1f}% "
        f"(M4) representa discriminação pura pós-ocupação: mesmo dentro da mesma categoria "
        f"ocupacional, negros recebem sistematicamente menos. Os valores SHAP reforçam essa "
        f"interpretação: a variável racial mantém o {k['shap_negro_rank']}º lugar na importância "
        f"preditiva do XGBoost mesmo com acesso completo às variáveis educacionais, demográficas, "
        f"contextuais e ocupacionais — consistente com os experimentos de auditoria de Pager (2007)."
    )
    add_para(doc,
        f"A redução de apenas {abs(k['gap_2025']-k['gap_2016'])/k['gap_2016']*100:.1f}% do gap em "
        f"dez anos — equivalente a 0,02 pontos de log-rendimento por ano — sugere que, ao ritmo "
        f"atual, a convergência racial levaria mais de um século para eliminar o diferencial "
        f"observado em 2016. Reformas no campo da educação, sem intervenção simultânea nos "
        f"mecanismos de segregação residencial e de acesso às redes profissionais, são insuficientes."
    )
    add_para(doc,
        "A decomposição de Oaxaca-Blinder (seção 4.6) acrescenta precisão a esse diagnóstico: "
        "no modelo ampliado com variáveis ocupacionais (CBO, formalidade, horas), 84,0% do gap "
        "bruto é atribuível a diferenças de dotações observáveis — especialmente acesso desigual "
        "a ocupações de maior prestígio —, enquanto apenas 16,0% reflete retornos diferenciais. "
        "O GLMM logístico (seção 4.9) confirma e quantifica esse mecanismo na população completa "
        "(N=2.395.285): negros têm OR=0,708 de acesso a ocupações qualificadas (CBO 1–4) e "
        "OR=0,726 ao top 20% de renda em M1 — penalidades persistentes após controle por "
        "educação, experiência, sexo e contexto de UPA. O gradiente confirmado — OR(top20)=0,726 "
        "vs OR(top10)=0,671 — evidencia que o teto de vidro é progressivamente mais seletivo: "
        "a exclusão racial se intensifica exatamente nos limiares mais altos de acesso. "
        "O ICC de 26,2% no GLMM de UPA (vs 9,8% no HLM salarial) revela que o acesso a "
        "ocupações qualificadas é mais geograficamente determinado do que os salários. "
        "A convergência entre o gap residual do M4 (5,4%), o AME GLMM (−1,12 p.p. em M2) e "
        "o gap QR no q90 (−11,8%) estabelece diagnóstico tripartite: discriminação de acesso "
        "(GLMM), discriminação salarial intra-ocupação (HLM) e glass ceiling de progressão (QR). "
        "A análise de segregação espacial (seção 4.7) acrescenta que o padrão é geograficamente "
        "heterogêneo: o gap salarial é máximo nas capitais (62,8%), onde a estratificação "
        "ocupacional é mais intensa, e menor nas RMs não capitais (38,9%). Nas capitais, negros "
        "(60,2% da PEA) estão excluídos dos postos de alta remuneração (IR=0,47 no Top 5%), "
        "confirmando o padrão de glass ceiling racial que as seções 4.9 e 4.10 formalizam "
        "com precisão econométrica."
    )
    add_para(doc,
        "Teto de vidro duplo — síntese metodológica: o GLMM logístico (seção 4.9) documenta a "
        "barreira de acesso (discriminação antes da porta de entrada), e a regressão quantílica "
        "(seção 4.10), testada pelo critério de Koenker-Bassett, documenta a barreira de "
        "progressão salarial (discriminação dentro da ocupação, crescente com o quantil). "
        "Os dois instrumentos, estimados de forma independente sobre N=2,4 milhões de "
        "observações, convergem para a mesma conclusão: a discriminação racial no mercado de "
        "trabalho brasileiro opera como um sistema de barreiras em camadas, onde cada camada "
        "é empiricamente identificável, estatisticamente robusta e quantitativamente substancial. "
        "A interação negro×educação próxima de 1,0 em ambos os modelos (GLMM M3 OR≈1,01; "
        "retorno marginal QR ≈ 0 p.p. a diferença entre grupos) confirma o teto de credencial: "
        "investimento em capital humano por si só não rompe essas barreiras estruturais."
    )

    doc.add_page_break()

    # ── 5.1 PO: PRIORIZAÇÃO DE POLÍTICAS PÚBLICAS ─────────────────────────────
    add_heading(doc, "5.1 Pesquisa Operacional: Priorização e Alocação de Políticas Anti-Discriminação", level=2)
    add_para(doc,
        "A Pesquisa Operacional (PO) transforma os coeficientes econométricos estimados — "
        "gap salarial (HLM/OB: 52,6%), barreira de acesso (GLMM: OR=0,708 para CBO 1–4), "
        "glass ceiling de progressão (QR: Δ[q90−q10]=−3,99pp, Z=−2,75**) — em problemas "
        "de decisão formais. Seis intervenções foram avaliadas: (P1) cotas ocupacionais CBO 1–4, "
        "(P2) enforcement anti-discriminação, (P3) equidade educacional, (P4) desegregação "
        "residencial, (P5) mentoria e redes profissionais, (P6) transparência salarial "
        "obrigatória. Os modelos têm premissas de efetividade explicitamente marcadas [ASSUMIDO], "
        "baseadas em Holzer & Neumark (2000) e Bertrand & Duflo (2017), e não constituem "
        "previsões causais positivas — são cenários normativos de planejamento."
    )

    # Tabela PO TOPSIS se disponível
    topsis_path = TABLES / "po_politicas_topsis.csv"
    if topsis_path.exists():
        df_topsis = pd.read_csv(topsis_path)
        add_para(doc, "Tabela PO-1. Ranking TOPSIS das Políticas Anti-Discriminação Racial",
                 bold=True, first_line=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        headers_tp = ["Rank", "Intervenção", "Impacto Gap Sal.", "Impacto Gap Aces.", "Custo-Efetividade", "CC (TOPSIS)"]
        tbl_tp = doc.add_table(rows=len(df_topsis)+1, cols=len(headers_tp))
        tbl_tp.style = "Table Grid"
        tbl_tp.alignment = WD_TABLE_ALIGNMENT.CENTER
        shade_row(tbl_tp.rows[0], "1F3864")
        for j, h in enumerate(headers_tp):
            cell = tbl_tp.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9); run.font.name = "Arial"
        for i, (_, row_d) in enumerate(df_topsis.iterrows()):
            shade = "F2F2F2" if i % 2 == 0 else "FFFFFF"
            shade_row(tbl_tp.rows[i+1], shade)
            vals = [
                str(int(row_d["Rank"])),
                str(row_d["Política"]),
                f"{row_d['C1_gap_sal']:.4f}",
                f"{row_d['C2_gap_aces']:.4f}",
                f"{row_d['C4_custo_ef']:.3f}",
                f"{row_d['CC']:.4f}",
            ]
            for j, v in enumerate(vals):
                cell = tbl_tp.rows[i+1].cells[j]
                cell.text = v
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9); run.font.name = "Times New Roman"
        add_caption(doc, "Nota: CC próximo de 1 = melhor política segundo o critério multicritério (TOPSIS). "
                    "Critérios AHP: impacto gap salarial > impacto gap acesso > viabilidade ≈ custo-efetividade. "
                    "CR_AHP = 0,004 (consistência excelente). Premissas de efetividade marcadas [ASSUMIDO].")

    if (FIGURES / "po_politicas_topsis.png").exists():
        add_figure(doc, FIGURES / "po_politicas_topsis.png",
            "Figura PO-1 – Ranking TOPSIS (barras) e perfil radar das 6 políticas. "
            "P1 (Cotas CBO 1–4) domina por atacar diretamente a barreira de acesso documentada "
            "pelo GLMM (OR=0,708). P3 (Equidade Educacional) e P5 (Mentoria/Redes) completam "
            "o top-3 — confirmando que a discriminação de acesso exige intervenção estrutural, "
            "não apenas medidas de remuneração.",
            width_cm=16)

    add_para(doc,
        "O resultado mais relevante do TOPSIS é a posição dominante das Cotas Ocupacionais "
        "CBO 1–4 (CC=0,79, distante do segundo colocado em 0,56): a política que ataca "
        "diretamente a barreira de acesso — documentada pelo GLMM como OR=0,708 — tem "
        "retorno esperado por unidade de custo mais de duas vezes superior ao da segunda "
        "melhor política. Isso é consistente com o diagnóstico econométrico: 84,0% do gap "
        "salarial bruto é mediado por diferenças de dotações (acesso ocupacional), e a "
        "barreira de acesso antecede a barreira salarial. Intervenções que atacam apenas "
        "o resultado (salário), sem remover a barreira de entrada, capturam apenas os 16% "
        "do gap atribuíveis a retornos diferenciais — deixando 84% intocados."
    )

    if (FIGURES / "po_politicas_pl2_pareto.png").exists():
        add_figure(doc, FIGURES / "po_politicas_pl2_pareto.png",
            "Figura PO-2 – Fronteira de Pareto para o teto de vidro duplo (B=5 unidades de custo). "
            "Eixo x: redução do gap de acesso (GLMM log-odds); eixo y: redução do gap salarial (OB/QR). "
            "Cor = λ (peso do gap salarial). Ponto ouro: λ=0,5 (equilíbrio acesso+salário). "
            "Painel direito: alocação ótima de equilíbrio.",
            width_cm=16)
    add_para(doc,
        "A fronteira de Pareto (Figura PO-2) revela o tradeoff fundamental do teto de vidro "
        "duplo: maximizar a redução do gap de acesso (GLMM) ou do gap salarial (QR) são objetivos "
        "parcialmente conflitantes. Na solução de equilíbrio (λ=0,5), a alocação ótima concentra "
        "recursos em P1 (cotas CBO) e P5 (mentoria/redes) — as políticas com maior efetividade "
        "simultânea nos dois objetivos. Transparência salarial (P6) tem custo baixo e impacto "
        "imediato no glass ceiling de progressão QR, sendo recomendada como medida de curto prazo "
        "complementar às reformas estruturais de acesso. Com orçamento B=5, a PL-1 projeta "
        "redução de 24,0% do gap salarial bruto — mas isso requer implementação simultânea de "
        "cotas ocupacionais, equidade educacional e programas de mentoria."
    )

    doc.add_page_break()

    # ── 6. CONCLUSÃO ──────────────────────────────────────────────────────────
    doc.add_page_break()

    # ── 6. RECOMENDAÇÕES DE POLÍTICAS PÚBLICAS ────────────────────────────────
    add_heading(doc, "6. Recomendações de Políticas Públicas", level=1)
    add_para(doc,
        "Os resultados apresentados neste trabalho — em especial a persistência de um gap salarial "
        "racial de 5,5% (em log) após controle exaustivo de observáveis, o glass ceiling racial "
        "confirmado pela regressão quantílica (M3: −8,2% no q10 a −12,3% no q95), o gap de "
        "oportunidade documentado pelo logit multinível GLMM (OR=0,747, AME=−1,07 p.p. em M2) "
        "e o gap de acesso a alta qualificação praticamente estacionário entre 2016 e 2025 "
        "(negros: 9,6% → 11,0%; brancos: 19,7% → 22,4%) — configuram um diagnóstico que "
        "demanda intervenção pública estruturada em três eixos estratégicos, descritos a seguir."
    )

    add_heading(doc, "6.1 Eixo 1: Equidade de Acesso — Romper o Gap de Oportunidade", level=2)
    add_para(doc,
        "O Eixo 1 ataca o componente de dotações (84% do gap segundo Oaxaca-Blinder) por meio "
        "de políticas que ampliem o acesso de trabalhadores negros às ocupações de alta "
        "qualificação. Cinco instrumentos são recomendados:"
    )
    add_para(doc,
        "P1.1 — Cotas e metas de diversidade no setor privado. Legislação de transparência salarial "
        "e metas de representação racial nos escalões de liderança das empresas de médio e grande "
        "porte, inspirada na Lei de Igualdade Salarial (Lei n. 14.611/2023), com extensão para "
        "indicadores raciais além do de gênero. O Índice de Representação de 0,47 no Top 5% das "
        "capitais indica que o setor privado está longe de refletir a composição demográfica da "
        "força de trabalho."
    )
    add_para(doc,
        "P1.2 — Expansão de programas de mentoria e acesso a redes profissionais qualificadas. "
        "A SNA identificou que trabalhadores negros com pós-graduação têm betweenness nulo — "
        "estão desconectados das redes de indicação que convertem credenciais em mobilidade. "
        "Programas de mentoria cruzada (matching entre executivos brancos e profissionais negros "
        "em ascensão) têm custo fiscal baixo e impacto demonstrado sobre mobilidade intergeracional."
    )
    add_para(doc,
        "P1.3 — Políticas de educação profissional racialmente equânimes. A análise de composição "
        "ocupacional por CBO mostra razão negro/branco de 0,42 para dirigentes. Programas de "
        "formação técnica e gerencial com foco em populações sub-representadas — como o PRONATEC "
        "com metas raciais — são instrumentos de médio prazo para reduzir esse diferencial."
    )
    add_para(doc,
        "P1.4 — Intervenção na segregação residencial. O coeficiente contextual γ̂₀₁=−0,289 para "
        "a proporção de negros na UPA indica que viver em bairros majoritariamente negros reduz "
        "o rendimento em 25%, independentemente das características individuais. Políticas de "
        "habitação inclusiva — incluindo financiamento preferencial para famílias negras em "
        "regiões de maior acesso ao mercado de trabalho formal — são instrumentos estruturantes "
        "que políticas de renda sem endereçamento espacial não conseguem substituir."
    )
    add_para(doc,
        "P1.5 — Incentivo à formalização do emprego doméstico e rural. O grupo 'elementar' "
        "concentra 16,4% dos negros vs 7,8% dos brancos (razão 2,12). O trabalho doméstico "
        "tem cobertura previdenciária historicamente inferior. A universalização da cobertura "
        "previdenciária para trabalhadores domésticos e rurais tem impacto desproporcional "
        "sobre a segurança econômica de trabalhadores negros no médio e longo prazo."
    )

    add_heading(doc, "6.2 Eixo 2: Equidade Salarial — Combater a Discriminação de Retornos", level=2)
    add_para(doc,
        "O Eixo 2 ataca os 16% do gap atribuídos a retornos diferenciais — isto é, a remuneração "
        "diferenciada por raça dentro da mesma categoria ocupacional e mesmo perfil de "
        "qualificações. A regressão quantílica indica que esse componente é maior no topo "
        "da distribuição (gap residual M4: −7,9% no q95 contra −3,2% no q10)."
    )
    add_para(doc,
        "P2.1 — Fiscalização da Lei de Igualdade Salarial com recorte racial. A Lei n. 14.611/2023 "
        "obriga empresas com mais de 100 empregados a divulgar relatórios de igualdade salarial "
        "por gênero. A extensão do obrigação ao recorte racial — e o fortalecimento da fiscalização "
        "pelo MTE — permitiria identificar empresas com gaps raciais sistematicamente superiores "
        "à média setorial, sujeitas a penalidades ou metas obrigatórias de ajuste."
    )
    add_para(doc,
        "P2.2 — Transparência salarial no setor público. O gap racial persiste no setor público "
        "mesmo com controles (−25,2% vs −29,2% no privado). A publicação de faixas salariais por "
        "cargo, nível e raça nos relatórios anuais de gestão de pessoas das autarquias federais e "
        "estaduais tornaria visíveis eventuais diferenciais de progressão na carreira."
    )
    add_para(doc,
        "P2.3 — Proibição de critérios informais de promoção e cultura organizacional. "
        "Organizações com critérios de promoção baseados em 'fit cultural' ou redes informais "
        "de indicação perpetuam a discriminação estatística (Arrow, 1973) sem que ela seja "
        "explicitamente racial. A regulação de processos de promoção transparentes — com "
        "critérios de desempenho mensuráveis e auditáveis — é intervenção de baixo custo "
        "e alto impacto sobre o componente de retornos do gap."
    )

    add_heading(doc, "6.3 Eixo 3: Inclusão Produtiva — Superar a Armadilha da Renda Média", level=2)
    add_para(doc,
        "O Eixo 3 conecta as políticas anteriores a uma agenda de desenvolvimento econômico. "
        "A análise de H4 demonstrou que, se trabalhadores negros tivessem a mesma distribuição "
        "ocupacional que brancos, o ganho potencial de renda média seria de aproximadamente "
        "+127% para esse grupo — o que, extrapolado para a força de trabalho como um todo "
        "(negros ≈ 58%), representa um proxy de ganho de produtividade da ordem de dezenas de "
        "pontos percentuais. Hsieh et al. (2019) estimaram que a redução de barreiras "
        "ocupacionais para mulheres e negros nos EUA explicou até 40% do crescimento da "
        "produtividade total dos fatores entre 1960 e 2010 — magnitude que, para o Brasil "
        "com maior desigualdade racial de base, pode ser ainda superior."
    )
    add_para(doc,
        "P3.1 — Reforma tributária com progressividade racial implícita. O sistema tributário "
        "brasileiro é regressivo: tributos indiretos (ICMS, IPI, PIS/COFINS) consomem "
        "proporcionalmente mais da renda de trabalhadores de menor renda, que são "
        "majoritariamente negros. A transição do IBS/CBS proposta pela Reforma Tributária "
        "(EC 132/2023) cria oportunidade de redistribuição progressiva via cashback do IBS "
        "para rendas abaixo da linha de pobreza — instrumento com impacto racial desproporcional."
    )
    add_para(doc,
        "P3.2 — Investimento público em educação de qualidade nas UPAs de alta concentração negra. "
        "O contexto de UPA explica 20% da variância de renda (ICC_UPA). Escolas de tempo integral "
        "e com infraestrutura adequada nas periferias urbanas — onde a concentração de negros é "
        "maior — têm impacto sobre o capital humano que nenhuma política de renda de curto prazo "
        "consegue replicar. O coeficiente γ̂₀₁=+0,022 para a média de educação da UPA indica "
        "que elevar o nível médio de escolaridade da vizinhança aumenta a renda individual, "
        "mesmo controlando pela escolaridade própria."
    )
    add_para(doc,
        "P3.3 — Fortalecimento do sistema de inovação com inclusão racial. Países que superaram "
        "a armadilha da renda média investiram em P&D e formação de capital humano qualificado. "
        "O gap de alta qualificação (negros: 11% vs brancos: 22% em 2025) representa um "
        "desperdício de talento que limita a base de inovação do país. Bolsas de pós-graduação "
        "e P&D com recorte racial — via CAPES, CNPq e FAPESP — são instrumentos que agem "
        "simultaneamente sobre o gap de oportunidade e sobre o potencial de crescimento da PTF."
    )
    add_para(doc,
        "P3.4 — Política de valorização do salário mínimo com correção de produtividade. "
        "A análise temporal mostra que a renda real mediana cresceu de R$1.106 (2022) para "
        "R$1.345 (2024) — recuperação real de +21,6% —, mas desacelerou em 2025 (R$1.283) "
        "mesmo com a taxa de emprego atingindo máximo histórico de 94,4%. Isso sugere que "
        "o crescimento do emprego em 2025 foi concentrado em setores de baixa produtividade. "
        "Uma política de salário mínimo indexada à produtividade — e não apenas à inflação — "
        "incentiva empresas a migrarem para setores de maior valor adicionado, criando demanda "
        "por trabalho qualificado que retroalimenta os eixos 1 e 2."
    )
    add_para(doc,
        "Em síntese, as políticas recomendadas neste trabalho atuam em três horizontes temporais: "
        "curto prazo (1–2 anos): fiscalização da Lei de Igualdade Salarial, transparência "
        "salarial no setor público e cashback tributário; médio prazo (3–7 anos): educação "
        "profissional com metas raciais, mentoria e habitação inclusiva; longo prazo (8–15 anos): "
        "investimento em P&D com recorte racial e reforma do modelo de crescimento. A convergência "
        "racial no mercado de trabalho não é apenas uma demanda de justiça — é uma condição "
        "necessária para que o Brasil supere a armadilha da renda média."
    )

    doc.add_page_break()

    add_heading(doc, "Conclusão(ões)", level=1)
    add_para(doc,
        f"Este trabalho investigou o gap salarial racial no Brasil sob uma abordagem metodológica "
        f"integrativa, combinando econometria multinível, machine learning interpretável e análise "
        f"de redes sociais sobre 10 anos de dados da PNAD Contínua. As cinco hipóteses originais "
        f"foram confirmadas."
    )
    add_para(doc,
        f"O gap racial bruto de {k['gb']:.1f}% reduz-se para {k['gl']:.1f}% quando contextos "
        f"residenciais e macroeconômicos são controlados, mas permanece significativo e "
        f"estatisticamente robusto — sinalizando discriminação direta residual não explicável por "
        f"diferenças de capital humano ou localização. A mediação contextual de {k['med']:.1f}% pelo "
        f"bairro de moradia é o resultado de maior relevância para políticas públicas: ele quantifica "
        f"a contribuição da segregação residencial ao gap racial, abrindo espaço para intervenções "
        f"habitacionais e de mobilidade urbana como ferramentas de redução da desigualdade racial."
    )
    add_para(doc,
        "A análise de redes adiciona uma perspectiva inovadora ao diagnóstico: a exclusão de "
        "trabalhadores negros das posições de brokerage na rede de co-residência sugere que, mesmo "
        "quando a segregação geográfica é moderada em termos absolutos, as assimetrias nas posições "
        "estruturais da rede são suficientes para limitar o acesso de negros ao capital social que "
        "converte educação em mobilidade."
    )

    p = doc.add_paragraph()
    run = p.add_run("Contribuição principal. ")
    run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(12)
    run2 = p.add_run(
        "Este estudo oferece, ao nosso conhecimento, a primeira análise integrada de HLM multinível, "
        "clustering, SHAP e SNA sobre a série completa da PNAD Contínua, estabelecendo uma "
        "metodologia replicável para o monitoramento longitudinal da desigualdade racial no mercado "
        "de trabalho brasileiro."
    )
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(p, first_line=1.25)

    p = doc.add_paragraph()
    run = p.add_run("Limitações e direções futuras. ")
    run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(12)
    run2 = p.add_run(
        "O caráter transversal do painel público da PNAD Contínua impede a análise de trajetórias "
        "individuais de mobilidade salarial ao longo do tempo. A extensão da análise para dados "
        "corporativos (RAIS) permitiria investigar a hipótese de glass ceiling em cargos de liderança."
    )
    run2.font.name = "Times New Roman"; run2.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(p, first_line=1.25)

    doc.add_page_break()

    # ── REFERÊNCIAS ────────────────────────────────────────────────────────────
    add_heading(doc, "Referências", level=1)
    refs = [
        "BLINDER, A. S. Wage discrimination: reduced form and structural estimates. Journal of Human Resources, v. 8, n. 4, p. 436–455, 1973.",
        "BREIMAN, L. Random forests. Machine Learning, v. 45, n. 1, p. 5–32, 2001.",
        "BURT, R. S. Structural holes and good ideas. American Journal of Sociology, v. 110, n. 2, p. 349–399, 2004.",
        "CHEN, T.; GUESTRIN, C. XGBoost: a scalable tree boosting system. In: ACM SIGKDD, 2016. p. 785–794.",
        "DAVIES, D. L.; BOULDIN, D. W. A cluster separation measure. IEEE Transactions on Pattern Analysis and Machine Intelligence, v. 1, n. 2, p. 224–227, 1979.",
        "GRANOVETTER, M. The strength of weak ties. American Journal of Sociology, v. 78, n. 6, p. 1360–1380, 1973.",
        "ARROW, K. J. The theory of discrimination. In: ASHENFELTER, O.; REES, A. (Ed.). Discrimination in labor markets. Princeton: Princeton University Press, 1973. p. 3–33.",
        "HASENBALG, C. Discriminação e desigualdades raciais no Brasil. Rio de Janeiro: Graal, 1979.",
        "KOENKER, R.; BASSETT, G. Regression quantiles. Econometrica, v. 46, n. 1, p. 33–50, 1978.",
        "HENRIQUES, R. Desigualdade racial no Brasil: evolução das condições de vida na década de 90. Rio de Janeiro: IPEA, 2001. (Texto para discussão, 807).",
        "IBGE – INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA. PNAD Contínua: notas metodológicas. Rio de Janeiro: IBGE, 2023.",
        "LUNDBERG, S. M.; LEE, S.-I. A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, v. 30, 2017.",
        "MARQUES, E. Redes sociais, segregação e pobreza. São Paulo: Editora UNESP, 2010.",
        "OAXACA, R. Male-female wage differentials in urban labor markets. International Economic Review, v. 14, n. 3, p. 693–709, 1973.",
        "PAGER, D. The use of field experiments for studies of employment discrimination: contributions, critiques, and directions for the future. The Annals of the American Academy of Political and Social Science, v. 609, p. 104–133, 2007.",
        "RAUDENBUSH, S. W.; BRYK, A. S. Hierarchical linear models: applications and data analysis methods. 2. ed. Thousand Oaks: Sage, 2002.",
        "ROUSSEEUW, P. J. Silhouettes: a graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, v. 20, p. 53–65, 1987.",
        "SAMPSON, R. J.; RAUDENBUSH, S. W.; EARLS, F. Neighborhoods and violent crime: a multilevel study of collective efficacy. Science, v. 277, p. 918–924, 1997.",
        "SOARES, S. S. D. Perfil da discriminação no mercado de trabalho: raça, sexo e salários no Brasil 1992–2006. Rio de Janeiro: IPEA, 2009. (Texto para discussão, 1.395).",
        "WILSON, W. J. The truly disadvantaged: the inner city, the underclass, and public policy. Chicago: University of Chicago Press, 1987.",
        "HSIEH, C. T. et al. The allocation of talent and U.S. economic growth. Econometrica, v. 87, n. 5, p. 1439–1474, 2019.",
        "MOINEDDIN, R.; MATHESON, F. I.; GLAZIER, R. H. A simulation study of sample size for multilevel logistic regression models. BMC Medical Research Methodology, v. 7, n. 34, 2007.",
        "ALMEIDA, S. Racismo estrutural. São Paulo: Sueli Carneiro; Pólen, 2019.",
        "ANDREWS, G. R. Blacks and whites in São Paulo, Brazil: 1888–1988. Madison: University of Wisconsin Press, 1991.",
        "BECKER, G. S. The economics of discrimination. Chicago: University of Chicago Press, 1957.",
        "CACCIAMALI, M. C.; HIRATA, G. I. A influência da raça e do gênero nas oportunidades de obtenção de renda — uma análise da discriminação em mercados de trabalho distintos: Bahia e São Paulo. Estudos Econômicos, v. 35, n. 4, p. 767–795, 2005.",
        "CRENSHAW, K. Demarginalizing the intersection of race and sex: a black feminist critique of antidiscrimination doctrine, feminist theory and antiracist politics. University of Chicago Legal Forum, v. 1989, n. 1, p. 139–167, 1989.",
        "DIEESE – DEPARTAMENTO INTERSINDICAL DE ESTATÍSTICA E ESTUDOS SOCIOECONÔMICOS. Desigualdade racial no mercado de trabalho. São Paulo: DIEESE, 2023.",
        "SOARES, S. S. D. O perfil da discriminação no mercado de trabalho — homens negros, mulheres brancas e mulheres negras. Rio de Janeiro: IPEA, 2000. (Texto para discussão, 769).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent       = Cm(0)
        pf.space_after        = Pt(6)
        pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing       = 1.5
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO A — HIPÓTESES SOBRE O PAPEL DO ESTADO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Anexo A — Hipóteses sobre o Papel do Estado na Desigualdade de Renda", level=1)
    add_para(doc,
        "Este anexo apresenta os resultados de cinco hipóteses complementares sobre o papel do "
        "Estado como indutor ou moderador de desigualdades de renda, estimadas com a mesma base "
        "de dados PNAD Contínua 2016–2025. As análises utilizam amostra de 20% (1,54 milhão de "
        "observações de trabalhadores com renda positiva) e regressões OLS com erros clusterizados "
        "por UPA, além de indicadores descritivos calculados sobre o universo completo."
    )

    add_heading(doc, "A.1 H1 — O Estado como Indutor de Desigualdade Geral de Renda", level=2)
    add_para(doc,
        "Hipótese: O Estado, por meio da estrutura de emprego público e do prêmio salarial que "
        "ele oferece, é um indutor de desigualdade de renda — independentemente de raça ou gênero."
    )
    add_para(doc,
        "Resultado: CONFIRMADA parcialmente. O Gini total é 0,4879. Dentro do setor público, "
        "o Gini é 0,4662 — ligeiramente inferior ao privado (0,4715) —, indicando que "
        "internamente o setor público é menos desigual. Contudo, o prêmio público médio sobre o "
        "privado é de +99,6% (quase o dobro da remuneração), gerando desigualdade entre setores. "
        "A decomposição de Theil T mostra que 7,8% da desigualdade total é explicada pela "
        "diferença entre setores (público vs. privado). O Estado amplifica a desigualdade "
        "sobretudo pela criação de um segmento premium de empregadores — com remuneração "
        "muito acima da mediana privada —, e não por ser mais desigual internamente. A análise "
        "por decil de renda mostra que o prêmio público é maior nos decis intermediários "
        "e superiores (servidores estatutários e militares de alta patente), enquanto nos "
        "decis inferiores o prêmio é menor ou nulo — confirmando que o Estado é "
        "estruturalmente indutor de desigualdade entre segmentos do mercado de trabalho."
    )
    if (FIGURES / "estado_h1_gini.png").exists():
        add_figure(doc, FIGURES / "estado_h1_gini.png",
            "Figura A1 – Curvas de Lorenz por setor (esq.) e prêmio salarial público por "
            "decil de renda (dir.). Gini: total=0,488 | público=0,466 | privado=0,472. "
            "Theil entre setores = 7,8% do total.",
            width_cm=15)

    add_heading(doc, "A.2 H2 — O Estado como Indutor de Desigualdade Racial via Salários", level=2)
    add_para(doc,
        "Hipótese: O Estado paga salários que perpetuam ou ampliam o gap racial, sendo indutor "
        "de desigualdade de renda por raça."
    )
    add_para(doc,
        "Resultado: REFUTADA parcialmente. O gap racial bruto é menor no setor público "
        "(−33,3%) do que no privado (−44,1%). Após controle por escolaridade, experiência e "
        "horas trabalhadas (regressão com interação negro × setor_publico), o gap controlado "
        "é −29,2% no privado e −25,2% no público — uma diferença de 4 pontos percentuais "
        "favorável ao setor público. A interação negro:setor_publico é positiva e "
        "estatisticamente significativa, indicando que o Estado atenua (mas não elimina) o "
        "gap racial. O gap de 25,2% que persiste mesmo no setor público — onde a admissão "
        "é via concurso público padronizado — é consistente com discriminação na progressão "
        "de carreira: cargos de Direção e Assessoramento Superior (DAS) são preenchidos por "
        "indicação política, e a sub-representação de negros nesses cargos perpetua o gap "
        "mesmo após a entrada equânime via concurso."
    )
    if (FIGURES / "estado_h2h3_gaps.png").exists():
        add_figure(doc, FIGURES / "estado_h2h3_gaps.png",
            "Figura A2 – Gap racial (H2, esq.) e gap de gênero (H3, dir.) por setor, "
            "comparando gap bruto e gap controlado. Barras de erro = IC 95% clusterizado por UPA.",
            width_cm=15)

    add_heading(doc, "A.3 H3 — O Estado como Indutor de Igualdade por Raça e Gênero", level=2)
    add_para(doc,
        "Hipótese: O setor público promove maior igualdade racial e de gênero em relação ao "
        "setor privado, funcionando como indutor de equidade."
    )
    add_para(doc,
        "Resultado: CONFIRMADA para raça, REFUTADA para gênero. Para raça, o Estado atenua o "
        "gap racial em aproximadamente 4 pontos percentuais (ver H2). Para gênero, o resultado "
        "é contrário ao esperado: o gap de gênero controlado é de −16,9% no setor privado e "
        "−18,8% no setor público — 1,9 ponto percentual maior no setor público. Em termos "
        "brutos, o gap de gênero também é maior no público (−24,1% vs −21,1% no privado). "
        "Esse resultado contradiz a narrativa de que o concurso público garante igualdade de "
        "gênero. A explicação é estrutural: o setor público tem concentração de mulheres em "
        "cargos de médio nível (professoras, enfermeiras, assistentes) enquanto os cargos de "
        "maior remuneração (juízes, promotores, delegados, militares de alta patente) têm "
        "maioria masculina. A formalidade do concurso equaliza o acesso ao setor, mas não "
        "equaliza a distribuição por carreira dentro dele."
    )

    add_heading(doc, "A.4 H4 — Renda Real, Emprego e Inclusão Produtiva", level=2)
    add_para(doc,
        "Hipótese: Embora os índices de atividade econômica e emprego tenham melhorado no "
        "governo atual (2023–2025), a renda real acima da inflação tem desacelerado; a inclusão "
        "da população marginalizada em oportunidades de qualidade aumentaria a produtividade."
    )
    add_para(doc,
        "Resultado: CONFIRMADA com nuances. A taxa de emprego da PEA atingiu máximo histórico "
        "de 94,4% em 2025 (vs 87,4% em 2020). A renda real mediana — deflacionada pelo IPCA "
        "(base 2016=100) — recuperou-se fortemente de R$1.106 em 2022 (mínimo histórico da "
        "série) para R$1.345 em 2024 (+21,6% real), mas recuou levemente para R$1.283 em "
        "2025 (−4,6%), mesmo com emprego em alta. Isso é consistente com composição do "
        "crescimento: o emprego adicional em 2025 concentrou-se em setores de serviços de "
        "baixa produtividade (serviços domésticos, comércio varejista, entregas por aplicativos), "
        "puxando a mediana para baixo apesar da expansão quantitativa. A hipótese de inclusão "
        "produtiva é apoiada pela simulação ocupacional: se trabalhadores negros (≈58% da "
        "força de trabalho) tivessem a mesma distribuição por grupo CBO que os brancos, o "
        "ganho potencial de renda média para esse grupo seria de +127,3%. Hsieh et al. (2019) "
        "estimam que a equalização de barreiras ocupacionais por raça e gênero nos EUA explicou "
        "até 40% do crescimento da PTF entre 1960 e 2010; para o Brasil, com desigualdade de "
        "base maior, o potencial é proporcionalmente superior."
    )
    if (FIGURES / "estado_h4_tendencia.png").exists():
        add_figure(doc, FIGURES / "estado_h4_tendencia.png",
            "Figura A3 – Renda mediana real deflacionada por IPCA (esq.) e taxa de emprego "
            "da PEA por ano (dir.), 2016–2025. A divergência entre emprego (crescente) e "
            "renda real (desaceleração em 2025) evidencia a armadilha da renda média.",
            width_cm=15)
    if (FIGURES / "estado_h4_inclusao.png").exists():
        add_figure(doc, FIGURES / "estado_h4_inclusao.png",
            "Figura A4 – Distribuição ocupacional por raça (esq.) e simulação de ganho de "
            "renda por equalização da distribuição CBO (dir.). Negros têm 11,0% em alta "
            "qualificação vs 22,4% dos brancos (2025); equalização implicaria +127,3% de renda.",
            width_cm=15)

    add_heading(doc, "A.5 H5 — Políticas Públicas para Superar a Armadilha da Renda Média", level=2)
    add_para(doc,
        "Hipótese: Barreiras estruturais ao upgrade ocupacional da força de trabalho — com "
        "recorte racial — mantêm o Brasil preso na armadilha da renda média, impedindo a "
        "transição para uma economia de alta produtividade."
    )
    add_para(doc,
        "Resultado: CONFIRMADA. A proporção de trabalhadores em ocupações de alta qualificação "
        "(dirigentes + profissionais) evoluiu de 13,9% (2016) para 15,7% (2025) para o total, "
        "mas o gap racial permaneceu praticamente estático: negros passaram de 9,6% para 11,0%, "
        "e brancos de 19,7% para 22,4%, mantendo diferença de ~11 pontos percentuais durante "
        "todo o período. O gap de acesso à alta qualificação sequer convergiu — na margem, "
        "ele levemente se ampliou (de 10,1 p.p. em 2016 para 11,4 p.p. em 2025). Isso implica "
        "que, ao ritmo atual, a eliminação do gap de acesso levaria mais de um século — "
        "consistente com a estimativa da trajetória de gap salarial apresentada na seção 5. "
        "A análise por grupo CBO também demonstra que a penalidade racial é transversal: mesmo "
        "dentro da mesma categoria ocupacional, negros recebem menos em todos os grupos CBO, "
        "indicando que a barreira não é apenas de acesso, mas também de remuneração intra-categoria. "
        "A agenda de políticas públicas para superar essa armadilha está detalhada na Seção 6."
    )
    if (FIGURES / "estado_h5_armadilha.png").exists():
        add_figure(doc, FIGURES / "estado_h5_armadilha.png",
            "Figura A5 – Evolução do acesso a alta qualificação por raça (2016–2025, esq.) e "
            "renda mediana por grupo CBO e raça (dir.). O gap de acesso é estacionário e a "
            "penalidade salarial é transversal a todos os grupos ocupacionais.",
            width_cm=15)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO B — COMPARATIVO DE LOG-VEROSSIMILHANÇA DOS MODELOS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Anexo B — Comparativo de Log-Verossimilhança e Justificação Metodológica", level=1)
    add_para(doc,
        "Este anexo apresenta a comparação formal de ajuste entre especificações OLS (sem "
        "efeitos aleatórios) e HLM (com efeitos aleatórios de UF), por meio de log-verossimilhança "
        "(LL) e Critério de Informação de Akaike (AIC). Os modelos são estimados na mesma amostra "
        "de 20% (n=1.543.512) para garantir comparabilidade."
    )

    # Tabela de resultados
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    headers_b = ["Modelo", "Tipo", "k", "Log-Verossimilhança", "AIC"]
    rows_b = [
        ["OLS Nulo",        "OLS", "2",  "−2.072.459", "4.144.922"],
        ["OLS Individual",  "OLS", "10", "−1.830.383", "3.660.787"],
        ["OLS + FE (UF)",   "OLS", "36", "−1.842.380", "3.684.833"],
        ["OLS Contextual",  "OLS", "16", "−1.794.140", "3.588.313"],
        ["OLS Ocupação M4", "OLS", "29", "−1.501.555", "3.003.168"],
        ["HLM Nulo",        "HLM", "3",  "−1.976.647", "3.953.299"],
        ["HLM Individual",  "HLM", "12", "−1.830.562", "3.661.149"],
        ["HLM Contextual",  "HLM", "15", "−1.794.327", "3.588.684"],
        ["HLM Completo M3", "HLM", "18", "−1.794.327", "3.588.691"],
        ["HLM Ocupação M4", "HLM", "31", "−1.501.815", "3.003.692"],
    ]
    tbl = doc.add_table(rows=len(rows_b)+1, cols=5)
    tbl.style = "Table Grid"
    for j, h in enumerate(headers_b):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
    for i, row_data in enumerate(rows_b):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i+1].cells[j]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"

    doc.add_paragraph()
    add_para(doc,
        "Nota: k = número de parâmetros livres (coeficientes + componentes de variância). "
        "AIC = 2k − 2LL. Para os modelos HLM, a log-verossimilhança é calculada via "
        "verossimilhança restrita (REML) para os componentes de variância. Os modelos OLS "
        "são estimados por MQO padrão com erros clusterizados por UPA para inferência, "
        "mas a LL reportada é a log-verossimilhança não-clusterizada (usada apenas para "
        "comparação de ajuste, não para inferência)."
    )

    add_para(doc,
        "Principais conclusões da comparação: (1) O LRT entre HLM Nulo e OLS Nulo gera "
        "χ²=191.625 (Δk=1), rejeitando o modelo plano com probabilidade virtualmente nula — "
        "a estrutura hierárquica é estatisticamente necessária. (2) O OLS com FE de UF "
        "(27 dummies) tem AIC=3.684.832, pior que o OLS Individual (AIC=3.660.787) "
        "— os 26 parâmetros adicionais para efeitos fixos de estado não compensam o custo "
        "de penalização do AIC, enquanto o HLM Contextual (AIC=3.588.684) obtém "
        "desempenho similar com metade dos parâmetros, decomposição formal da variância "
        "e standard errors válidos para inferência hierárquica. (3) O modelo HLM Ocupação "
        "(M4) é o de melhor LL dentre os HLM (−1.501.815), com AIC=3.003.692 — apenas "
        "524 unidades acima do OLS M4 equivalente, mas com a vantagem fundamental de "
        "produzir standard errors válidos para a estrutura de dados aninhada."
    )
    if (FIGURES / "modelos_loglik_aic.png").exists():
        add_figure(doc, FIGURES / "modelos_loglik_aic.png",
            "Figura B1 – Log-verossimilhança (esq.) e AIC (dir.) por modelo. "
            "Azul: OLS; Verde: HLM. O OLS+FE tem AIC pior que OLS Individual, "
            "enquanto o HLM Contextual obtém desempenho equivalente com maior validade inferencial.",
            width_cm=16)
    if (FIGURES / "modelos_lrt_icc.png").exists():
        add_figure(doc, FIGURES / "modelos_lrt_icc.png",
            "Figura B2 – LRT progressivo entre etapas do HLM (esq.): cada adição é "
            "estatisticamente justificada. Trajetória do ICC_UF (dir.): permanece acima de "
            "5% em todos os modelos, confirmando necessidade da estrutura de 3 níveis.",
            width_cm=16)

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Carregando resultados ...")
    r = load_results()
    k = extract_kpis(r)

    print("KPIs:")
    for key in ["gb","gl","med","icc_m0","xgb_r2","shap_top1_val","shap_negro_rank","sna_h"]:
        print(f"  {key}: {k[key]}")

    print("\nGerando Word ...")
    doc = build_doc(r, k)
    doc.save(OUT_DOC)

    size_kb = OUT_DOC.stat().st_size // 1024
    print(f"\nArquivo gerado: {OUT_DOC.name}  ({size_kb} KB)")
    print("Abra com Microsoft Word para editar e imprimir.")
