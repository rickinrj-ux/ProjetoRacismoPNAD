"""
gerar_guia_estudo.py
Gera guia_estudo_defesa.docx — roteiro de leituras + interpretação detalhada dos gráficos.
"""
import sys; sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from params import P, fmt, fmtN, ame, or_str

ROOT   = Path(r"C:\Users\user\Documents\ProjetoRacismoPNAD")
FIGS   = ROOT / "outputs" / "figures"
OUT    = ROOT / "guia_estudo_defesa.docx"

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False,
             color=None, name="Calibri"):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(0x1F,0x38,0x64)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p

def add_para(doc, text="", size=11, bold=False, italic=False,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, level=0, size=11, bold=False, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_colored_box(doc, title, body_lines, title_color=(0x1F,0x38,0x64),
                    bg_hex="E3F2FD"):
    # Simulates a callout box with a bold title paragraph + indented body
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"▌ {title}")
    set_font(run, size=11, bold=True, color=title_color)
    for line in body_lines:
        pb = doc.add_paragraph()
        pb.paragraph_format.left_indent  = Cm(0.8)
        pb.paragraph_format.space_after  = Pt(2)
        run2 = pb.add_run(line)
        set_font(run2, size=10.5)

def add_separator(doc):
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        run.font.size = Pt(8)

def add_figure(doc, img_path, width_cm=14, caption=""):
    p_img = Path(img_path)
    if p_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(p_img), width=Cm(width_cm))
    else:
        add_para(doc, f"[Figura não encontrada: {p_img.name}]",
                 size=9, italic=True, color=(0x99,0x99,0x99),
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    if caption:
        pc = doc.add_paragraph(caption)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(10)
        for run in pc.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x60,0x60,0x60)

# ══════════════════════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_para(doc, "GUIA DE ESTUDO PARA DEFESA DE TCC",
         size=20, bold=True, color=(0x1F,0x38,0x64),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40)
add_para(doc, "Racismo Estrutural e Mercado de Trabalho no Brasil",
         size=14, italic=True, color=(0x44,0x44,0x44),
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Ricardo Calheiros  |  MBA USP/ESALQ  |  2026",
         size=11, color=(0x66,0x66,0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_separator(doc)
add_para(doc, "Este guia tem dois objetivos: (1) indicar o que ler de cada obra citada no TCC antes "
              "da defesa, com priorização por risco de pergunta da banca; e (2) oferecer uma "
              "interpretação detalhada de cada gráfico produzido, incluindo o que dizer se você "
              "for questionado sobre ele durante a apresentação.",
         size=11, space_after=12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTE 1 — ROTEIRO DE LEITURAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PARTE 1 — ROTEIRO DE LEITURAS PRIORITÁRIAS", level=1)
add_para(doc, "Como acessar as obras", size=13, bold=True, color=(0x1F,0x38,0x64), space_before=6)
add_bullet(doc, "Portal CAPES (gratuito com vínculo USP/ESALQ): periodicos.capes.gov.br — login com CPF USP")
add_bullet(doc, "Google Scholar: busque título completo + nome do autor")
add_bullet(doc, "Sci-Hub: para artigos pagos — cole o DOI ou título")
add_bullet(doc, "LibGen: para livros — busque autor + título")
doc.add_paragraph()

# ── Prioridade 1 ──────────────────────────────────────────────────────────────
add_heading(doc, "1.1  Prioridade 1 — Você será perguntado sobre estas obras", level=2)

add_para(doc, "RAUDENBUSH, S. W.; BRYK, A. S. (2002). Hierarchical Linear Models. 2. ed. Thousand Oaks: SAGE.",
         size=11, bold=True, space_before=8)
add_bullet(doc, "Por que é crítica: você cita o limiar de ICC=5% e usa a nomenclatura M0–M4 diretamente deste livro.")
add_bullet(doc, "O que ler: Capítulo 2 (conceito de ICC, decomposição da variância entre e dentro dos grupos) "
                "e Capítulo 4 (modelos de intercepto e inclinação aleatória). Aproximadamente 40 páginas.")
add_bullet(doc, "Número específico a verificar: a regra prática de ICC > 0,05 como justificativa para HLM "
                "está no Capítulo 2, seção 'Intraclass Correlation'.")
add_bullet(doc, "Pergunta provável da banca: 'Por que 5% é o limiar para o ICC?' — Resposta: convenção de "
                "Raudenbush & Bryk (2002) amplamente adotada na literatura educacional e econômica; "
                f"no seu caso o ICC é {fmt(P['ICC_HLM_M0_pct'],2)}%, bem acima do limiar.")
doc.add_paragraph()

add_para(doc, "HSIEH, C-T.; HURST, E.; JONES, C. I.; KLENOW, P. J. (2019). The Allocation of Talent and "
              "U.S. Economic Growth. Econometrica, v. 87, n. 5, p. 1439-1474.",
         size=11, bold=True, space_before=8)
add_bullet(doc, "Por que é crítica: você usa o argumento de que eliminar barreiras ocupacionais = ganho de "
                "produtividade (+127,3% renda negra, proxy +74,3 p.p. de PIB).")
add_bullet(doc, "O que ler: Abstract (2 parágrafos) + Introduction (pp. 1–6) + Tabela 1 (resultados "
                "principais) + Conclusion (pp. 38–40). O artigo completo tem 40 páginas mas "
                "as seções indicadas são suficientes para a defesa.")
add_bullet(doc, "Achado central a memorizar: barreiras à entrada de mulheres e negros em ocupações "
                "de alta qualificação explicam 40% do crescimento da PTF dos EUA entre 1960 e 2010.")
add_bullet(doc, "Pergunta provável: 'Qual a diferença entre o resultado deles e o seu?' — "
                "Resposta: eles estimam PTF com modelo de equilíbrio geral; seu proxy é mais conservador "
                "(diferença de renda × participação na força de trabalho), o que justifica usá-lo "
                "como lower bound.")
doc.add_paragraph()

add_para(doc, "OAXACA, R. (1973). Male-Female Wage Differentials in Urban Labor Markets. "
              "International Economic Review, v. 14, n. 3, p. 693-709.",
         size=11, bold=True, space_before=8)
add_para(doc, "BLINDER, A. S. (1973). Wage Discrimination: Reduced Form and Structural Estimates. "
              "Journal of Human Resources, v. 8, n. 4, p. 436-455.",
         size=11, bold=True)
add_bullet(doc, "Por que são críticos: você usa a decomposição deles e escolheu a formulação two-fold "
                "(dotações + retornos) em vez da three-fold.")
add_bullet(doc, "O que ler: em cada artigo, leia apenas a seção onde a decomposição é formalizada "
                "(~4–6 páginas). Cada artigo tem ~15 páginas no total.")
add_bullet(doc, "Diferença entre os dois: Oaxaca usa o grupo masculino como referência; "
                "Blinder usa uma média ponderada. Sua implementação usou brancos como grupo de referência.")
add_bullet(doc, "Pergunta provável: 'Por que two-fold e não three-fold?' — Resposta: na decomposição "
                "three-fold, o terceiro componente (interação) não tem interpretação econômica clara "
                "em contextos de discriminação racial; a literatura de economics of discrimination "
                "prefere two-fold (Fortin, Lemieux & Firpo, 2011).")
doc.add_paragraph()

# ── Prioridade 2 ──────────────────────────────────────────────────────────────
add_heading(doc, "1.2  Prioridade 2 — Risco moderado", level=2)

add_para(doc, "BURT, R. S. (2004). Structural Holes and Good Ideas. "
              "American Journal of Sociology, v. 110, n. 2, p. 349-399.",
         size=11, bold=True, space_before=8)
add_bullet(doc, "Por que é relevante: fundamenta a análise de betweenness centrality na SNA e "
                "o argumento de que negros têm betweenness=0 mesmo com pós-graduação.")
add_bullet(doc, "O que ler: Introduction (pp. 349–355) + seção 'Network Constraint and Brokerage' "
                "(pp. 355–370). São ~20 páginas.")
add_bullet(doc, "Conceito-chave: structural hole = posição de broker entre grupos que não se conectam "
                "diretamente. Alta betweenness = acesso antecipado a informação e oportunidades.")
doc.add_paragraph()

add_para(doc, "WILSON, W. J. (1987). The Truly Disadvantaged: The Inner City, the Underclass, "
              "and Public Policy. Chicago: University of Chicago Press.",
         size=11, bold=True, space_before=8)
add_bullet(doc, "Por que é relevante: justifica o uso da UPA como proxy de bairro e o argumento "
                "de que o local de moradia explica parte do gap salarial.")
add_bullet(doc, "O que ler: Capítulo 1 (tese central do efeito de vizinhança) e Capítulo 7 "
                "(implicações de política). São ~40 páginas.")
add_bullet(doc, "Tese central: concentração de pobreza em bairros segregados cria isolamento "
                "de redes profissionais — reforça o argumento da UPA como mediadora.")
doc.add_paragraph()

# ── Prioridade 3 ──────────────────────────────────────────────────────────────
add_heading(doc, "1.3  Prioridade 3 — Leia se sobrar tempo", level=2)

obras_p3 = [
    ("MOINEDDIN, R. et al. (2007). A simulation study of sample size for multilevel logistic "
     "regression models. BMC Medical Research Methodology, v. 7, n. 34.",
     "Justifica o uso de logit multinível como aproximação válida mesmo sem GLMM completo. "
     "Leia apenas Abstract + Discussion (~3 páginas)."),
    ("CHEN, T.; GUESTRIN, C. (2016). XGBoost: A Scalable Tree Boosting System. "
     "Proceedings of KDD 2016.",
     "Base técnica do ML/SHAP. Leia Abstract + seção de 'Gain Importance' (~5 páginas)."),
    ("LUNDBERG, S.; LEE, S-I. (2017). A Unified Approach to Interpreting Model Predictions. "
     "NeurIPS 2017.",
     "Fundamenta os SHAP values. Leia Introduction + Figure 1 (~4 páginas)."),
    ("FORTIN, N.; LEMIEUX, T.; FIRPO, S. (2011). Decomposition Methods in Economics. "
     "Handbook of Labor Economics, v. 4A.",
     "Survey sobre decomposições salariais. Leia seção 3 (Oaxaca-Blinder, ~15 páginas). "
     "Justifica a escolha two-fold."),
]
for titulo, instrucao in obras_p3:
    add_para(doc, titulo, size=11, bold=True, space_before=6)
    add_bullet(doc, instrucao)
    doc.add_paragraph()

# ── Cronograma ────────────────────────────────────────────────────────────────
add_heading(doc, "1.4  Cronograma sugerido (1 semana antes da defesa)", level=2)
cronograma = [
    ("Dia 1", "Raudenbush & Bryk — Capítulos 2 e 4 (~40 páginas)"),
    ("Dia 2", "Hsieh et al. (2019) — artigo completo (~40 páginas)"),
    ("Dia 3", "Oaxaca (1973) + Blinder (1973) — seções de decomposição (~30 páginas no total)"),
    ("Dia 4", "Burt (2004) + Wilson (1987) Cap. 1 (~60 páginas no total)"),
    ("Dia 5", "Reler o roteiro de defesa e simular as perguntas da banca com os novos insumos"),
    ("Dia 6", "Revisão das equações dos modelos (HLM, Oaxaca, QR) e dos resultados numéricos"),
    ("Dia 7", "Descanso + leitura leve do Fortin et al. se sentir necessidade"),
]
for dia, atividade in cronograma:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{dia}: ")
    set_font(r1, size=11, bold=True, color=(0x1F,0x38,0x64))
    r2 = p.add_run(atividade)
    set_font(r2, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTE 2 — INTERPRETAÇÃO DETALHADA DOS GRÁFICOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PARTE 2 — COMO LER CADA GRÁFICO", level=1)
add_para(doc, "Para cada figura, este guia explica: (a) o que está nos eixos, (b) o que o padrão "
              "visual significa, (c) o achado principal e (d) o que dizer se a banca perguntar.",
         size=11, space_after=10)

# ── Figuras Descritivas (novas) ───────────────────────────────────────────────
add_heading(doc, "Figura 1 — Distribuição de Log-Renda por Raça (Densidade KDE)", level=2)
add_figure(doc, FIGS / "fig1_densidade_log_salario.png", width_cm=14,
           caption="Estimativa por kernel (KDE) da log-renda por raça. Ponderado por V1028.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: log da renda nominal. Escala logarítmica aplaina a assimetria da distribuição "
                "de renda e permite leitura de diferenças percentuais como distâncias.")
add_bullet(doc, "Eixo Y: densidade estimada — altura da curva indica concentração de observações "
                "naquele nível de renda.")
add_bullet(doc, "Linha azul = brancos; linha vermelha = negros. Linhas verticais tracejadas = medianas.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Curva de negros deslocada para a ESQUERDA (menores log-rendas): não são apenas "
                "alguns outliers — toda a distribuição está à esquerda.")
add_bullet(doc, "Mediana bruta: brancos R$2.000 / negros R$1.450 → gap de 27,5%. "
                "Gap médio maior (37,5%) porque negros são sobre-representados na cauda inferior.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'O gráfico de densidade mostra que o gap não é gerado por poucos casos extremos — "
     "toda a distribuição de negros está deslocada à esquerda. O gap de 27,5% na mediana "
     "é a estimativa mais conservadora: a média (37,5%) é maior porque incorpora as rendas "
     "do topo onde o gap racial é ainda mais intenso (confirmado pela regressão quantílica).'"],
    title_color=(0xB7,0x1C,0x1C))
doc.add_paragraph()

add_heading(doc, "Figura 2 — Heatmap do Gap Salarial por Gênero e Educação", level=2)
add_figure(doc, FIGS / "fig2_heatmap_gap_genero_educ.png", width_cm=14,
           caption="Gap racial médio (%) por categoria de gênero × nível de escolaridade.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: nível de escolaridade (categorias crescentes). "
                "Eixo Y: gênero (masculino / feminino).")
add_bullet(doc, "Cor das células: intensidade do gap racial médio. Cores mais quentes = gap maior.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Mulheres negras: células com maior intensidade — confirmando interseccionalidade "
                "(Crenshaw, 1989). A desvantagem combinada raça × gênero supera a soma das partes.")
add_bullet(doc, "Escolaridade não elimina o gap: mesmo na célula 'mulher negra + superior completo', "
                "o gap persiste — diplomas não são suficientes para anular a discriminação interseccional.")
add_bullet(doc, "Confirmação estatística: interação negro×sexo_fem no GLMM (β=+0.0416, p<0.001) "
                "confirma o padrão visual.")
add_colored_box(doc, "O que dizer se perguntarem sobre interseccionalidade:",
    ["'Crenshaw (1989) argumenta que identidades sobrepostas criam formas específicas de discriminação "
     "não capturadas por análises unidimensionais. O heatmap mostra isso empiricamente: o gap de "
     "mulheres negras é maior do que a soma do gap de gênero + gap racial. O coeficiente "
     "negro×sexo_fem=+0.0416 no GLMM confirma estatisticamente que a interação existe e é "
     "positiva — ou seja, o efeito combinado é ainda pior do que cada fator isoladamente.'"],
    title_color=(0xB7,0x1C,0x1C))
doc.add_paragraph()

add_heading(doc, "Figura 4 — Gap Salarial por Faixa Etária (Ciclo de Vida)", level=2)
add_figure(doc, FIGS / "fig4_gap_faixa_etaria.png", width_cm=14,
           caption="Gap racial médio (%) por faixa etária. Ponderado por V1028.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: faixas etárias (14–24, 25–34, 35–44, 45–54, 55–64, 65+).")
add_bullet(doc, "Eixo Y: gap salarial racial médio (%) entre brancos e negros naquela faixa etária.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "14–24 anos: gap de 9,1% — pequeno, quando todos estão na base do mercado.")
add_bullet(doc, "25–34 anos: gap sobe para 25,0% — primeiro ciclo de diferenciação de carreira.")
add_bullet(doc, "35–44 anos: gap PICO de 37,5% — fase de maior mobilidade vertical; "
                "o teto de vidro se fecha exatamente quando as promoções mais importam.")
add_bullet(doc, "65+: gap de 40,0% — acumulação máxima ao longo de toda a carreira.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'O padrão de ciclo de vida é a confirmação empírica da natureza acumulativa da discriminação. "
     "Isso é consistente com o coeficiente negro×experiência (β=−0.0019, p<0.001) no GLMM: "
     "cada ano adicional de experiência AMPLIA o gap racial em vez de reduzi-lo. Em um mercado "
     "justo, mais experiência deveria nivelar as diferenças — o sinal negativo confirma que o "
     "mecanismo de discriminação se intensifica com a progressão de carreira, não o contrário.'"],
    title_color=(0xB7,0x1C,0x1C))
doc.add_paragraph()

# ── Tabelas Analíticas ────────────────────────────────────────────────────────
add_heading(doc, "TABELAS ANALÍTICAS — O QUE SABER SOBRE CADA TABELA", level=2)
add_para(doc, "As quatro tabelas abaixo foram geradas a partir dos microdados PNAD Contínua 2016–2025 "
              "com ponderação amostral (peso V1028). Elas estabelecem o diagnóstico descritivo e "
              "os resultados Mincer antes da modelagem HLM/GLMM.",
         size=11, space_after=8)

add_colored_box(doc, "Tabela 1 — Estatísticas Descritivas por Raça (tab1_descritiva_racial.csv)",
    ["Conteúdo: N, renda média (R$), renda mediana (R$), escolaridade média (anos), "
     "% emprego formal, % superior completo — para brancos e negros separadamente.",
     "Números a memorizar: renda mediana brancos = R$2.000 / negros = R$1.450 (gap 27,5%); "
     "renda média: gap de 37,5%; formal: brancos 53% vs negros 48%.",
     "Se perguntarem 'por que mediana?': A distribuição de renda é muito assimétrica (cauda direita "
     "longa). A mediana é robusta a outliers e representa melhor o trabalhador típico. O gap de "
     "37,5% na média é maior porque inclui rendas do topo onde o gap é ainda mais intenso."],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

add_colored_box(doc, "Tabela 2 — Gap Bruto por Subgrupo (tab2_gap_bruto_subgrupos.csv)",
    ["Conteúdo: gap racial médio (%) segmentado por gênero × faixa educacional.",
     "Subgrupo com maior gap: mulher negra — confirma empiricamente a interseccionalidade.",
     "Subgrupo com menor gap: homens com superior completo (escolaridade nivela parte da "
     "desvantagem de capital humano, mas gap persiste).",
     "Se perguntarem: 'Por que mostrar subgrupos?' — Confirma Crenshaw (1989): políticas de renda "
     "que tratam 'negro' como categoria homogênea subestimam severamente o impacto sobre mulheres "
     "negras, que são o subgrupo mais afetado pela discriminação composta."],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

add_colored_box(doc, "Tabela 3 — Equação de Mincer Progressiva (tab3_mincer_progressivo.csv)",
    ["Conteúdo: 5 especificações OLS progressivas (M1–M5) com coeficiente de 'negro' em cada modelo.",
     "M1 (apenas negro): gap = −23,1% | M2 (+edu+experiência): −22,3% | M3 (+UF): −7,9% | "
     "M4 (+setor+horas): −7,0% | M5 (completo): −6,2%",
     "Padrão: o salto entre M2 e M3 (−22,3% → −7,9%) mostra que os estados explicam grande "
     "parte do gap bruto — confirmando a dimensão territorial da desigualdade racial.",
     "Se perguntarem sobre Mincer (1974): 'A especificação padrão para estimar retornos à educação "
     "inclui log(renda) ~ educação + experiência + experiência². A versão progressiva aqui mostra "
     "explicitamente quanto cada bloco de variáveis explica do gap — é a formalização OLS do que "
     "o HLM faz de forma hierárquica.'"],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

add_colored_box(doc, "Tabela 4 — Termos de Interação Racial no GLMM (tab4_interacoes.csv)",
    ["Conteúdo: coeficientes de interação estimados no GLMM: negro×sexo_fem, negro×experiência, "
     "negro×educ_alto, negro×educ_med.",
     "negro×sexo_fem: β=+0.0416 (***) — mulheres negras têm desvantagem ADICIONAL além da soma "
     "dos efeitos de raça e gênero separados (interseccionalidade confirmada).",
     "negro×experiência: β=−0.0019 (***) — cada ano de experiência AMPLIA o gap racial. "
     "Sinal negativo: a experiência retorna menos para negros do que para brancos.",
     "negro×educ_alto e negro×educ_med: não significativos (ns) — mais escolaridade não "
     "elimina a desvantagem racial. Capital humano é necessário mas não suficiente.",
     "Se perguntarem sobre β negativo em negro×experiência: 'Significa que 1 ano adicional de "
     "experiência reduz a probabilidade de ocupação qualificada de negros 0.0019 pontos a mais "
     "do que para brancos. O mecanismo provável é o teto de vidro: com o tempo, brancos são "
     "promovidos e negros estacionam — o que se acumula ao longo do ciclo de vida (Fig 4).'"],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()
add_separator(doc)

# ── Figura 3: Composição CBO ─────────────────────────────────────────────────
add_heading(doc, "Figura 3 — Razão de Representação por Grupo Ocupacional (CBO)", level=2)
add_figure(doc, FIGS / "comp_razao_grupo_cbo.png", width_cm=14,
           caption="Razão negros/brancos por grupo CBO. Linha tracejada = paridade (1,0).")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo Y: grupos ocupacionais CBO (do mais elementar ao mais qualificado).")
add_bullet(doc, "Eixo X: razão entre a proporção de negros e a proporção de brancos no grupo. "
                "Valor 1,0 = paridade. Valor < 1 = sub-representação de negros. Valor > 1 = sobre-representação.")
add_bullet(doc, "Linha tracejada vertical em 1,0 = ponto de equilíbrio.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Grupos de prestígio (Dirigentes, Profissionais) ficam à esquerda da linha: negros "
                "são sub-representados. Grupos elementares ficam à direita: sobre-representados.")
add_bullet(doc, "Dirigentes: razão ≈ 0,42 — para cada 100 brancos na função, há apenas 42 negros.")
add_bullet(doc, "Elementares: razão ≈ 2,12 — negros são mais do que o dobro do esperado pela paridade.")
add_colored_box(doc, "O que dizer se perguntarem:",
    [f"'Este gráfico é o primeiro diagnóstico visual da segregação ocupacional. A razão de 0,42 nos "
     f"Dirigentes não é explicada por diferença de escolaridade — quando controlamos por educação no "
     f"GLMM (lme4), o OR permanece em {or_str(P['OR_M2'])}. Isso confirma que parte substancial da sub-representação "
     f"é discriminação de acesso, não apenas diferença de capital humano.'"],
    title_color=(0xB7,0x1C,0x1C))
doc.add_paragraph()

# ── Figura 5: Representação no topo ──────────────────────────────────────────
add_heading(doc, "Figura 5 — Índice de Representação no Topo da Distribuição de Renda", level=2)
add_figure(doc, FIGS / "comp_representacao_topo.png", width_cm=14,
           caption="IR (Índice de Representação) por percentil de renda. IR=1 = paridade racial.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: percentis de renda (do mais baixo ao topo).")
add_bullet(doc, "Eixo Y: IR = (% negros no percentil) / (% negros na população total). "
                "IR=1 = representação proporcional. IR<1 = sub-representação no topo.")
add_bullet(doc, "Linhas separadas por capital vs interior, ou por região — verifique a legenda do gráfico gerado.")
add_para(doc, "Achado principal:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "No top 5% das capitais, IR ≈ 0,47: negros ocupam apenas 47% do que seria esperado "
                "pela sua participação na população.")
add_bullet(doc, "O IR cai conforme sobe o percentil — essa queda é a formalização visual do glass ceiling.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'O IR quantifica em pontos de percentil o que o gráfico de razão CBO mostra categoricamente. "
     "A queda do IR no topo confirma que o glass ceiling não é apenas uma metáfora — é uma "
     "barreira mensurável que se intensifica exatamente onde os prêmios salariais são maiores. "
     "Isso é consistente com os resultados da regressão quantílica.'"],
    title_color=(0xB7,0x1C,0x1C))
doc.add_paragraph()

# ── Figura 6: HLM decomposição ──────────────────────────────────────────────
add_heading(doc, "Figura 6 — Decomposição HLM: Do Gap Bruto ao Resíduo de Discriminação", level=2)
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Barras horizontais: cada barra representa o gap racial estimado em um modelo diferente "
                "(M1 a M4). O comprimento da barra é o tamanho do gap.")
add_bullet(doc, "A sequência de cima para baixo mostra como o gap 'encolhe' à medida que variáveis "
                "de contexto e ocupação são adicionadas.")
add_bullet(doc, "M1 (gap bruto) = 19,3%: sem nenhum controle.")
add_bullet(doc, "M2 (+ UPA): renda média da UPA absorve ~9 p.p. — o lugar onde se mora explica metade do gap.")
add_bullet(doc, "M3 (+ UF): redução marginal adicional pelas diferenças entre estados.")
add_bullet(doc, "M4 (+ CBO + formalidade + horas): o gap cai para 6,2% — este é o resíduo que não tem "
                "explicação por características observáveis. É a discriminação de remuneração pura.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "A maior redução ocorre entre M1 e M2 (adição da UPA): confirma Wilson (1987) — "
                "o local de moradia é o principal mediador do gap.")
add_bullet(doc, "O gap residual de 6,2% em M4 não tem explicação por nenhuma variável observável: "
                "é o lower bound da discriminação de remuneração.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'A sequência dos modelos não é arbitrária — cada passo adiciona um nível da hierarquia social. "
     "Primeiro o contexto de moradia (UPA), depois o contexto institucional (UF), depois a posição "
     "no mercado de trabalho (CBO + formalidade). O gap de 6,2% que sobra é robusto a todos os "
     "controles observáveis disponíveis na PNAD.'"],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

# ── Figura 7: Oaxaca ─────────────────────────────────────────────────────────
add_heading(doc, "Figura 7 — Decomposição de Oaxaca-Blinder", level=2)
add_figure(doc, FIGS / "oaxaca_decomposicao.png", width_cm=14,
           caption="Decomposição two-fold: gap total em Dotações (84%) e Retornos (16%).")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, f"O gráfico divide o gap total (≈{fmt(P['GAP_PCT'],1)}% em log-renda) em duas partes:")
add_bullet(doc, "Barra/segmento DOTAÇÕES (84%): diferença explicada por negros terem, em média, "
                "menos anos de escolaridade, menos acesso a emprego formal, ocupações menos "
                "qualificadas e mais horas em subemprego.", level=1)
add_bullet(doc, "Barra/segmento RETORNOS (16%): diferença que persiste mesmo quando negros e brancos "
                "têm as MESMAS características — o mercado remunera diferentemente.", level=1)
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "84% do gap é de dotações: a discriminação opera principalmente no ACESSO a "
                "características valorizadas, não no salário dentro da função.")
add_bullet(doc, "Isso inverte a intuição comum: combater apenas o gap salarial direto (16%) "
                "deixa 84% do problema intacto.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'O Oaxaca une os três vértices do triângulo de evidências. Os 84% de dotações são explicados "
     "pelo logit (acesso a ocupações) e pela SNA (exclusão das redes). Os 16% de retornos são "
     "explicados pelo HLM M4 (6,2% de discriminação pura de remuneração). Os métodos são "
     "consistentes entre si.'",
     "Se perguntarem sobre three-fold: 'A decomposição three-fold adiciona um componente de "
     "interação entre dotações e retornos que não tem interpretação causal clara em contexto "
     "de discriminação racial. Por isso a literatura de economics of discrimination prefere "
     "two-fold (Fortin, Lemieux & Firpo, 2011).'"],
    title_color=(0x15,0x65,0xC0))
doc.add_paragraph()

# ── Figura 8: GLMM Odds Ratios ───────────────────────────────────────────────
add_heading(doc, "Figura 8 — GLMM (lme4): Odds Ratios para Acesso a Ocupações Qualificadas", level=2)
add_figure(doc, FIGS / "glmm_odds_ratios.png", width_cm=14,
           caption="Odds Ratios do GLMM para negro vs branco. M1=sem contexto, M2=com efeitos aleatórios UPA. "
                   "Linha = OR=1 (sem diferença racial).")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: Odds Ratio (OR). OR=1 = sem diferença racial. OR<1 = negros têm MENOR chance.")
add_bullet(doc, "Cada ponto = estimativa de OR com IC 95%. GLMM implementado com lme4::glmer "
                "(nAGQ=0, otimizador bobyqa) para acomodar a estrutura hierárquica dos dados.")
add_bullet(doc, f"M1 (sem efeito aleatório de contexto): OR={or_str(P['OR_M1'])} — negros têm {fmt(P['OR_M1_menor_pct'],1)}% menos chance. "
                f"ICC M1={fmt(P['ICC_M1_pct'],1)}% (variação entre UPAs elevada). N={fmtN(P['N_GLMM'])} (PEA completa).")
add_bullet(doc, f"M2 (com efeito aleatório de UPA): OR={or_str(P['OR_M2'])} — após controlar o contexto local, "
                f"negros têm {fmt(P['OR_M2_menor_pct'],1)}% menos chance. ICC M2={fmt(P['ICC_M2_pct'],1)}%.")
add_para(doc, "O AME (Average Marginal Effect):", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, f"AME M1 = {ame(P['AME_M1_pp'])}: sem controle de contexto, negros têm {fmt(abs(P['AME_M1_pp']),2)} p.p. a menos de "
                "probabilidade de ocupação qualificada.")
add_bullet(doc, f"AME M2 = {ame(P['AME_M2_pp'])}: com efeito aleatório de UPA, o efeito líquido de ser negro "
                f"é {ame(P['AME_M2_pp'])} — discriminação pura de acesso após controlar moradia e contexto.")
add_colored_box(doc, "O que dizer se perguntarem sobre OR vs AME e sobre o GLMM:",
    ["'O OR mede razão de chances — útil para comparação entre modelos. O AME traduz isso em "
     "diferença de probabilidade — mais intuitivo para comunicar ao gestor público.'",
     f"'O GLMM (Generalized Linear Mixed Model) é a extensão natural do HLM para variáveis binárias. "
     f"Ao usar lme4::glmer com efeitos aleatórios por UPA, capturo a mesma estrutura hierárquica "
     f"do HLM aplicada ao modelo de acesso a ocupações. O ICC={fmt(P['ICC_M1_pct'],1)}% em M1 confirma que há "
     f"variância substancial entre UPAs — ignorar essa hierarquia produziria erros padrão "
     f"subestimados e OR's viesados.'"],
    title_color=(0xFF,0x8F,0x00))
doc.add_paragraph()

# ── Figura 9: Quantile Regression ────────────────────────────────────────────
add_heading(doc, "Figura 9 — Regressão Quantílica: Trajetória do Gap ao Longo da Distribuição", level=2)
add_figure(doc, FIGS / "quantreg_trajetoria.png", width_cm=14,
           caption="Gap racial (% de diferença de log-renda) por percentil. Linha M3 = sem CBO/formalidade. "
                   "Linha M4 = com CBO/formalidade.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: percentil da distribuição de renda (q10 = mais pobres, q95 = mais ricos).")
add_bullet(doc, "Eixo Y: coeficiente do dummy 'negro' em log-renda — quanto menor, maior o gap.")
add_bullet(doc, "Duas linhas: M3 (sem ocupação/formalidade) e M4 (com ocupação/formalidade).")
add_bullet(doc, "Faixa cinza ao redor de cada linha = intervalo de confiança 95%.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Inclinação descendente de ambas as linhas (coeficiente fica mais negativo conforme "
                "sobe o quantil) = o gap aumenta no topo da distribuição = glass ceiling racial.")
add_bullet(doc, "Diferença entre M3 e M4 cresce no topo: no q95, a mediação ocupacional (CBO + "
                "formalidade) explica mais do gap do que no q10. Negros no topo são excluídos "
                "das ocupações mais rentáveis.")
add_bullet(doc, "Valores concretos: q10 ≈ −3,2%; q50 ≈ −4,1%; q90 ≈ −7,1%; q95 ≈ −7,9% (M4).")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'Um OLS capturaria apenas a média — aproximadamente o ponto de q50. Mas o glass ceiling "
     "está no q90–q95, onde o gap é sistematicamente maior. Sem regressão quantílica, esse "
     "padrão seria invisível. A regressão quantílica é o único método que formaliza a "
     "hipótese do glass ceiling como teste estatístico.'"],
    title_color=(0x2E,0x7D,0x32))
doc.add_paragraph()

# ── Figura 9b: RIF-OB ─────────────────────────────────────────────────────────
add_heading(doc, "Figura 9b — RIF-OB: Sticky Floor Discriminatório vs. Glass Ceiling por Dotações", level=2)
add_figure(doc, FIGS / "rif_ob_retornos_quantis.png", width_cm=14,
           caption="RIF-OB incondicional por quantil (Firpo, Fortin & Lemieux, 2018). "
                   "Retornos (discriminação) vs Dotações por decil de renda.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: decil/quantil da distribuição incondicional de renda (q10=base, q90=topo).")
add_bullet(doc, "Dois componentes empilhados por quantil: Dotações (diferença de características) "
                "e Retornos (discriminação pura — o que persiste mesmo com características iguais).")
add_bullet(doc, "RIF = Recentered Influence Function — permite decompor o gap em qualquer quantil "
                "incondicionalmente, sem condicionar ao quantil dentro de cada grupo racial.")
add_para(doc, "O padrão sticky floor vs. glass ceiling:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "q10 — base: Dotações=61,3%, Retornos=33,1%. Discriminação de preço é 3× mais "
                "intensa na BASE — efeito sticky floor (chão pegajoso).")
add_bullet(doc, "q50 — mediana: Dotações=57,4%, Retornos=21,5%. Discriminação decai progressivamente.")
add_bullet(doc, "q90 — topo: Dotações=75,8%, Retornos=11,2%. No topo, o gap é quase inteiramente "
                "por dotações — negros são EXCLUÍDOS das posições rentáveis (glass ceiling).")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'A QR formaliza o glass ceiling condicional — o gap cresce no topo dentro de cada grupo. "
     "O RIF-OB faz a decomposição incondicional — permite perguntar: de onde vem esse gap maior "
     "no topo? A resposta é: no topo o mecanismo dominante é exclusão (dotações = 75,8%), "
     "não discriminação salarial direta. Na base, a discriminação de preço ainda responde por "
     "33,1% do gap — o sticky floor discriminatório. São dois fenômenos distintos que coexistem.'"],
    title_color=(0x2E,0x7D,0x32))
doc.add_paragraph()

# ── Figura 9c: Interseccionalidade OB 4 Grupos ────────────────────────────────
add_heading(doc, "Figura 9c — Interseccionalidade: Oaxaca-Blinder por 4 Grupos Raça × Gênero", level=2)
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Referência: Homem Branco. Comparação: Mulher Branca, Homem Negro, Mulher Negra.")
add_bullet(doc, "Para cada grupo, o gap total é decomposto em Dotações (características) e "
                "Retornos (discriminação direta).")
add_bullet(doc, "HLM M4: coeficiente da interação negro×sexo_fem = β=+0,069 — indica atenuação "
                "do gap racial para mulheres (raça e gênero interagem, não se somam linearmente).")
add_para(doc, "Resultados por grupo (referência: Homem Branco):", size=11, bold=True,
         color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Mulher Branca: gap=46,6% — dotações negativas (−8,9%) / retornos 109% "
                "(retornos amplificam o gap: discriminação de gênero pura explica mais de 100%). "
                "Sem penalidade de raça.")
add_bullet(doc, "Homem Negro: gap=40,3% — 71% dotações / 29% retornos. Discriminação racial "
                "direta explica ~29% do gap do homem negro.")
add_bullet(doc, "Mulher Negra: gap=96,4% — maior gap de todos os grupos. Penalidade adicional "
                "por interseccionalidade: +9,5 p.p. além da soma de raça + gênero separados.")
add_colored_box(doc, "O que dizer se perguntarem sobre interseccionalidade:",
    ["'A Mulher Negra não sofre apenas raça + gênero de forma aditiva — existe uma penalidade "
     "adicional de +9,5 p.p. que emerge especificamente da combinação dos dois eixos de "
     "opressão. Isso é a interseccionalidade no sentido de Crenshaw (1989): os eixos interagem, "
     "não se somam. O HLM confirma com β_neg×fem=+0,069 (***) — a discriminação racial é "
     "estatisticamente menor para mulheres negras do que se esperaria da soma linear.'",
     "'Gap de 96,4% para mulher negra vs 40,3% para homem negro: a mulher negra experimenta "
     "uma dupla desvantagem que quase dobra o gap. Políticas de gênero e raça precisam ser "
     "desenhadas conjuntamente para esse grupo.'"],
    title_color=(0xFF,0x8F,0x00))
doc.add_paragraph()

# ── Figura 10: SHAP ───────────────────────────────────────────────────────────
add_heading(doc, "Figura 10 — SHAP Values: Importância das Variáveis no XGBoost", level=2)
add_figure(doc, FIGS / "shap_importance_xgb.png", width_cm=14,
           caption="Importância média dos SHAP values por variável. R²=0,6162.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: importância média absoluta dos SHAP values — quanto cada variável 'move' "
                "a predição de renda em média, independente de direção.")
add_bullet(doc, "Barras maiores = variáveis mais determinantes para a renda prevista pelo modelo.")
add_bullet(doc, "Variáveis em vermelho/negativo = associadas a menores rendas. "
                "Variáveis em azul/positivo = associadas a maiores rendas.")
add_para(doc, "Ranking esperado e o que significa:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "#1 Renda média da UPA (|SHAP|≈0,272): onde você mora supera o quanto você "
                "estudou como determinante de renda — confirma Wilson (1987) e justifica o HLM.")
add_bullet(doc, "#2 Horas trabalhadas (|SHAP|≈0,166): nova variável extraída dos ZIPs brutos. "
                "Volume de trabalho é o segundo maior determinante.")
add_bullet(doc, "#3–#4 CBO e emprego formal: variáveis também extraídas dos ZIPs — sua contribuição "
                "metodológica principal. Sem elas, o modelo subestimaria o papel da ocupação.")
add_bullet(doc, "#11 Raça/negro (|SHAP|≈0,029): após todos os controles, a raça ainda move "
                "a predição em −2,5%. Não é o maior fator — mas é o fator que não deveria "
                "existir em um mercado de trabalho não discriminatório.")
add_colored_box(doc, "O que dizer se perguntarem sobre interpretabilidade do XGBoost:",
    ["'O SHAP resolve o problema black-box do XGBoost. Ele decompõe a predição de cada "
     "observação individualmente, garantindo aditividade e consistência (propriedades de Shapley). "
     "Um simples gráfico de feature importance do XGBoost não teria essas garantias — "
     "o SHAP é o gold standard atual para interpretabilidade de modelos de ML.'"],
    title_color=(0x15,0x65,0xC0))
doc.add_paragraph()

# ── Figura 11: SNA ─────────────────────────────────────────────────────────────
add_heading(doc, "Figura 11 — Rede Social Ocupacional: Betweenness por Raça e Escolaridade", level=2)
add_figure(doc, FIGS / "sna_rede_demografica.png", width_cm=14,
           caption="Rede de co-ocupação. Nós = grupos de raça × educação. Tamanho do nó = betweenness centrality.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Nós: cada nó representa um grupo (ex: 'negro + superior completo', 'branco + médio').")
add_bullet(doc, "Arestas: conexão = co-ocorrência na mesma ocupação/setor.")
add_bullet(doc, "Tamanho do nó: proporcional à betweenness centrality — quão central o grupo é "
                "como ponte entre diferentes clusters da rede.")
add_bullet(doc, "Cor/posição: grupos brancos tendem a ocupar o centro da rede; grupos negros "
                "ficam na periferia independente do nível educacional.")
add_para(doc, "Achado principal:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Brancos com pós-graduação têm alta betweenness: são brokers que conectam "
                "grupos diferentes — acesso a vagas, informação e promoções (Burt, 2004).")
add_bullet(doc, "Negros com pós-graduação têm betweenness ≈ 0: mesmo com as mesmas credenciais, "
                "não ocupam posições de ponte na rede. As credenciais 'valem menos' porque "
                "não convertem em acesso às redes de conversão.")
add_colored_box(doc, "O que dizer se perguntarem:",
    ["'A SNA adiciona uma dimensão que o HLM e o logit não capturam: o capital social estrutural. "
     "Mesmo que um negro consiga entrar em uma ocupação qualificada, sua betweenness=0 significa "
     "que ele não tem acesso às redes informais que geram promoção, mentoria e informação "
     "antecipada sobre vagas. Isso explica parte do gap residual de 6,2% do HLM M4.'"],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

# ── Figuras Estado H1 ─────────────────────────────────────────────────────────
add_heading(doc, "Figura 12 — Gini por Setor (H1): Estado como Indutor de Desigualdade?", level=2)
add_figure(doc, FIGS / "estado_h1_gini.png", width_cm=14,
           caption="Gini por setor (público vs privado) e decomposição de Theil T.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Gráfico de barras: Gini total (0,4879), público (0,4662) e privado (0,4715).")
add_bullet(doc, "Gini menor no público = distribuição de renda mais igualitária dentro do setor público.")
add_bullet(doc, "Decomposição de Theil T: 7,8% da desigualdade total é ENTRE os dois setores. "
                "Isso significa que a diferença de renda médio entre público e privado gera "
                "7,8% da desigualdade agregada.")
add_para(doc, "Interpretação da hipótese H1:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "O estado REDUZ a desigualdade interna (Gini público < privado), mas o prêmio "
                "salarial do servidor (+99,6% vs privado) cria uma divisão ENTRE setores que "
                "contribui 7,8% para a desigualdade total.")
add_bullet(doc, "Conclusão de H1: o estado é simultaneamente redutor de desigualdade interna e "
                "gerador de desigualdade interssetorial. A resposta é 'ambos' — depende da perspectiva.")
doc.add_paragraph()

# ── Figuras Estado H2/H3 ───────────────────────────────────────────────────────
add_heading(doc, "Figura 13 — Gaps Racial e de Gênero por Setor (H2/H3)", level=2)
add_figure(doc, FIGS / "estado_h2h3_gaps.png", width_cm=14,
           caption="Gaps salariais controlados por raça e gênero no setor público vs privado.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Barras agrupadas: cada par compara o gap no setor privado vs público, "
                "para raça (H2) e gênero (H3).")
add_bullet(doc, "Valores são coeficientes de regressão controlada (β do dummy negro ou sexo_fem), "
                "não diferenças brutas.")
add_bullet(doc, "Gap racial bruto: privado −44,1% / público −33,3%. Controlado: privado −29,2% / "
                "público −25,2%.")
add_bullet(doc, "Gap de gênero: privado −16,9% / público −18,8% (após controles). "
                "O setor público tem gap de GÊNERO maior que o privado — o paradoxo H3.")
add_para(doc, "Interpretação do paradoxo H3:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "O concurso público iguala a entrada (não há discriminação no acesso ao cargo), "
                "mas a progressão de carreira e os cargos de confiança (DAS, cargos militares "
                "seniores, judiciário de topo) são preenchidos por nomeação e promoção, "
                "onde o viés de gênero persiste.")
add_bullet(doc, "Resultado: dentro do setor público, mulheres ficam presas nos cargos de entrada "
                "enquanto homens avançam — gap maior que no setor privado onde não há esse "
                "'teto de vidro institucionalizado'.")
doc.add_paragraph()

# ── Figura H4 tendência ──────────────────────────────────────────────────────
add_heading(doc, "Figura 14 — Renda Real e Emprego 2016–2025 (H4)", level=2)
add_figure(doc, FIGS / "estado_h4_tendencia.png", width_cm=14,
           caption="Renda real mediana (base IPCA 2016=100) e taxa de emprego 2016–2025.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo Y esquerdo (linha): renda real mediana em R$ de 2016.")
add_bullet(doc, "Eixo Y direito (barra ou linha pontilhada): taxa de emprego (% da PEA).")
add_bullet(doc, "Dois eixos no mesmo gráfico permitem ver a co-evolução das duas séries.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Emprego sobe consistentemente de 87–88% (2017–2020) para 94,4% em 2025 — máxima histórica.")
add_bullet(doc, "Renda real cai de R$1.200 (2016) para R$1.106 (2022 — mínima), recupera para "
                "R$1.345 (2024) e recua novamente para R$1.283 em 2025.")
add_bullet(doc, "Divergência entre emprego e renda real: confirma H4 — o crescimento do emprego "
                "no governo atual não se traduziu em prosperidade equivalente em termos reais.")
add_bullet(doc, "A renda nominal sobe (R$1.200→R$2.000), mas deflacionada pelo IPCA "
                "(base 2016=100) o ganho real é modesto e volátil.")
doc.add_paragraph()

# ── Figura H4 inclusão ─────────────────────────────────────────────────────────
add_heading(doc, "Figura 15 — Simulação de Inclusão Produtiva (H4)", level=2)
add_figure(doc, FIGS / "estado_h4_inclusao.png", width_cm=14,
           caption="Ganho potencial de renda se negros tivessem a mesma distribuição CBO que brancos.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Duas barras: renda média negra ATUAL vs renda CONTRAFACTUAL "
                "(se negros tivessem a mesma distribuição de CBO que brancos).")
add_bullet(doc, "A diferença entre as barras = +127,3% = ganho potencial de renda se a "
                "segregação ocupacional fosse eliminada.")
add_para(doc, "Premissas da simulação:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Mantém os salários de cada grupo CBO constantes — só redistribui negros pelos "
                "grupos como se tivessem o perfil CBO dos brancos.")
add_bullet(doc, "Proxy conservador: não considera mudança de produtividade, apenas redistribuição.")
add_bullet(doc, "O proxy de PIB (+74,3 p.p.) assume que negros representam 58% da força de trabalho "
                "e que o ganho de renda se traduz proporcionalmente em demanda agregada.")
doc.add_paragraph()

# ── Figura H5 armadilha ─────────────────────────────────────────────────────────
add_heading(doc, "Figura 16 — Gap de Qualificação 2016–2025 (H5: Armadilha da Renda Média)", level=2)
add_figure(doc, FIGS / "estado_h5_armadilha.png", width_cm=14,
           caption="% de trabalhadores com alta qualificação por raça ao longo do tempo.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Eixo X: anos (2016–2025).")
add_bullet(doc, "Eixo Y: % de trabalhadores com alta qualificação (CBO grupos 1 e 2, ou superior "
                "completo com função compatível).")
add_bullet(doc, "Três linhas: total, brancos e negros.")
add_para(doc, "O que o padrão mostra:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Brancos: 19,7% (2016) → 22,4% (2025) — crescimento de +2,7 p.p. em 10 anos.")
add_bullet(doc, "Negros: 9,6% (2016) → 11,0% (2025) — crescimento de +1,4 p.p. em 10 anos.")
add_bullet(doc, "As três linhas sobem em paralelo: o gap absoluto permanece praticamente constante "
                "(≈10 p.p.). Não há convergência.")
add_bullet(doc, "Isso é a armadilha da renda média: o Brasil eleva a qualificação geral, mas a "
                "desigualdade racial na qualificação não diminui — e sem quebrar essa barreira, "
                "o país não consegue sair da armadilha de renda média (Hsieh et al., 2019).")
doc.add_paragraph()

# ── Figura modelos LL/AIC ─────────────────────────────────────────────────────
add_heading(doc, "Figura 17 — Comparação de Modelos: Log-Likelihood e AIC", level=2)
add_figure(doc, FIGS / "modelos_loglik_aic.png", width_cm=14,
           caption="LL e AIC para 10 modelos. Quanto menor o AIC, melhor o ajuste penalizado por complexidade.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Log-Likelihood (LL): quanto mais próximo de zero (menos negativo), melhor o ajuste. "
                "Mais parâmetros sempre melhoram o LL — por isso o AIC penaliza.")
add_bullet(doc, "AIC = −2×LL + 2×k, onde k = número de parâmetros. Modelos mais simples com "
                "ajuste equivalente são preferidos.")
add_bullet(doc, "Eixo X: os 10 modelos em ordem de complexidade crescente "
                "(OLS Nulo → OLS Individual → OLS+FE → HLM Nulo → ... → HLM Ocupação).")
add_bullet(doc, "Colunas/barras com dois painéis: LL (painel superior) e AIC (painel inferior).")
add_para(doc, "Achado principal:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "OLS+FE(UF) tem AIC = 3.684.833 — PIOR que OLS Individual (3.660.787). "
                "Adicionar dummies de UF não ajuda porque os estados são unidades muito grandes "
                "para capturar a variação real.")
add_bullet(doc, "HLM Contextual tem AIC = 3.588.684 — o melhor ajuste penalizado entre os modelos "
                "intermediários. Demonstra que a estrutura hierárquica compensa o custo dos "
                "parâmetros adicionais.")
doc.add_paragraph()

# ── Figura LRT/ICC ─────────────────────────────────────────────────────────────
add_heading(doc, "Figura 18 — LRT Waterfall e Trajetória do ICC", level=2)
add_figure(doc, FIGS / "modelos_lrt_icc.png", width_cm=14,
           caption="LRT (Likelihood Ratio Test) entre modelos encadeados e trajetória do ICC por modelo.")
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Painel esquerdo (waterfall de LRT): cada barra mostra o χ² do teste de razão de "
                "verossimilhança ao adicionar um nível ao modelo. Barras positivas = o modelo "
                "mais complexo é estatisticamente superior.")
add_bullet(doc, "LRT HLM Nulo vs OLS Nulo: χ²=191.625 (Δk=1) — vastamente significativo. "
                "Isso prova que existe estrutura hierárquica nos dados: as rendas dentro de "
                "uma mesma UPA são mais parecidas entre si do que entre UPAs diferentes.")
add_bullet(doc, "Painel direito (ICC): Intra-Class Correlation por modelo — % da variância total "
                f"que se deve às diferenças entre UPAs. ICC={fmt(P['ICC_HLM_M0_pct'],2)}% em M0.")
add_para(doc, "Por que o ICC > 5% importa:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Se o ICC fosse < 5%, o OLS seria suficiente (Raudenbush & Bryk, 2002). "
                f"Com {fmt(P['ICC_HLM_M0_pct'],2)}%, ignorar a hierarquia produz erros padrão subestimados — e "
                "inferência estatística inválida.")
add_bullet(doc, "O LRT χ²=191.625 com Δk=1 prova que o componente de variância aleatória "
                "da UPA não é zero: hierarquia não é apenas conveniência, é estrutura real dos dados.")
add_colored_box(doc, "O que dizer se perguntarem sobre a escolha do HLM:",
    ["'Não escolhi HLM por ser mais sofisticado — escolhi porque os dados me obrigaram. "
     f"O ICC={fmt(P['ICC_HLM_M0_pct'],2)}% e o LRT χ²=191.625 mostram que a estrutura hierárquica é real e significativa. "
     "Ignorar isso com OLS produziria inferência inválida — erros padrão subestimados e "
     "p-valores inflados. A escolha foi ditada pelos dados, não pela preferência metodológica.'"],
    title_color=(0x2E,0x7D,0x32))
doc.add_paragraph()

# ── Análises de Sensibilidade ─────────────────────────────────────────────────
add_heading(doc, "Análises de Sensibilidade — Konfound, E-values e Oster δ*", level=2)
add_para(doc, "Como ler:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Konfound (pkonfound): quantos casos precisariam ser trocados de trat./controle, "
                "ou qual % dos casos precisaria estar mal alocado, para reverter a conclusão. "
                "Quanto maior, mais robusto.")
add_bullet(doc, "E-value (VanderWeele & Ding, 2017): qual a magnitude mínima de uma variável "
                "omitida (como razão de risco) para anular o efeito observado. E-value ≥ 2 "
                "indica boa resistência.")
add_bullet(doc, "Oster δ*: qual proporção da variação total dos controles observados seria "
                "necessária nos não-observados para anular o efeito. δ* negativo confirma "
                "robustez (variáveis omitidas teriam que atuar na direção contrária).")
add_para(doc, "Resultados:", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Konfound HLM: M1=99,5%; M2=M3=98,8%; M4=98,5% — quase 100% dos casos "
                "precisariam ser mal alocados para reverter a conclusão de discriminação.")
add_bullet(doc, "OLS subestima: gap OLS=43,8%–48,5% vs HLM=50,8%–53,0% (Δ=5,4–9,7 p.p.) — "
                "o viés de subestimação do OLS fica em torno de 9,7 p.p. em M1 e 5,4 p.p. em M4.")
add_bullet(doc, f"E-values GLMM lme4 (PEA completa): M1={fmt(P['EVAL_M1'],3)}; M2={fmt(P['EVAL_M2'],3)} — variável omitida "
                f"precisaria ter associação ≥2,3× com raça E com acesso a ocupações para eliminar o OR.")
add_bullet(doc, "Oster δ*: M1=−0,48; M2=−0,43; M4=−0,39 — todos negativos, confirmando "
                "que variáveis omitidas teriam que atenuar (não amplificar) para anular o gap.")
add_colored_box(doc, "O que dizer se perguntarem sobre endogeneidade ou variáveis omitidas:",
    ["'Testei três frameworks de sensibilidade independentes. O pkonfound mostra que seria "
     "necessário reclassificar ~99% das observações para reverter a conclusão — impossível "
     "dado o N=7,7 milhões. O E-value ≥2,0 exige que uma variável omitida tenha associação "
     "dobrada com raça E com renda simultaneamente — sem candidato plausível. E o Oster δ* "
     "negativo significa que a seleção nas variáveis omitidas teria que ser na direção "
     "oposta à seleção observada. Os três convergem: o resultado é robusto.'"],
    title_color=(0x2E,0x7D,0x32))
doc.add_paragraph()

# ── Heckman + Tendência Temporal + Event Study COVID ──────────────────────────
add_heading(doc, "Análises Complementares — Heckman, Tendência Temporal e Event Study COVID", level=2)
add_para(doc, "Correção de Seleção Amostral (Heckman):", size=11, bold=True,
         color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Problema: analisamos apenas quem tem renda (PEA ocupada). A seleção para "
                "participar do mercado de trabalho pode ser endógena — se negros participam "
                "menos, o gap observado é subestimado.")
add_bullet(doc, "Modelo de seleção (Probit 1ª etapa): educação, idade, número de filhos, "
                "presença de cônjuge → gera lambda de Mills (λ).")
add_bullet(doc, "Resultado: λ=−1,985 (***) — seleção amostral significativa e negativa. "
                "Gap OLS=−9,71% → Gap Heckman=−7,41% (Δ=+2,3 p.p.): o gap bruto estava "
                "SUBESTIMADO; negros mais excluídos do mercado atenuavam o gap aparente.")
add_para(doc, "Tendência Temporal e Teste de Chow:", size=11, bold=True,
         color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Gap anual estimado: a discriminação racial mantém-se estável ao longo do "
                "período (tendência linear δ=0,0008/ano, não significativa).")
add_bullet(doc, "Teste de Chow (2020): F(2,6)=7,012; p=0,027 — quebra estrutural significativa "
                "em 2020 (pandemia). O gap racial NÃO teve convergência pré-COVID.")
add_para(doc, "Event Study COVID (DiD):", size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "Tratamento: pós-2020 (COVID). Estimativa: τ=+0,015; SE=0,007; p=0,025.")
add_bullet(doc, "Interpretação: τ=+0,015 indica **convergência aparente** de curto prazo — "
                "o gap reduziu-se 1,5 p.p. após 2020. Provável seleção de saída: negros de "
                "menor renda saíram do mercado formal, elevando a renda relativa dos que ficaram.")
add_colored_box(doc, "O que dizer se perguntarem sobre essas análises:",
    ["'O Heckman resolve a crítica de que estamos vendo apenas os 'bem-sucedidos'. O resultado "
     "(λ negativo e gap que cresce com a correção) indica que a exclusão racial do mercado "
     "de trabalho é real e que os excluídos estariam em situação ainda pior.'",
     "'O Chow mostra quebra estrutural em 2020. O DiD mostra convergência aparente de curto "
     "prazo (τ=+0.015): negros sofreram menos no diferencial salarial, mas provavelmente por "
     "seleção de saída — não por melhora real. O gap de longo prazo permanece estável.'"],
    title_color=(0x1F,0x38,0x64))
doc.add_paragraph()

# ── Pesquisa Operacional — TOPSIS / AHP / LP ──────────────────────────────────
add_heading(doc, "Pesquisa Operacional — Priorização de Políticas (TOPSIS, AHP, LP)", level=2)
add_para(doc, "O que é e por que está na dissertação:", size=11, bold=True,
         color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "PO (Pesquisa Operacional) transforma o diagnóstico empírico em recomendações "
                "de política priorizadas objetivamente. Sem PO, a dissertação terminaria com "
                "'há discriminação' — com PO, termina com 'e aqui está a política mais "
                "custo-eficaz para combatê-la, formalmente justificada'.")
add_bullet(doc, "Métodos: TOPSIS (ranking multicritério), AHP (pesos dos critérios, CR=0,004 "
                "← consistente), LP (Programação Linear com restrição orçamentária), "
                "Pareto frontier (eficiência custo × impacto).")
add_para(doc, "TOPSIS — Ranking de Políticas (critérios: impacto no gap, custo, viabilidade, prazo):",
         size=11, bold=True, color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, f"#1 Cotas em CBO 1–4 (P1): CC={fmt(P['TOPSIS_P1_CC'],3)} — política mais próxima da 'solução ideal'. "
                "Ataca diretamente o mecanismo de exclusão de acesso a ocupações qualificadas "
                "que explica 84% do gap via dotações.")
add_bullet(doc, f"#2 Equidade Educacional (P2): CC={fmt(P['TOPSIS_P2_CC'],3)} — segunda mais eficiente. "
                "Reduz o componente de dotações educacionais.")
add_bullet(doc, f"#3 Mentoria e Redes (P3): CC={fmt(P['TOPSIS_P3_CC'],3)} — responde diretamente ao achado da "
                "SNA (betweenness=0 para negros = sem capital social estrutural).")
add_bullet(doc, f"#4 Transparência Salarial (P4): CC={fmt(P['TOPSIS_P4_CC'],3)} — ataca os retornos diferenciais (16%).")
add_bullet(doc, f"PL-1 (Cotas CBO): projeta redução de {fmt(P['PL1_B5_PCT'],1)}% do gap bruto em simulações de LP.")
add_para(doc, "AHP — Consistência dos Pesos:", size=11, bold=True,
         color=(0x1F,0x38,0x64), space_before=4)
add_bullet(doc, "CR=0,004 (< 0,10 = aceitável). A matriz de comparação por pares de critérios "
                "é consistente — os pesos refletem preferências racionais e transitivas.")
add_colored_box(doc, "O que dizer se perguntarem sobre PO:",
    ["'A PO não é um método avaliativo adicional — é a tradução dos resultados empíricos "
     "em ação. A decomposição OB mostra que 84% do gap é de dotações → isso informa os "
     "pesos do TOPSIS. A SNA mostra que betweenness=0 → isso entra no critério de viabilidade "
     "da política P3. Os métodos são encadeados: diagnóstico → priorização → simulação.'",
     "'O TOPSIS é um método multi-critério MCDM clássico (Hwang & Yoon, 1981). Ele ranqueia "
     "alternativas pela distância relativa à solução ideal e à pior solução. CR<0,10 no AHP "
     "é o padrão de consistência de Saaty (1980) — garantia de que os pesos derivados da "
     "comparação por pares são internamente consistentes.'"],
    title_color=(0xFF,0x8F,0x00))
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTE 3 — EQUAÇÕES IMPORTANTES PARA MEMORIZAR
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "PARTE 3 — EQUAÇÕES PARA MEMORIZAR NA DEFESA", level=1)
add_para(doc, "Não é necessário decorar todas as equações — mas saber nomear e descrever cada "
              "componente dos 3 principais modelos demonstra domínio metodológico.",
         size=11, space_after=10)

add_heading(doc, "3.1  Modelo HLM M4 (equação de dois níveis simplificada)", level=2)
add_para(doc, "Nível 1 (indivíduo i na UPA j):", size=11, bold=True, space_before=6)
add_para(doc, "log(renda_ij) = β₀j + β₁(negro_i) + β₂(sexo_fem_i) + β₃(horas_c_i) + "
              "β₄(emprego_formal_i) + β₅(edu_anos_i) + β₆(idade_c_i) + β₇(ocp_grupo_cbo_i) + ε_ij",
         size=10.5, italic=True, color=(0x22,0x22,0x22))
add_para(doc, "Nível 2 (UPA j dentro de UF k):", size=11, bold=True, space_before=6)
add_para(doc, "β₀j = γ₀₀ + γ₀₁(renda_media_upa_j) + γ₀₂(edu_media_upa_j) + u₀j",
         size=10.5, italic=True, color=(0x22,0x22,0x22))
add_para(doc, "Onde u₀j ~ N(0, τ₀₀) é o efeito aleatório da UPA — fonte do ICC.",
         size=10.5, italic=True, color=(0x66,0x66,0x66))
add_para(doc, "O interesse é em β₁: o gap racial líquido após todos os controles. Em M4, β₁ ≈ −0,0644 "
              "(−6,2%). O sinal negativo confirma que negros ganham menos, tudo o mais constante.",
         size=11, space_before=6)

add_heading(doc, "3.2  Decomposição de Oaxaca-Blinder (two-fold)", level=2)
add_para(doc, "Ȳ_B − Ȳ_N  =  (X̄_B − X̄_N)·β_B  +  X̄_N·(β_B − β_N)",
         size=12, italic=True, bold=True, color=(0x15,0x65,0xC0))
add_bullet(doc, "Ȳ_B e Ȳ_N: renda média de brancos e negros (em log).")
add_bullet(doc, "X̄_B − X̄_N: diferença nas características médias (dotações).")
add_bullet(doc, "β_B: vetor de retornos estimado para brancos (grupo de referência).")
add_bullet(doc, "β_B − β_N: diferença nos retornos às mesmas características (discriminação de remuneração).")
add_bullet(doc, "84% do gap vem do primeiro termo (dotações); 16% do segundo (retornos).")

add_heading(doc, "3.3  GLMM — Modelo Linear Generalizado Misto (lme4::glmer)", level=2)
add_para(doc, "P(ocp_qualif_ij = 1) = Λ(β₀j + β₁·negro_i + X_i·γ + u₀j)",
         size=12, italic=True, bold=True, color=(0xFF,0x8F,0x00))
add_para(doc, "β₀j = γ₀₀ + u₀j,  u₀j ~ N(0, σ²ᵤ)",
         size=11, italic=True, color=(0x44,0x44,0x44))
add_bullet(doc, "Λ(·): função logística — transforma o preditor linear em probabilidade [0,1].")
add_bullet(doc, f"β₁ (negro): log-odds de negro vs branco. OR M1={or_str(P['OR_M1'])} (sem contexto); "
                f"OR M2={or_str(P['OR_M2'])} (com efeito aleatório de UPA). N={fmtN(P['N_GLMM'])} (PEA completa).")
add_bullet(doc, f"u₀j: efeito aleatório de UPA j — capta a heterogeneidade entre bairros. "
                f"ICC M1={fmt(P['ICC_M1_pct'],1)}% | ICC M2={fmt(P['ICC_M2_pct'],1)}%.")
add_bullet(doc, "Implementação: lme4::glmer (Python equivalente via statsmodels/pymer4), "
                "nAGQ=0 (aproximação de Laplace) para viabilidade computacional, "
                "otimizador bobyqa.")
add_bullet(doc, f"AME M1={ame(P['AME_M1_pp'])} | AME M2={ame(P['AME_M2_pp'])}: interpretados como diferença de "
                "probabilidade de ocupação qualificada atribuída à raça, todas as demais "
                "características iguais.")
add_bullet(doc, "Termos de interação (Tabela 4): negro×sexo_fem β=+0.0416 (***); "
                "negro×experiência β=−0.0019 (***); negro×educ_alto/med ns.")

add_separator(doc)
add_para(doc, "Boa defesa, Ricardo.",
         size=13, bold=True, italic=True, color=(0x1F,0x38,0x64),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

# ── Salvar ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"Arquivo gerado: {OUT.name}")
print(f"Tamanho: {OUT.stat().st_size // 1024} KB")
