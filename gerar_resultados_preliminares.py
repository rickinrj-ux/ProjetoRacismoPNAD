"""
gerar_resultados_preliminares.py
================================
Monta o documento de RESULTADOS PRELIMINARES no modelo MBA USP/Esalq
(Template Resultados Preliminares_PT), puxando todos os números de params.py
(fonte única de verdade). Estrutura: Título, Autores, Resumo, Palavras-chave,
Introdução, Material e Métodos, Resultados Preliminares, Considerações,
Referências. Formatação: Times New Roman 12, espaçamento 1,5, justificado,
títulos de seção em negrito e alinhados à esquerda. Sem os textos de instrução.

Saída: Resultados_Preliminares_TCC.docx
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from params import P, fmt, fmtN, or_str, ame

sys.stdout.reconfigure(encoding="utf-8")
ROOT = Path(__file__).parent
FIG  = ROOT / "outputs" / "figures"
OUT  = ROOT / "Resultados_Preliminares_TCC.docx"

def g(k, d=0.0): return P.get(k, d)
def pa(v, dec=1): return fmt(abs(v), dec)   # |valor| em pt-BR

doc = Document()
# Defaults de estilo (USP/Esalq)
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
st.paragraph_format.space_after = Pt(6)
for s in doc.sections:
    s.top_margin = Cm(3); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)


def titulo(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"

def centro(t, size=12, italic=False, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size); r.font.name = "Times New Roman"

def secao(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

def sub(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(t); r.bold = True; r.italic = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

def par(t):
    p = doc.add_paragraph(t); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def kv(rotulo, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(rotulo); r.bold = True; r.font.name = "Times New Roman"
    r2 = p.add_run(txt); r2.font.name = "Times New Roman"

def figura(nome, legenda, w=15):
    path = FIG / nome
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Cm(w))
    else:
        p.add_run(f"[figura: {nome}]").italic = True
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(legenda); rc.font.size = Pt(10); rc.font.name = "Times New Roman"
    c.paragraph_format.space_after = Pt(10)


# ── Cabeçalho ─────────────────────────────────────────────────────────────────
titulo("A desigualdade racial no mercado de trabalho brasileiro é estrutural, "
       "contextual e geograficamente heterogênea")
doc.add_paragraph()
centro("Ricardo Calheiros¹*; Edilson José Rodrigues²", size=12)
centro("¹* MBA em Data Science e Analytics, USP/Esalq. E-mail: rickinrj@gmail.com", size=10)
centro("² Orientador. MBA USP/Esalq.", size=10)
doc.add_paragraph()

# ── Resumo ────────────────────────────────────────────────────────────────────
secao("Resumo")
par(f"A desigualdade racial no mercado de trabalho brasileiro foi investigada com a PNAD "
    f"Contínua (2016–2025; {fmtN(int(g('N_GLMM',7694198)))} observações da população "
    f"economicamente ativa), integrando econometria multinível, aprendizado de máquina, "
    f"análise de redes e pesquisa operacional. Os resultados preliminares indicam que a "
    f"discriminação opera em camadas: uma barreira de acesso a ocupações qualificadas "
    f"(GLMM logístico, OR={or_str(g('OR_M2',0.69))}), uma penalidade salarial residual e um "
    f"teto de vidro que se agrava no topo da distribuição (OR de acesso ao decil superior "
    f"= {or_str(g('OR_TOP10_M1',0.45))}). Mais de metade do diferencial salarial é mediada "
    f"pelo contexto de moradia (UPA), e modelos de inclinação aleatória mostram que tanto a "
    f"penalidade de acesso quanto a salarial variam significativamente entre os estados. A "
    f"pesquisa operacional traduz o diagnóstico em priorização de políticas, com ganho de "
    f"focalização territorial.")
kv("Palavras-chave: ", "gap salarial racial; modelos hierárquicos; teto de vidro; "
   "segregação residencial; pesquisa operacional.")

# ── Introdução ────────────────────────────────────────────────────────────────
secao("Introdução")
par("A persistência da desigualdade racial no mercado de trabalho brasileiro é um dos "
    "fenômenos mais documentados das ciências sociais nacionais, desde a tese do "
    "preconceito estrutural de Hasenbalg (1979) até as decomposições contemporâneas do "
    "rendimento (Henriques, 2001; Soares, 2009). Apesar do avanço educacional das últimas "
    "décadas, trabalhadores negros seguem auferindo rendimentos inferiores e ocupando, em "
    "menor proporção, as posições de maior prestígio e remuneração — um padrão que a "
    "explicação meritocrática, baseada apenas em capital humano, não dá conta de prever.")
par("A literatura internacional sugere que a discriminação não é um evento único na "
    "contratação, mas um sistema de barreiras que se reforçam: exclusão de acesso a "
    "ocupações qualificadas (Pager, 2007), efeitos de vizinhança e segregação residencial "
    "(Wilson, 1987; Sampson, 1997) e exclusão das redes sociais que convertem credenciais "
    "em oportunidades (Granovetter, 1973; Burt, 2004). Tratados isoladamente, esses "
    "mecanismos têm sido difíceis de quantificar de forma integrada e em escala nacional.")
par("Este trabalho enfrenta essa lacuna combinando, sobre a série completa da PNAD "
    "Contínua, modelos lineares hierárquicos (HLM), modelos logísticos multiníveis (GLMM), "
    "decomposição de Oaxaca-Blinder, regressão quantílica, aprendizado de máquina com "
    "valores SHAP, análise de redes sociais (SNA) e pesquisa operacional. A convergência "
    "de métodos independentes sobre o mesmo conjunto de dados permite distinguir os "
    "mecanismos da desigualdade e testar sua robustez.")
par("O objetivo deste trabalho é identificar, isolar e quantificar os mecanismos da "
    "desigualdade racial de rendimentos e de acesso ocupacional no Brasil, avaliando sua "
    "heterogeneidade geográfica e setorial e traduzindo o diagnóstico em recomendações de "
    "política priorizadas por métodos de pesquisa operacional.")

# ── Material e Métodos ────────────────────────────────────────────────────────
secao("Material e Métodos")
sub("Base de dados")
par(f"Utilizou-se a Pesquisa Nacional por Amostra de Domicílios Contínua (PNAD Contínua/"
    f"IBGE), série completa de 2016 a 2025, totalizando {fmtN(int(g('N_GLMM',7694198)))} "
    f"observações da população economicamente ativa com rendimento positivo após o "
    f"tratamento dos dados. As variáveis incluíram raça (preto/pardo agregados em 'negro'), "
    f"gênero, idade, escolaridade, jornada, vínculo, grupo ocupacional (CBO) e indicadores "
    f"de contexto agregados por unidade primária de amostragem (UPA, proxy de bairro) e por "
    f"unidade da federação (UF). Os dados são de acesso público; nenhum indivíduo é "
    f"identificável.")
sub("Modelos multiníveis: HLM e GLMM")
par("Estimaram-se modelos lineares hierárquicos (HLM) para o logaritmo do rendimento, com "
    "efeitos aleatórios de UPA e UF, decompondo o gap bruto em mediação contextual e "
    "penalidade residual (Raudenbush; Bryk, 2002). O acesso a ocupações qualificadas e ao "
    "topo da distribuição de renda foi modelado por GLMM logístico (lme4::glmer, R), com "
    "efeito aleatório de UPA. Modelos de inclinação aleatória (random slope) de raça e de "
    "gênero por UF testaram a heterogeneidade geográfica do efeito (teste de fronteira de "
    "Stram e Lee, 1994), estimados também separadamente para os setores público e privado.")
sub("Decomposição e regressão quantílica")
par("A decomposição de Oaxaca-Blinder (Oaxaca, 1973; Blinder, 1973) separou o gap em "
    "componentes de dotações e de retornos. A regressão quantílica (Koenker; Bassett, 1978) "
    "avaliou a trajetória do gap ao longo da distribuição de rendimentos, caracterizando o "
    "teto de vidro.")
sub("Aprendizado de máquina e interpretabilidade")
par("Random Forest (Breiman, 2001) e XGBoost (Chen; Guestrin, 2016) foram ajustados para "
    "prever o rendimento, com interpretação por valores SHAP (Lundberg; Lee, 2017), "
    "quantificando a contribuição de cada variável e suas interações de forma "
    "não-paramétrica — uma camada de triangulação e robustez frente aos modelos "
    "econométricos.")
sub("Redes sociais e pesquisa operacional")
par("A análise de redes sociais (SNA) modelou a co-presença ocupacional como grafo, "
    "medindo isolamento estrutural (constraint de Burt) e capacidade de corretagem "
    "(betweenness) por grupo. Por fim, a pesquisa operacional (TOPSIS, AHP e programação "
    "linear; Saaty, 1980; Hwang; Yoon, 1981) traduziu os coeficientes confirmados em "
    "priorização e alocação ótima de políticas, inclusive em versão regionalizada por UF.")

# ── Resultados Preliminares ───────────────────────────────────────────────────
secao("Resultados Preliminares")
par("Os resultados preliminares sustentam uma tese central: a desigualdade racial no mercado "
    "de trabalho brasileiro não é um evento pontual de contratação, mas um SISTEMA DE BARREIRAS "
    "EM CAMADAS — acesso, remuneração e redes — ancorado no território e geograficamente "
    "heterogêneo, que persiste mesmo quando se controla a escolaridade e mesmo dentro do Estado. "
    "As subseções a seguir percorrem esse sistema: partem da falha da explicação meritocrática, "
    "isolam as três camadas que a substituem, demonstram seu caráter sistêmico e, por fim, "
    "traduzem o diagnóstico em prescrição.")

sub("A falha da explicação meritocrática")
par(f"Se o rendimento fosse função apenas do capital humano, controlar escolaridade, "
    f"experiência e jornada deveria dissolver o gap racial — o que não ocorre. A decomposição "
    f"de Oaxaca-Blinder atribui cerca de {pa(g('DOT_PCT',83.5),0)}% do diferencial a diferenças "
    f"de dotações e apenas {pa(g('RET_PCT',16.5),0)}% a retornos diferenciais. O ponto decisivo, "
    f"porém, está na COMPOSIÇÃO dessas dotações (Figura 1): os fatores que mais explicam o gap "
    f"não são educacionais, mas CONTEXTUAIS e OCUPACIONAIS — a proporção de negros na UPA e o "
    f"acesso a grupos ocupacionais de prestígio. As dotações não são, portanto, mérito neutro: "
    f"são o produto das barreiras que as subseções seguintes isolam. O capital humano explica "
    f"pouco; o que importa é onde se nasce e a que ocupações se tem acesso.")
figura("oaxaca_por_variavel.png",
       "Figura 1. Decomposição de Oaxaca-Blinder por variável: o gap vem sobretudo do contexto "
       "(proporção de negros na UPA) e da ocupação — não da escolaridade.", w=15)

sub("Camada 1 — A barreira de acesso e o teto de vidro")
par(f"A primeira camada é a exclusão da PORTA DE ENTRADA. O GLMM logístico multinível estima "
    f"que um trabalhador negro idêntico a um branco tem cerca de "
    f"{fmt(round((1-g('OR_M2',0.69))*100),0)}% menos chance de ocupar um cargo qualificado "
    f"(OR={or_str(g('OR_M2',0.69))}; {ame(g('AME_M2_pp',-4.84))}). A barreira não é uniforme ao "
    f"longo da hierarquia: ela se agrava no topo (Figura 2) — o odds ratio de acesso cai de "
    f"{or_str(g('OR_TOP20_M1',0.536))} no quintil superior de renda para "
    f"{or_str(g('OR_TOP10_M1',0.4533))} no decil superior. É um teto de vidro de ACESSO: quanto "
    f"mais valiosa a posição, mais opaco o filtro racial. O elevado ICC de UPA "
    f"({pa(g('ICC_M1_pct',22.2),1)}%) já antecipa que essa barreira tem raiz territorial — o fio "
    f"que conecta as camadas.")
figura("glmm_glassceil_forest.png",
       "Figura 2. GLMM — odds ratios de acesso por desfecho: o gradiente decrescente rumo ao "
       "topo da renda caracteriza o teto de vidro de acesso (OR < 1 = barreira).", w=14)

sub("Camada 2 — O território como eixo da desigualdade")
par("Por que as dotações são desiguais? A segunda camada responde: pelo LUGAR. No modelo "
    "hierárquico, mais da metade do gap salarial é mediada pelo contexto de moradia (UPA), e o "
    "aprendizado de máquina confirma-o de forma independente e não-paramétrica: entre todos os "
    "preditores do rendimento, a RENDA MÉDIA DO BAIRRO é o mais importante (Figura 3) — acima da "
    "escolaridade e de qualquer atributo individual. A raça, isoladamente, tem peso preditivo "
    "direto modesto justamente porque opera ATRAVÉS do território: a segregação residencial "
    "histórica converte-se em segregação de renda atual. Em uma frase: o bairro prediz mais que "
    "o diploma. É esse eixo territorial que dá unidade às demais camadas.")
figura("shap_importance_xgb.png",
       "Figura 3. Importância SHAP (XGBoost, R²≈0,62): a renda média da UPA (bairro) é o "
       "principal determinante do rendimento, evidenciando o eixo territorial da desigualdade.", w=14)

sub("Camada 3 — As redes que excluem")
par("A terceira camada explica por que superar as duas primeiras ainda não basta. A análise de "
    "redes sociais mostra que os grupos negros — inclusive os de alta escolaridade — ocupam "
    "posições de maior isolamento estrutural (constraint de Burt) e betweenness próximo de zero: "
    "ficam fora das redes de indicação que convertem credenciais em oportunidades (Figura 4). "
    "Assim, o retorno do diploma negro é menor não por mérito, mas pela ausência do capital "
    "social que, historicamente, é majoritariamente branco (Granovetter, 1973; Burt, 2004). As "
    "três camadas — acesso, território e redes — reforçam-se mutuamente.")
figura("sna_constraint_vs_renda.png",
       "Figura 4. Análise de redes: isolamento estrutural (constraint de Burt) × renda por grupo "
       "(raça × educação) — grupos negros combinam maior isolamento e menor renda.", w=13)

sub("O caráter sistêmico: heterogeneidade geográfica, Estado e gênero")
par(f"Três achados confirmam que se trata de um sistema, e não de fatores avulsos. Primeiro, a "
    f"discriminação é GEOGRAFICAMENTE HETEROGÊNEA: modelos de inclinação aleatória (real, via "
    f"lme4 e MixedLM) rejeitam o efeito homogêneo entre estados (p<0,001), tanto no acesso quanto "
    f"no salário — o gap salarial racial varia entre UFs de {pa(g('RS_GAP_LO_PCT',-19.8),1)}% a "
    f"{pa(g('RS_GAP_HI_PCT',-2.2),1)}% (Figura 5). Segundo, o ESTADO não dissolve a barreira: o "
    f"setor público atenua o gap SALARIAL racial ({pa(g('HRS_PUB_GAP_PCT',-7.8),1)}% vs "
    f"{pa(g('HRS_PRIV_GAP_PCT',-11.2),1)}% no privado) e o homogeneíza entre UFs, mas NÃO atenua a "
    f"barreira de ACESSO (OR={or_str(g('GRS_PUB_OR',0.705))} vs {or_str(g('GRS_PRIV_OR',0.695))}) "
    f"— o concurso equaliza a remuneração de quem entra, não a entrada. Terceiro, o GÊNERO segue "
    f"padrão próprio e complementar: mulheres acessam mais ocupações qualificadas "
    f"(OR={fmt(g('GGE_OCP_OR',2.01),2)}) mas muito menos o topo de renda "
    f"(OR={or_str(g('GGE_TOP10_OR',0.522))}) — teto de vidro de remuneração, com heterogeneidade "
    f"geográfica ainda maior que a racial.")
figura("mapa_po_regional.png",
       "Figura 5. Penalidade salarial racial por estado (BLUP do random slope): mais escuro = "
       "maior desvantagem; estrelas marcam as UFs prioritárias para a focalização da política.", w=11)

sub("Hipóteses complementares: o papel ambíguo do Estado")
par("Um conjunto de hipóteses adicionais testa diretamente a hipótese alternativa de que o "
    "Estado promoveria — ou, ao menos, não dissolveria — a desigualdade racial. Os resultados "
    "são ambíguos e refinam a tese central. Como indutor de igualdade GERAL, o Estado funciona: "
    "o setor público reduz levemente a desigualdade de renda (índice de Gini inferior ao do "
    "privado; H1). Mas como dissolvente da barreira RACIAL, não: embora o público atenue o gap "
    "racial (H2, Figura 6) — coerente com a camada sistêmica anterior —, ele PIORA o gap de "
    "gênero (−20,9% no público contra −18,6% no privado; H3): o concurso equaliza a entrada, mas "
    "a promoção às posições de liderança reflete, ou amplifica, os vieses do setor privado. No "
    "agregado (H4), o país combina pleno emprego histórico com renda real estagnada — produz "
    "ocupação, não prosperidade igualitária. E a armadilha da qualificação persiste (H5): a "
    "proporção de trabalhadores negros em ocupações de alta qualificação (≈11%) é cerca de "
    "metade da de brancos (≈22%), gap estável na década. Em síntese, o Estado mitiga a "
    "desigualdade GERAL, mas não a RACIAL — sustentando parcialmente a hipótese alternativa e "
    "reforçando o caráter sistêmico do problema.")
figura("estado_h2h3_gaps.png",
       "Figura 6. Hipóteses do Estado (H2/H3): gap racial e de gênero por setor (gaps "
       "controlados) — o público atenua o racial, mas agrava o de gênero.", w=15)

sub("Da diagnose à prescrição: pesquisa operacional")
par(f"Se a desigualdade é um sistema em camadas e territorialmente desigual, a política eficaz "
    f"precisa ser multidimensional e focalizada — exatamente o que a pesquisa operacional "
    f"formaliza. O ranqueamento multicritério (TOPSIS) aponta as cotas de acesso a ocupações "
    f"qualificadas como intervenção dominante (CC={fmt(g('TOPSIS_P1_CC',0.83),3)}; Figura 7), "
    f"coerente com o fato de a maior barreira ser de acesso. A versão regionalizada mostra, "
    f"ainda, que concentrar o orçamento nas UFs de maior penalidade (lideradas por "
    f"{g('RPO_WORST_UF','DF')}, {pa(g('RPO_WORST_GAP_PCT',-20.9),1)}%) rende até "
    f"{pa(g('RPO_GANHO_B9',68.2),0)}% mais redução do gap do que a distribuição uniforme — "
    f"fechando o ciclo entre o diagnóstico das camadas e a alocação ótima de recursos.")
figura("po_politicas_topsis.png",
       "Figura 7. Pesquisa operacional (TOPSIS): ranking multicritério das seis políticas; cotas "
       "de acesso a ocupações qualificadas lideram, em linha com a barreira de acesso.", w=15)

# ── Considerações preliminares ────────────────────────────────────────────────
secao("Considerações Preliminares")
par("Tomados em conjunto, os resultados parciais convergem para a tese central: a desigualdade "
    "racial no trabalho brasileiro é um SISTEMA DE BARREIRAS EM CAMADAS — acesso, remuneração e "
    "redes — que se reforçam e têm como eixo comum o TERRITÓRIO (o bairro prediz mais que o "
    "diploma). Esse sistema persiste sob controle de escolaridade, agrava-se no topo da "
    "distribuição, varia entre estados e não é dissolvido pelo Estado, que equaliza a "
    "remuneração de quem entra, mas não o acesso. A consequência prescritiva é direta: políticas "
    "unidimensionais e uniformes são insuficientes; a pesquisa operacional indica intervenção "
    "multidimensional, centrada no acesso e focalizada nas UFs de maior penalidade. As próximas "
    "etapas incluem o refinamento das análises de robustez e a consolidação das recomendações.")

# ── Referências ───────────────────────────────────────────────────────────────
secao("Referências")
refs = [
    "BLINDER, A. S. Wage discrimination: reduced form and structural estimates. Journal of Human Resources, v. 8, n. 4, p. 436-455, 1973.",
    "BREIMAN, L. Random forests. Machine Learning, v. 45, n. 1, p. 5-32, 2001.",
    "BURT, R. S. Structural holes and good ideas. American Journal of Sociology, v. 110, n. 2, p. 349-399, 2004.",
    "CHEN, T.; GUESTRIN, C. XGBoost: a scalable tree boosting system. In: KDD, 2016. p. 785-794.",
    "GRANOVETTER, M. The strength of weak ties. American Journal of Sociology, v. 78, n. 6, p. 1360-1380, 1973.",
    "HASENBALG, C. Discriminação e desigualdades raciais no Brasil. Rio de Janeiro: Graal, 1979.",
    "HENRIQUES, R. Desigualdade racial no Brasil: evolução das condições de vida na década de 90. Rio de Janeiro: IPEA, 2001. (Texto para discussão, 807).",
    "HWANG, C. L.; YOON, K. Multiple attribute decision making: methods and applications. Berlin: Springer, 1981.",
    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Pesquisa Nacional por Amostra de Domicílios Contínua. Rio de Janeiro: IBGE, 2016-2025.",
    "KOENKER, R.; BASSETT, G. Regression quantiles. Econometrica, v. 46, n. 1, p. 33-50, 1978.",
    "LUNDBERG, S. M.; LEE, S. A unified approach to interpreting model predictions. In: NeurIPS, 2017. p. 4765-4774.",
    "OAXACA, R. Male-female wage differentials in urban labor markets. International Economic Review, v. 14, n. 3, p. 693-709, 1973.",
    "PAGER, D. Marked: race, crime, and finding work in an era of mass incarceration. Chicago: University of Chicago Press, 2007.",
    "RAUDENBUSH, S. W.; BRYK, A. S. Hierarchical linear models: applications and data analysis methods. 2. ed. Thousand Oaks: Sage, 2002.",
    "SAATY, T. L. The analytic hierarchy process. New York: McGraw-Hill, 1980.",
    "SAMPSON, R. J.; RAUDENBUSH, S. W.; EARLS, F. Neighborhoods and violent crime. Science, v. 277, p. 918-924, 1997.",
    "SOARES, S. S. D. Perfil da discriminação no mercado de trabalho: raça, sexo e salários no Brasil 1992-2006. Rio de Janeiro: IPEA, 2009. (Texto para discussão, 1395).",
    "STRAM, D. O.; LEE, J. W. Variance components testing in the longitudinal mixed effects model. Biometrics, v. 50, n. 4, p. 1171-1177, 1994.",
    "WILSON, W. J. The truly disadvantaged: the inner city, the underclass, and public policy. Chicago: University of Chicago Press, 1987.",
]
for r in sorted(refs):
    p = doc.add_paragraph(r); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs: run.font.size = Pt(11)

doc.save(str(OUT))
print(f"Arquivo gerado: {OUT.name}")
print(f"Tamanho: {OUT.stat().st_size // 1024} KB")
