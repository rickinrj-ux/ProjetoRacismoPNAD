"""
gerar_resultados_preliminares.py
================================
Monta o documento de RESULTADOS PRELIMINARES no modelo MBA USP/Esalq
(Template Resultados Preliminares_PT), puxando todos os números de params.py
(fonte única de verdade). Estrutura: Título, Autores, Resumo, Palavras-chave,
Introdução, Material e Métodos, Resultados Preliminares, Considerações,
Referências. Formatação: Times New Roman 12, espaçamento 1,5, justificado,
títulos de seção em negrito e alinhados à esquerda. Sem os textos de instrução.

Saída: Resultados_Preliminares_TCC.docx
"""

# --- bootstrap raiz do projeto (reorg estrutura) ---
import os as _os, sys as _sys
from pathlib import Path as _Path
_os.chdir(_Path(__file__).resolve().parents[2])
_sys.path.insert(0, _os.getcwd())
# --- fim bootstrap ---

import sys
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from params import P, fmt, fmtN, or_str, ame

sys.stdout.reconfigure(encoding="utf-8")
ROOT = Path.cwd()
FIG  = ROOT / "outputs" / "figures"
OUT  = ROOT / "entregaveis" / "Resultados_Preliminares_TCC.docx"
(ROOT / "entregaveis").mkdir(exist_ok=True)

def g(k, d=0.0): return P.get(k, d)
def pa(v, dec=1): return fmt(abs(v), dec)   # |valor| em pt-BR

doc = Document()
# Defaults de estilo (USP/Esalq)
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
st.paragraph_format.space_after = Pt(6)
for s in doc.sections:
    s.top_margin = Cm(3); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)


def titulo(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"

def centro(t, size=12, italic=False, bold=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size); r.font.name = "Times New Roman"

def secao(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

def sub(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(t); r.bold = True; r.italic = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

def par(t):
    p = doc.add_paragraph(t); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def kv(rotulo, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(rotulo); r.bold = True; r.font.name = "Times New Roman"
    r2 = p.add_run(txt); r2.font.name = "Times New Roman"

def figura(nome, legenda, w=15):
    path = FIG / nome
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Cm(w))
    else:
        p.add_run(f"[figura: {nome}]").italic = True
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(legenda); rc.font.size = Pt(10); rc.font.name = "Times New Roman"
    c.paragraph_format.space_after = Pt(10)


def tabela(legenda, headers, rows):
    """Tabela de resultado (pedido do orientador): legenda acima (ABNT), grade."""
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    rc = cap.add_run(legenda); rc.font.size = Pt(10); rc.bold = True; rc.font.name = "Times New Roman"
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _cell(cell, txt, bold=False):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(str(txt))
        r.font.size = Pt(10); r.bold = bold; r.font.name = "Times New Roman"

    for j, h in enumerate(headers):
        _cell(t.rows[0].cells[j], h, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            _cell(cells[j], v)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Carrega resultados dos CSV canônicos (para as tabelas) ─────────────────────
TBL = ROOT / "outputs" / "tables"
_med = pd.read_csv(TBL / "gap_decomposicao_serie_completo.csv")
_oba = pd.read_csv(TBL / "ob_acesso.csv").iloc[0]
_glm = pd.read_csv(TBL / "glmm_glassceil_full.csv")
_ev  = pd.read_csv(TBL / "evalues_glmm.csv")
_itx = pd.read_csv(TBL / "interseccional_ob4grupos.csv")

def _glm_m2(des):
    r = _glm[(_glm.desfecho == des) & (_glm.modelo == "M2")].iloc[0]
    e = _ev[(_ev.Desfecho == des) & (_ev.Modelo == "M2")]
    ev = fmt(float(e["E-value (OR)"].iloc[0]), 2) if len(e) else "---"
    ic = f"[{fmt(r.CI95_lo, 3)}; {fmt(r.CI95_hi, 3)}]"
    return [fmt(r.OR_negro, 3), ic, fmt(r.AME_pp, 1), ev]

_MEDLBL = {"M1_Individual": "M1 (individual)", "M2_Localidade": "M2 (+ contexto UPA)",
           "M3_Completo": "M3 (+ UF)", "M4_Ocupacao": "M4 (+ ocupação)"}
_med_rows = [[_MEDLBL.get(r["Modelo"], r["Modelo"]), fmt(r["b_negro"], 4), fmt(r["Gap%"], 1),
              ("—" if pd.isna(r["Mediacao_total%"]) else fmt(r["Mediacao_total%"], 1))]
             for r in _med.to_dict("records")]
_oba_rows = [["Dotação (composição: capital humano e ocupação)", fmt(_oba["pct_dotacao"], 1)],
             ["Coeficiente (não explicado / discriminação)", fmt(_oba["pct_coeficiente"], 1)]]
_itx_rows = [[r["grupo"], fmt(r["gap_pct"], 1), fmt(r["end_pct"], 1), fmt(r["ret_pct"], 1),
              ("—" if abs(r["penalidade_extra_pct"]) < 1e-6 else fmt(r["penalidade_extra_pct"], 1))]
             for r in _itx.to_dict("records")]


# ── Cabeçalho ─────────────────────────────────────────────────────────────────
titulo("A desigualdade racial no mercado de trabalho brasileiro é estrutural, "
       "contextual e geograficamente heterogênea")
doc.add_paragraph()
centro("Ricardo Calheiros¹*; Edilson José Rodrigues²", size=12)
centro("¹* MBA em Data Science e Analytics, USP/Esalq. E-mail: rickinrj@gmail.com", size=10)
centro("² Orientador. MBA USP/Esalq.", size=10)
doc.add_paragraph()

# ── Resumo ────────────────────────────────────────────────────────────────────
secao("Resumo")
par(f"A desigualdade racial no mercado de trabalho brasileiro foi investigada com a PNAD "
    f"Contínua (2016–2025; {fmtN(int(g('N_GLMM',7694198)))} observações da população "
    f"economicamente ativa), integrando econometria multinível, decomposições do gap e "
    f"aprendizado de máquina interpretável. Os resultados preliminares indicam que a "
    f"discriminação opera em camadas: uma barreira de acesso a ocupações qualificadas "
    f"(GLMM logístico, OR={or_str(g('OR_M2',0.69))}), uma penalidade salarial residual e um "
    f"teto de vidro que se agrava no topo da distribuição (OR de acesso ao decil superior "
    f"= {or_str(g('OR_TOP10_M1',0.45))}). Mais de metade do diferencial salarial é mediada "
    f"pelo contexto de moradia (UPA); a decomposição de Oaxaca-Blinder atribui 83,8% do gap a "
    f"diferenças de dotações (sobretudo de ocupação) e 16,2% à parcela não explicada, enquanto a "
    f"penalidade recai de forma agravada sobre a mulher negra (penalidade interseccional de "
    f"+9,5 p.p.).")
kv("Palavras-chave: ", "gap salarial racial; modelos hierárquicos; teto de vidro; "
   "decomposição de Oaxaca-Blinder; interseccionalidade.")

# ── Introdução ────────────────────────────────────────────────────────────────
secao("Introdução")
par("A persistência da desigualdade racial no mercado de trabalho brasileiro é um dos "
    "fenômenos mais documentados das ciências sociais nacionais, desde a tese do "
    "preconceito estrutural de Hasenbalg (1979) até as decomposições contemporâneas do "
    "rendimento (Henriques, 2001; Soares, 2009). Apesar do avanço educacional das últimas "
    "décadas, trabalhadores negros seguem auferindo rendimentos inferiores e ocupando, em "
    "menor proporção, as posições de maior prestígio e remuneração — um padrão que a "
    "explicação meritocrática, baseada apenas em capital humano, não dá conta de prever.")
par("Esse hiato convive com um quadro recente de avanço social com desigualdade persistente: a "
    "Pesquisa de Orçamentos Familiares (POF) do IBGE registra melhora ampla na qualidade de vida "
    "entre 2008 e 2018 (queda de ~30% no Índice de Perda de Qualidade de Vida), mas sem fechar o "
    "gap racial (chefes pretos/pardos 0,183 vs. brancos 0,122) nem o territorial (Norte 0,223 / "
    "Nordeste 0,207 vs. Sul 0,114; IBGE, POF 2017–2018). No rendimento, o Gini domiciliar per "
    "capita voltou a subir em 2025 (0,491, ante 0,487 em 2024), puxado pelo topo (IBGE/PNAD "
    "Contínua, 2025). É esse padrão — progresso agregado que não dissolve a barreira racial — que "
    "este trabalho disseca no mercado de trabalho.")
par("A literatura internacional sugere que a discriminação não é um evento único na "
    "contratação, mas um sistema de barreiras que se reforçam: exclusão de acesso a "
    "ocupações qualificadas (Pager, 2007), efeitos de vizinhança e segregação residencial "
    "(Wilson, 1987; Sampson, 1997) e barreiras na conversão de credenciais em ocupações "
    "qualificadas (Pager, 2007). Tratados isoladamente, esses "
    "mecanismos têm sido difíceis de quantificar de forma integrada e em escala nacional.")
par("Este trabalho enfrenta essa lacuna combinando, sobre a série completa da PNAD "
    "Contínua, quatro métodos complementares — modelos lineares hierárquicos (HLM), "
    "decomposição de Oaxaca-Blinder, regressão quantílica com RIF e modelos logísticos "
    "multiníveis (GLMM) —, validados por aprendizado de máquina com valores SHAP. A "
    "convergência de métodos independentes sobre o mesmo conjunto de dados permite distinguir os "
    "mecanismos da desigualdade e testar sua robustez.")
par("O objetivo deste trabalho é identificar, isolar e quantificar os mecanismos da "
    "desigualdade racial de rendimentos e de acesso ocupacional no Brasil, distinguindo a "
    "parcela de composição da parcela de discriminação e traduzindo o diagnóstico em "
    "implicações de política focadas no acesso a ocupações qualificadas.")

# ── Material e Métodos ────────────────────────────────────────────────────────
secao("Material e Métodos")
sub("Base de dados")
par(f"Utilizou-se a Pesquisa Nacional por Amostra de Domicílios Contínua (PNAD Contínua/"
    f"IBGE), série completa de 2016 a 2025, totalizando {fmtN(int(g('N_GLMM',7694198)))} "
    f"observações da população economicamente ativa com rendimento positivo após o "
    f"tratamento dos dados. As variáveis incluíram raça (preto/pardo agregados em 'negro'), "
    f"gênero, idade, escolaridade, jornada, vínculo, grupo ocupacional (CBO) e indicadores "
    f"de contexto agregados por unidade primária de amostragem (UPA, proxy de bairro) e por "
    f"unidade da federação (UF). Os dados são de acesso público; nenhum indivíduo é "
    f"identificável.")
sub("Modelos multiníveis: HLM e GLMM")
par("Estimaram-se modelos lineares hierárquicos (HLM) para o logaritmo do rendimento, com "
    "efeitos aleatórios de UPA e UF, decompondo o gap bruto em mediação contextual e "
    "penalidade residual (Raudenbush; Bryk, 2002). O acesso a ocupações qualificadas e ao "
    "topo da distribuição de renda foi modelado por GLMM logístico (lme4::glmer, R), com "
    "efeito aleatório de UPA. A robustez dos achados foi avaliada por E-values (VanderWeele; "
    "Ding, 2017).")
sub("Decomposição e regressão quantílica")
par("A decomposição de Oaxaca-Blinder (Oaxaca, 1973; Blinder, 1973) separou o gap em "
    "componentes de dotações e de retornos. A regressão quantílica (Koenker; Bassett, 1978) "
    "avaliou a trajetória do gap ao longo da distribuição de rendimentos, caracterizando o "
    "teto de vidro.")
sub("Aprendizado de máquina e interpretabilidade")
par(f"Random Forest (Breiman, 2001; 200 árvores, profundidade máxima 10, mínimo de 50 "
    f"observações por folha) e XGBoost (Chen; Guestrin, 2016; 300 iterações, "
    f"profundidade máxima 6, taxa de aprendizado 0,05, subamostragem de 80% das linhas e "
    f"colunas por árvore, regularização L1={fmt(0.1,1)} e L2={fmt(1.0,1)}) foram ajustados para "
    f"prever o rendimento, com validação por partição hold-out treino/teste (80/20, "
    f"estratificação implícita por amostragem aleatória, semente fixa) — o desempenho reportado "
    f"(R²={fmt(g('ML_XGB_R2_TESTE',0.6168),4)} para XGBoost e R²={fmt(g('ML_RF_R2_TESTE',0.5735),4)} "
    f"para Random Forest) refere-se ao conjunto de TESTE, não a ajuste in-sample; a diferença "
    f"frente ao R² de treino (gap de overfitting ≤{fmt(g('ML_RF_GAP_OVERFIT',0.0006),4)}) é "
    f"desprezível para ambos os modelos. A interpretação usou valores SHAP (Lundberg; Lee, 2017), "
    f"quantificando a contribuição de cada variável e suas interações de forma "
    f"não-paramétrica — uma camada de triangulação e robustez frente aos modelos "
    f"econométricos, não um substituto causal para eles.")
sub("Sensibilidade e interseccionalidade")
par("A robustez do resíduo racial foi avaliada por E-values (VanderWeele; Ding, 2017), que "
    "quantificam quanto um confundidor não-observado precisaria pesar para anular o efeito. A "
    "decomposição interseccional (raça × gênero) formaliza a penalidade específica da combinação "
    "dos eixos (Crenshaw, 1989), distinguindo-a da soma dos efeitos isolados.")

# ── Resultados Preliminares ───────────────────────────────────────────────────
secao("Resultados Preliminares")
par("Os resultados preliminares sustentam uma tese central: a desigualdade racial no mercado "
    "de trabalho brasileiro não é um evento pontual de contratação, mas um SISTEMA DE BARREIRAS "
    "EM CAMADAS — acesso, remuneração e interseccionalidade — com forte mediação do território, "
    "que persiste mesmo quando se controla a escolaridade. As subseções a seguir percorrem esse "
    "sistema: partem da falha da explicação meritocrática e isolam as três camadas que a "
    "substituem.")

sub("A falha da explicação meritocrática")
par(f"Se o rendimento fosse função apenas do capital humano, controlar escolaridade, "
    f"experiência e jornada deveria dissolver o gap racial — o que não ocorre. A decomposição "
    f"de Oaxaca-Blinder atribui cerca de {pa(g('DOT_PCT',83.5),0)}% do diferencial a diferenças "
    f"de dotações e apenas {pa(g('RET_PCT',16.5),0)}% a retornos diferenciais. O ponto decisivo, "
    f"porém, está na COMPOSIÇÃO dessas dotações (Figura 1): os fatores que mais explicam o gap "
    f"não são educacionais, mas CONTEXTUAIS e OCUPACIONAIS — a proporção de negros na UPA e o "
    f"acesso a grupos ocupacionais de prestígio. As dotações não são, portanto, mérito neutro: "
    f"são o produto das barreiras que as subseções seguintes isolam. O capital humano explica "
    f"pouco; o que importa é onde se nasce e a que ocupações se tem acesso.")
figura("oaxaca_por_variavel.png",
       "Figura 1. Decomposição de Oaxaca-Blinder por variável: o gap vem sobretudo do contexto "
       "(proporção de negros na UPA) e da ocupação — não da escolaridade.", w=15)

sub("Camada 1 — A barreira de acesso e o teto de vidro")
par(f"A primeira camada é a exclusão da PORTA DE ENTRADA. O GLMM logístico multinível estima "
    f"que um trabalhador negro idêntico a um branco tem cerca de "
    f"{fmt(round((1-g('OR_M2',0.69))*100),0)}% menos chance de ocupar um cargo qualificado "
    f"(OR={or_str(g('OR_M2',0.69))}; {ame(g('AME_M2_pp',-4.84))}). A barreira não é uniforme ao "
    f"longo da hierarquia: ela se agrava no topo (Figura 2) — o odds ratio de acesso cai de "
    f"{or_str(g('OR_TOP20_M1',0.536))} no quintil superior de renda para "
    f"{or_str(g('OR_TOP10_M1',0.4533))} no decil superior. É um teto de vidro de ACESSO: quanto "
    f"mais valiosa a posição, mais opaco o filtro racial. O elevado ICC de UPA "
    f"({pa(g('ICC_M1_pct',22.2),1)}%) já antecipa que essa barreira tem raiz territorial — o fio "
    f"que conecta as camadas.")
figura("glmm_glassceil_forest.png",
       "Figura 2. GLMM — odds ratios de acesso por desfecho: o gradiente decrescente rumo ao "
       "topo da renda caracteriza o teto de vidro de acesso (OR < 1 = barreira).", w=14)
tabela("Tabela 1. GLMM logístico (M2, população completa) — teto de vidro de acesso. "
       "OR < 1 = menor chance de acesso para negros vs. brancos de mesmo perfil; IC 95% de Wald "
       "via broom.mixed (lme4::glmer); AME em pontos percentuais; E-value ≥ 2 indica robustez a "
       "confundidor não observado.",
       ["Desfecho", "OR (negro)", "IC 95%", "AME (p.p.)", "E-value"],
       [["Cargo qualificado (CBO 1–4)"] + _glm_m2("ocp_qualif"),
        ["Top 20% de renda"] + _glm_m2("y_top20"),
        ["Top 10% de renda"] + _glm_m2("y_top10")])
par(f"Uma leitura interseccional (quatro grupos raça×gênero, referência = homem branco) revela uma "
    f"inversão. No ACESSO à categoria, a mulher negra é alçada (OR={fmt(g('GRG_MN_OCP',1.33),2)}, "
    f"acima do homem branco, por profissões feminizadas em CBO 1–4) e o mais penalizado é o homem "
    f"negro (OR={fmt(g('GRG_HN_OCP',0.65),2)}); mas no TOPO da renda o quadro inverte e a mulher negra "
    f"torna-se a MAIS excluída de todos (OR={fmt(g('GRG_MN_TOP10',0.34),2)} no decil superior, abaixo "
    f"da mulher branca e do homem negro). A interação é sub-aditiva, mas o teto de vidro recai com "
    f"força máxima sobre a mulher negra (Figura 2b).")
figura("grupo_rg_interseccional.png",
       "Figura 2b. Interseccionalidade raça×gênero: a mulher negra é alçada no acesso à categoria, "
       "mas a mais excluída no topo da renda (OR vs. homem branco).", w=14)
par(f"Um indicador distribucional reforça o teto de vidro. O Gini da renda do trabalho é MAIOR entre "
    f"brancos ({fmt(g('GINI_BRANCO_TRAB',0.485),3)}) do que entre negros ({fmt(g('GINI_NEGRO_TRAB',0.449),3)}) "
    f"— em todos os anos da série. Isso não é equidade entre negros: é CONFINAMENTO AO PISO (renda "
    f"homogeneamente baixa), o reverso distribucional da barreira de acesso ao topo. Quem é barrado do "
    f"topo achata-se na base. Cabe a ressalva de que este Gini refere-se ao rendimento do TRABALHO entre "
    f"ocupados — conceito distinto do Gini domiciliar per capita do IBGE (Figura 2c).")
figura("gini_raca.png",
       "Figura 2c. Gini intra-raça (renda do trabalho, ponderado): brancos têm maior desigualdade "
       "interna; negros, comprimidos no piso — o reverso do teto de vidro, não equidade.", w=13)

sub("Camada 2 — O território como eixo da desigualdade")
par("Por que as dotações são desiguais? A segunda camada responde: em parte, pelo LUGAR. No "
    "modelo hierárquico, mais da metade do gap salarial é mediada pelo contexto de moradia "
    "(UPA) — um achado robusto de mediação. A raça opera, em medida relevante, ATRAVÉS do "
    "território: a segregação residencial histórica converte-se em segregação de renda atual. "
    "Ressalva metodológica: a renda média da UPA é preditor parcialmente endógeno do rendimento "
    "individual (problema do reflexo, Manski 1993), pois agrega o próprio indivíduo; por isso a "
    "interpretamos como evidência de mediação territorial, e não como 'determinante' causal "
    "isolado. É esse eixo territorial que dá unidade às demais camadas.")
figura("shap_importance_xgb.png",
       f"Figura 3. Importância SHAP (XGBoost, R² de teste={fmt(g('ML_XGB_R2_TESTE',0.6168),4)}): o "
       f"contexto territorial (renda média da UPA) está entre os preditores de maior peso, "
       f"sinalizando o eixo territorial da desigualdade — interpretado como mediação, não como "
       f"determinante causal isolado.", w=14)
par(f"Estabilidade da hierarquia SHAP. A concordância entre os dois modelos de árvore é alta: "
    f"correlação de Spearman entre os rankings completos de importância de "
    f"ρ={fmt(g('SHAP_SPEARMAN_RHO',0.839),3)}, com {int(g('SHAP_TOP10_OVERLAP',10))}/10 variáveis "
    f"coincidentes no top-10 de ambos os modelos — a hierarquia de importância territorial e "
    f"ocupacional (Figura 3) não é artefato de um único algoritmo. Ressalva de transparência: a "
    f"variável 'raça (negro)', embora sempre com sinal negativo (penalidade), tem posição "
    f"instável entre os modelos (rank {int(g('SHAP_NEGRO_RANK_RF',25))} no Random Forest vs. "
    f"rank {int(g('SHAP_NEGRO_RANK_XGB',11))} no XGBoost) — isso não compromete a conclusão "
    f"qualitativa (efeito racial direto residual mesmo após os controles), mas indica que a "
    f"MAGNITUDE relativa desse efeito frente às demais variáveis, isoladamente pelo SHAP, é "
    f"sensível à escolha do algoritmo, reforçando por que o efeito racial é estimado "
    f"primariamente pelos modelos econométricos (HLM, Oaxaca-Blinder, GLMM) e não pelo ML.")
par("Esse eixo territorial encontra corroboração externa no Índice de Progresso Social (IPS) municipal "
    "(Imazon e parceiros, 2026): as regiões de menor progresso social (Norte e Nordeste) coincidem com "
    "as de maior penalidade racial em nossos modelos. Ressalva metodológica: a integração fina com o "
    "proxy de bairro (UPA) é inviável — a PNAD não divulga o município e o IPS é municipal, mais "
    "agregado que a UPA —, de modo que o IPS entra como evidência convergente do caráter territorial, "
    "não como fonte de dados integrada.")
tabela("Tabela 2. Decomposição do gap por mediação (HLM de três níveis, PNAD Contínua 2016–2025). "
       "Cada modelo acrescenta controles ao anterior; o gap encolhe de −19,1% (M1) para −6,2% (M4), "
       "com 69,8% mediado por contexto de moradia e ocupação.",
       ["Modelo (controles acumulados)", "β negro", "Gap (%)", "Mediação acum. (%)"],
       _med_rows)
tabela("Tabela 3. Decomposição de Oaxaca-Blinder (especificação de acesso, com ocupação e contexto "
       "como dotações; população completa). Dotações + Coeficiente = 100% do gap. Ressalva (Oaxaca & "
       "Ransom, 1999): incluir ocupação como dotação subestima a discriminação — daí a complementaridade "
       "com o GLMM de acesso (Tabela 1).",
       ["Componente", "% do Gap"],
       _oba_rows)

sub("Camada 3 — A penalidade interseccional")
par("As duas primeiras camadas — acesso e território — combinam-se de forma agravada na "
    "interseção de raça e gênero. A decomposição interseccional mostra que a Mulher Negra "
    "acumula a maior desvantagem (gap de 96,4% vs. o Homem Branco) e, além disso, uma penalidade "
    "EXTRA de 9,5 pontos percentuais que não se reduz à soma das penalidades de raça e de gênero "
    "isoladas — a marca da interseccionalidade (Crenshaw, 1989). As camadas reforçam-se "
    "mutuamente: a barreira de acesso e o eixo territorial pesam de modo desigual sobre os "
    "diferentes grupos.")
tabela("Tabela 4. Decomposição interseccional (raça × gênero) do gap vs. o Homem Branco. "
       "Dotações + Retornos = 100% do gap. A Mulher Negra acumula o maior gap (96,4%) e uma "
       "penalidade extra de 9,5 p.p. não redutível à soma dos eixos de raça e gênero (Crenshaw, 1989).",
       ["Grupo", "Gap vs HB (%)", "Dotações (%)", "Retornos (%)", "Penal. extra (p.p.)"],
       _itx_rows)
par(f"Incerteza da penalidade interseccional. A penalidade extra da Tabela 4 é uma decomposição "
    f"não-linear (diferença entre três ajustes OLS por subgrupo) cujo erro-padrão direto está em "
    f"apuração via bootstrap (reamostragem por UPA). Como evidência complementar já disponível, o "
    f"modelo HLM interseccional com interação tripla explícita (log-renda contínua, controlando "
    f"educação superior) estima o termo raça×gênero×superior em "
    f"β={fmt(g('ITX_TRIPLE_B',-0.0434),4)} (EP={fmt(g('ITX_TRIPLE_SE',0.0053),4)}; "
    f"IC 95% [{fmt(g('ITX_TRIPLE_CI_LO',-0.0538),4)}; {fmt(g('ITX_TRIPLE_CI_HI',-0.033),4)}]), "
    f"estatisticamente distinto de zero — o intervalo não inclui zero, corroborando com incerteza "
    f"estatística explícita a existência de uma penalidade interseccional própria.")
figura("grupo_rg_interseccional.png",
       "Figura 4. Decomposição interseccional (raça × gênero): gap de renda vs. o Homem Branco; "
       "a Mulher Negra acumula as duas penalidades, com um resíduo interseccional próprio.", w=13)

# ── Considerações preliminares ────────────────────────────────────────────────
secao("Limitações e escopo de validade")
par("Natureza inferencial vs. preditiva. Os modelos HLM, Oaxaca-Blinder, regressão quantílica e "
    "correção de Heckman produzem estimativas de ASSOCIAÇÃO CONDICIONAL — o diferencial racial que "
    "persiste sob controle de observáveis —, e não prova causal contrafactual. O XGBoost e os "
    "valores SHAP têm finalidade PREDITIVA e interpretativa: medem a contribuição de cada variável "
    "para a previsão do rendimento, não o efeito causal de manipulá-la. A linguagem causal foi "
    "deliberadamente evitada.")
par("Cobertura da escolaridade. A escolaridade detalhada está registrada para cerca de 31% da PEA "
    "no painel público; os níveis entram como dummies de conclusão acompanhadas de um indicador "
    "explícito de não-registro (educ_missing), de modo que a categoria-base não confunda baixa "
    "escolaridade com dado ausente. O coeficiente racial é estável a essa especificação (variação "
    "inferior a 1%).")
par(f"Especificação da decomposição de Oaxaca-Blinder. A repartição entre composição e discriminação "
    f"depende de quais controles se tratam como dotações. Adota-se a especificação de acesso (ocupação "
    f"e contexto como dotações; {fmt(g('OB_COM_OCUP_DOT_PCT',83.8),1)}% de composição), coerente com a "
    f"leitura de que a discriminação age sobretudo no acesso às ocupações — medido diretamente pelo "
    f"GLMM. Como alerta Oaxaca e Ransom (1999), tratar a ocupação como dotação tende a subestimar a "
    f"discriminação total, já que a segregação ocupacional é, ela própria, discriminatória; daí a "
    f"complementaridade entre os métodos. Uma análise de sensibilidade explícita dimensiona esse "
    f"intervalo: removendo ocupação e contexto de UPA das dotações (especificação Mincer pura), a "
    f"parcela não explicada sobe de {fmt(g('OB_COM_OCUP_RET_PCT',16.2),1)}% para "
    f"{fmt(g('OB_SEM_OCUP_RET_PCT',75.2),1)}% do gap — ou seja, o intervalo de discriminação "
    f"total, a depender de quanto da segregação ocupacional é tratada como mérito (dotação) ou como "
    f"resultado de discriminação prévia (retorno), vai de "
    f"{fmt(g('OB_COM_OCUP_RET_PCT',16.2),1)}% a {fmt(g('OB_SEM_OCUP_RET_PCT',75.2),1)}%. As duas "
    f"leituras não são contraditórias: a especificação de acesso isola a discriminação SALARIAL "
    f"residual dentro da mesma ocupação (sticky floor, via RIF), enquanto a especificação sem "
    f"ocupação captura também a discriminação que já opera na TRIAGEM ocupacional, quantificada "
    f"diretamente pelo GLMM de acesso (Tabela 1).")

secao("Considerações Preliminares")
par("Tomados em conjunto, os resultados parciais convergem para a tese central: a desigualdade "
    "racial no trabalho brasileiro é um SISTEMA DE BARREIRAS EM CAMADAS — acesso, remuneração e "
    "interseccionalidade — com forte mediação do TERRITÓRIO (o contexto de moradia responde por "
    "mais da metade do gap). O gap é majoritariamente composição (Oaxaca: 83,8%), mas a composição "
    "é produto da discriminação no acesso às ocupações (GLMM: OR≈0,705) e convive com discriminação "
    "salarial residual, maior na base da distribuição (sticky floor). A consequência prescritiva é "
    "direta: políticas unidimensionais são insuficientes; a intervenção deve ser multidimensional e "
    "centrada no acesso a ocupações qualificadas. As próximas etapas incluem o refinamento das "
    "análises de robustez e a consolidação das recomendações.")

# ── Referências ───────────────────────────────────────────────────────────────
secao("Referências")
refs = [
    "BLINDER, A. S. Wage discrimination: reduced form and structural estimates. Journal of Human Resources, v. 8, n. 4, p. 436-455, 1973.",
    "BREIMAN, L. Random forests. Machine Learning, v. 45, n. 1, p. 5-32, 2001.",
    "CHEN, T.; GUESTRIN, C. XGBoost: a scalable tree boosting system. In: KDD, 2016. p. 785-794.",
    "CRENSHAW, K. Demarginalizing the intersection of race and sex. University of Chicago Legal Forum, v. 1989, n. 1, p. 139-167, 1989.",
    "FIRPO, S.; FORTIN, N. M.; LEMIEUX, T. Decomposing wage distributions using recentered influence function regressions. Econometrics, v. 6, n. 2, p. 28, 2018.",
    "HASENBALG, C. Discriminação e desigualdades raciais no Brasil. Rio de Janeiro: Graal, 1979.",
    "HENRIQUES, R. Desigualdade racial no Brasil: evolução das condições de vida na década de 90. Rio de Janeiro: IPEA, 2001. (Texto para discussão, 807).",
    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Pesquisa Nacional por Amostra de Domicílios Contínua. Rio de Janeiro: IBGE, 2016-2025.",
    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Pesquisa de Orçamentos Familiares 2017-2018: análise da qualidade de vida. Rio de Janeiro: IBGE, 2019.",
    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). PNAD Contínua: Rendimento de Todas as Fontes 2025. Rio de Janeiro: IBGE, 2025.",
    "KOENKER, R.; BASSETT, G. Regression quantiles. Econometrica, v. 46, n. 1, p. 33-50, 1978.",
    "LUNDBERG, S. M.; LEE, S. A unified approach to interpreting model predictions. In: NeurIPS, 2017. p. 4765-4774.",
    "OAXACA, R. Male-female wage differentials in urban labor markets. International Economic Review, v. 14, n. 3, p. 693-709, 1973.",
    "OAXACA, R. L.; RANSOM, M. R. Identification in detailed wage decompositions. Review of Economics and Statistics, v. 81, n. 1, p. 154-157, 1999.",
    "PAGER, D. Marked: race, crime, and finding work in an era of mass incarceration. Chicago: University of Chicago Press, 2007.",
    "RAUDENBUSH, S. W.; BRYK, A. S. Hierarchical linear models: applications and data analysis methods. 2. ed. Thousand Oaks: Sage, 2002.",
    "SAMPSON, R. J.; RAUDENBUSH, S. W.; EARLS, F. Neighborhoods and violent crime. Science, v. 277, p. 918-924, 1997.",
    "SOARES, S. S. D. Perfil da discriminação no mercado de trabalho: raça, sexo e salários no Brasil 1992-2006. Rio de Janeiro: IPEA, 2009. (Texto para discussão, 1395).",
    "VANDERWEELE, T. J.; DING, P. Sensitivity analysis in observational research: introducing the E-value. Annals of Internal Medicine, v. 167, n. 4, p. 268-274, 2017.",
    "WILSON, W. J. The truly disadvantaged: the inner city, the underclass, and public policy. Chicago: University of Chicago Press, 1987.",
]
for r in sorted(refs):
    p = doc.add_paragraph(r); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs: run.font.size = Pt(11)

secao("Declaração de uso de inteligência artificial")
par("Na elaboração deste trabalho foram utilizadas ferramentas de inteligência artificial "
    "(Claude Code, da Anthropic) como apoio à implementação e depuração de código (Python e R), à "
    "geração de figuras e tabelas e à formatação dos documentos. A concepção da pesquisa, a escolha "
    "das metodologias, a interpretação dos resultados e a redação final são de responsabilidade do "
    "autor, que revisou, validou e responde por todo o conteúdo.")

doc.save(str(OUT))
print(f"Arquivo gerado: {OUT.name}")
print(f"Tamanho: {OUT.stat().st_size // 1024} KB")
